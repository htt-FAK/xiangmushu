import io
import json
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

import config
from eval_judge import judge_payload
from eval_paths import API_RESULT
from server import app


PASS_THRESHOLD = 80


def _sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("自动评测模板", level=1)
    doc.add_paragraph("项目概述：{{PROJECT_OVERVIEW}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_api_ai_scenarios():
    client = TestClient(app)
    scenarios = []

    res = client.get("/api/template/list")
    assert res.status_code == 200
    data = res.json()
    assert isinstance(data.get("templates"), list)
    scenarios.append({"name": "GET /api/template/list", "ok": True, "detail": data})

    res = client.post(
        "/api/template/analyze",
        files={
            "file": (
                "ai_eval_anchor_template.docx",
                _sample_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    data = res.json()
    assert data["ok"] is True
    assert data["count"] >= 1
    assert data["mode"] == "anchor"
    assert data["tasks"][0]["location_hint"]["anchor"] == "{{PROJECT_OVERVIEW}}"
    scenarios.append(
        {
            "name": "POST /api/template/analyze",
            "ok": True,
            "count": data["count"],
            "mode": data["mode"],
        }
    )

    res = client.get("/api/kb/list")
    assert res.status_code == 200
    data = res.json()
    assert isinstance(data, list)
    assert all("slug" in item and "label" in item for item in data)
    scenarios.append({"name": "GET /api/kb/list", "ok": True, "count": len(data)})

    res = client.get("/api/kb/sources", params={"slug": "kb1"})
    assert res.status_code == 200
    data = res.json()
    assert isinstance(data.get("sources"), list)
    assert isinstance(data.get("chunk_count"), int)
    assert isinstance(data.get("source_count"), int)
    scenarios.append(
        {
            "name": "GET /api/kb/sources",
            "ok": True,
            "source_count": data["source_count"],
            "chunk_count": data["chunk_count"],
        }
    )

    res = client.get("/")
    assert res.status_code == 200
    assert "page-kb" in res.text
    assert "page-tpl" in res.text
    assert "page-gen" in res.text
    scenarios.append({"name": "GET /", "ok": True, "contains_spa_pages": True})

    res = client.post(
        "/api/template/analyze",
        files={
            "file": (
                "empty.docx",
                b"",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    data = res.json()
    assert data["ok"] is False
    assert "error" in data
    scenarios.append({"name": "POST /api/template/analyze (empty file)", "ok": True, "error": data["error"]})

    res = client.post(
        "/api/template/analyze",
        files={
            "file": (
                "not_a_docx.txt",
                b"This is plain text, not a docx file.",
                "text/plain",
            )
        },
    )
    assert res.status_code == 200
    data = res.json()
    assert data["ok"] is False
    assert "error" in data
    scenarios.append({"name": "POST /api/template/analyze (non-docx)", "ok": True, "error": data["error"]})

    res = client.post(
        "/api/template/analyze",
        files={
            "file": (
                "corrupted.docx",
                b"PK\x03\x04" + b"\x00" * 100,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    data = res.json()
    assert data["ok"] is False
    assert "error" in data
    scenarios.append({"name": "POST /api/template/analyze (corrupted docx)", "ok": True, "error": data["error"]})

    result = {
        "target": "api",
        "pass_threshold": PASS_THRESHOLD,
        "scenarios": scenarios,
        "functional_passed": True,
        "config": {
            "template_dir": str(Path(config.TEMPLATE_DIR).resolve()),
            "historical_dir": str(Path(config.HISTORICAL_DIR).resolve()),
        },
    }
    result["judge"] = judge_payload(result, target="api")
    API_RESULT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    assert API_RESULT.exists()
