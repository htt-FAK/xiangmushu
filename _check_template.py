from docx import Document

doc = Document(r'data\templates\智能体应用开发实践.docx')

print("=== 前15个段落 ===")
for i, p in enumerate(doc.paragraphs[:15]):
    print(f"[{i}] text={repr(p.text[:50])}")
    # 检查是否有图片/形状
    for j, r in enumerate(p.runs[:3]):
        draws = r._element.xpath(".//w:drawing")
        if draws:
            print(f"  run[{j}] 含 {len(draws)} 个 drawing")

print("\n=== 页眉检查 ===")
for i, s in enumerate(doc.sections):
    hdr = s.header
    for j, p in enumerate(hdr.paragraphs[:3]):
        if p.text.strip() or p.runs:
            print(f"Sec{i} header[{j}]={repr(p.text[:40])}")
            for r in p.runs[:2]:
                if r._element.xpath(".//w:drawing"):
                    print(f"  含 drawing")
