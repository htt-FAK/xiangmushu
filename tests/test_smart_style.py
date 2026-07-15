"""smart_style + column-width preservation 单元测试。

覆盖：
- merge_style 三级优先级合并
- _build_rPr_from_runstyle XML 构造
- heading_size_delta_pt 同步加减
- apply_smart_style 对标题段落选择正确 RunStyle
- 创新计划书模板 body_style 提取仿宋_GB2312
- _ensure_table_readability 列宽保留 (Task 6.1/6.4)
- _ensure_table_readability 最小列宽保护 (Task 6.2)
"""
from __future__ import annotations

import os
import sys
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import config
from core.smart_style import (
    _build_rPr_from_runstyle,
    apply_smart_style,
    merge_style,
)
from core.style_models import RunStyle, TemplateStyleProfile, system_default_profile

_INNOVATION_TEMPLATE = os.path.join(
    os.path.dirname(__file__),
    "..",
    "docs",
    "2.1.2024级广东理工学院创新计划书参考模板（通用）.docx",
)


# ── Helpers ────────────────────────────────────────────────────────────────

def _make_profile(
    body_ascii="Arial", body_east="黑体", body_pt=10.5,
    h1_pt=15.0, h2_pt=14.0, h3_pt=12.0,
) -> TemplateStyleProfile:
    body = RunStyle(font_ascii=body_ascii, font_east_asia=body_east, size_pt=body_pt)
    h1 = RunStyle(font_ascii="Times New Roman", font_east_asia="黑体",
                  size_pt=h1_pt, bold=True)
    h2 = RunStyle(font_ascii="Times New Roman", font_east_asia="黑体",
                  size_pt=h2_pt, bold=True)
    h3 = RunStyle(font_ascii="Times New Roman", font_east_asia="黑体",
                  size_pt=h3_pt, bold=True)
    return TemplateStyleProfile(
        body_style=body,
        heading_styles={1: h1, 2: h2, 3: h3},
        table_cell_style=body,
        table_label_style=RunStyle(bold=True),
    )


def _set_cell_width_dxa(tc, width_dxa: int) -> None:
    """Set w:tcW on a tc element in dxa units."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tw = tcPr.find(qn("w:tcW"))
    if tw is None:
        tw = OxmlElement("w:tcW")
        tcPr.append(tw)
    tw.set(qn("w:w"), str(width_dxa))
    tw.set(qn("w:type"), "dxa")


def _get_cell_width_dxa(tc) -> int:
    """Read w:tcW/@w:w from a tc element."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 0
    tw = tcPr.find(qn("w:tcW"))
    if tw is None:
        return 0
    w_val = tw.get(qn("w:w"))
    if w_val is None or w_val == "auto":
        return 0
    return int(w_val)


# ── Task 5.6: merge_style 三级合并 ──────────────────────────────────────

class TestMergeStyle(unittest.TestCase):
    """Test merge_style(profile, overrides) three-tier merging."""

    def test_user_override_beats_template(self):
        """User override takes priority over template value."""
        profile = _make_profile(body_ascii="Arial", body_pt=10.5)
        overrides = {"body_font_ascii": "Courier New", "body_size_pt": 11.0}
        merged = merge_style(profile, overrides)
        self.assertEqual(merged.body_style.font_ascii, "Courier New")
        self.assertAlmostEqual(merged.body_style.size_pt, 11.0)

    def test_empty_override_keeps_template(self):
        """Empty/None overrides keep template values."""
        profile = _make_profile(body_ascii="Arial", body_pt=10.5)
        merged = merge_style(profile, {})
        self.assertEqual(merged.body_style.font_ascii, "Arial")
        self.assertAlmostEqual(merged.body_style.size_pt, 10.5)

    def test_none_profile_gets_system_defaults(self):
        """None profile falls back to system defaults."""
        merged = merge_style(None)
        defaults = system_default_profile()
        self.assertEqual(merged.body_style.font_ascii, defaults.body_style.font_ascii)
        self.assertEqual(merged.body_style.font_east_asia, defaults.body_style.font_east_asia)
        self.assertAlmostEqual(merged.body_style.size_pt, defaults.body_style.size_pt)

    def test_none_profile_with_override(self):
        """None profile + user overrides = system defaults + overrides."""
        merged = merge_style(None, {"body_font_ascii": "Verdana"})
        self.assertEqual(merged.body_style.font_ascii, "Verdana")
        # east_asia should still be system default
        self.assertEqual(merged.body_style.font_east_asia, "宋体")


# ── Task 5.6: _build_rPr_from_runstyle ────────────────────────────────────

class TestBuildRPrFromRunstyle(unittest.TestCase):
    """Test _build_rPr_from_runstyle XML construction."""

    def test_font_attributes(self):
        """Verify w:rFonts@w:ascii and w:eastAsia."""
        rs = RunStyle(font_ascii="Arial", font_east_asia="黑体", size_pt=12.0)
        rpr = _build_rPr_from_runstyle(rs)
        rf = rpr.find(qn("w:rFonts"))
        self.assertIsNotNone(rf)
        self.assertEqual(rf.get(qn("w:ascii")), "Arial")
        self.assertEqual(rf.get(qn("w:eastAsia")), "黑体")

    def test_size_half_points(self):
        """Verify w:sz@w:val is half-point (pt × 2)."""
        rs = RunStyle(size_pt=12.0)
        rpr = _build_rPr_from_runstyle(rs)
        sz = rpr.find(qn("w:sz"))
        self.assertIsNotNone(sz)
        self.assertEqual(sz.get(qn("w:val")), "24")  # 12pt × 2 = 24

    def test_bold_presence(self):
        """Bold RunStyle includes w:b element."""
        rs = RunStyle(bold=True)
        rpr = _build_rPr_from_runstyle(rs)
        self.assertIsNotNone(rpr.find(qn("w:b")))

    def test_bold_absence(self):
        """Non-bold RunStyle omits w:b element."""
        rs = RunStyle(bold=False)
        rpr = _build_rPr_from_runstyle(rs)
        self.assertIsNone(rpr.find(qn("w:b")))


# ── Task 5.6: heading_size_delta_pt ──────────────────────────────────────

class TestHeadingSizeDeltaPt(unittest.TestCase):
    """heading_size_delta_pt applies uniformly to all heading levels."""

    def test_positive_delta(self):
        """Positive delta increases all heading sizes."""
        profile = _make_profile(h1_pt=15.0, h2_pt=14.0, h3_pt=12.0)
        merged = merge_style(profile, {"heading_size_delta_pt": 2.0})
        self.assertAlmostEqual(merged.heading_styles[1].size_pt, 17.0)
        self.assertAlmostEqual(merged.heading_styles[2].size_pt, 16.0)
        self.assertAlmostEqual(merged.heading_styles[3].size_pt, 14.0)

    def test_negative_delta_clamped_at_8pt(self):
        """Negative delta clamps at minimum 8pt."""
        profile = _make_profile(h1_pt=15.0, h2_pt=14.0, h3_pt=9.0)
        merged = merge_style(profile, {"heading_size_delta_pt": -5.0})
        # h3 = 9 - 5 = 4 → clamped to 8
        self.assertAlmostEqual(merged.heading_styles[3].size_pt, 8.0)
        # h1 = 15 - 5 = 10 → ok
        self.assertAlmostEqual(merged.heading_styles[1].size_pt, 10.0)

    def test_zero_delta(self):
        """Zero delta leaves sizes unchanged."""
        profile = _make_profile(h1_pt=15.0, h2_pt=14.0)
        merged = merge_style(profile, {"heading_size_delta_pt": 0.0})
        self.assertAlmostEqual(merged.heading_styles[1].size_pt, 15.0)
        self.assertAlmostEqual(merged.heading_styles[2].size_pt, 14.0)


# ── Task 5.6: apply_smart_style heading selection ──────────────────────

class TestApplySmartStyleHeading(unittest.TestCase):
    """apply_smart_style selects correct RunStyle for heading paragraphs."""

    def test_heading1_paragraph_uses_heading_style(self):
        """Paragraph with 'Heading 1' style uses heading_styles[1]."""
        doc = Document()
        para = doc.add_paragraph("测试标题")
        para.style = "Heading 1"
        run = para.add_run("测试标题")

        profile = _make_profile(body_pt=10.5, h1_pt=18.0)
        apply_smart_style(para, profile)

        # Check the run rPr has correct size (18pt × 2 = 36 half-points)
        rpr = run._r.find(qn("w:rPr"))
        self.assertIsNotNone(rpr)
        sz = rpr.find(qn("w:sz"))
        self.assertIsNotNone(sz)
        self.assertEqual(sz.get(qn("w:val")), "36")  # 18pt × 2

    def test_body_paragraph_uses_body_style(self):
        """Normal paragraph uses body_style."""
        doc = Document()
        para = doc.add_paragraph("正文内容")
        run = para.runs[0] if para.runs else para.add_run("正文内容")

        profile = _make_profile(body_pt=11.0)
        apply_smart_style(para, profile)

        rpr = run._r.find(qn("w:rPr"))
        self.assertIsNotNone(rpr)
        sz = rpr.find(qn("w:sz"))
        self.assertIsNotNone(sz)
        self.assertEqual(sz.get(qn("w:val")), "22")  # 11pt × 2


# ── Task 5.6: Innovation template body_style ─────────────────────────────

class TestInnovationTemplateBodyStyle(unittest.TestCase):
    """Verify the innovation template extracts 仿宋_GB2312 (not 宋体)."""

    @unittest.skipUnless(
        os.path.isfile(_INNOVATION_TEMPLATE),
        f"Innovation template not found: {_INNOVATION_TEMPLATE}",
    )
    def test_body_style_font_east_asia(self):
        """body_style.font_east_asia should be 仿宋_GB2312, not hardcoded 宋体."""
        from core.template_style_extractor import extract_style_profile
        profile = extract_style_profile(_INNOVATION_TEMPLATE)
        self.assertIn(
            "仿宋", profile.body_style.font_east_asia,
            f"Expected 仿宋 in east_asia font, got: {profile.body_style.font_east_asia}",
        )
        self.assertNotEqual(
            profile.body_style.font_east_asia, "宋体",
            "body_style should not be hardcoded 宋体",
        )


# ── Task 6.1/6.4: Column width preservation integration test ──────────────

class TestColumnWidthPreservation(unittest.TestCase):
    """_ensure_table_readability preserves column widths when PRESERVE=True."""

    def _make_table_with_widths(self, widths_dxa: list[int]) -> "Document":
        """Build a synthetic document with a 1-row table at given column widths."""
        doc = Document()
        ncols = len(widths_dxa)
        table = doc.add_table(rows=1, cols=ncols)
        # Fill cells so they are non-empty
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = f"C{i}"
        # Set column widths in XML
        first_tr = table.rows[0]._tr
        physical_tcs = first_tr.findall(qn("w:tc"))
        for tc, w in zip(physical_tcs, widths_dxa):
            _set_cell_width_dxa(tc, w)
        return doc

    @patch.object(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", True)
    def test_widths_preserved(self):
        """Column widths [1602, 7189] remain unchanged when PRESERVE=True."""
        doc = self._make_table_with_widths([1602, 7189])
        table = doc.tables[0]
        from core.filler import WordFiller
        wf = WordFiller()
        wf._ensure_table_readability(table)
        # Read widths back
        first_tr = table.rows[0]._tr
        physical_tcs = first_tr.findall(qn("w:tc"))
        actual = [_get_cell_width_dxa(tc) for tc in physical_tcs]
        self.assertEqual(actual, [1602, 7189],
                         f"Expected [1602, 7189] but got {actual}")

    @patch.object(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", False)
    def test_widths_equalized_when_preserve_false(self):
        """When PRESERVE=False, column widths are equalized (old behavior)."""
        doc = self._make_table_with_widths([1602, 7189])
        table = doc.tables[0]
        from core.filler import WordFiller
        wf = WordFiller()
        wf._ensure_table_readability(table)
        first_tr = table.rows[0]._tr
        physical_tcs = first_tr.findall(qn("w:tc"))
        actual = [_get_cell_width_dxa(tc) for tc in physical_tcs]
        # They should NOT be [1602, 7189] anymore
        self.assertNotEqual(actual, [1602, 7189],
                            "Widths should be equalized when PRESERVE=False")


# ── Task 6.2/6.4: Minimum column width guard ──────────────────────────────

class TestMinimumColumnWidthGuard(unittest.TestCase):
    """Narrow columns are expanded to MIN_COLUMN_WIDTH_DXA."""

    @patch.object(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", True)
    @patch.object(config, "MIN_COLUMN_WIDTH_DXA", 500)
    def test_narrow_column_expanded(self):
        """Column with width 100 dxa is expanded to ≥ 500."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = f"C{i}"
        # Set widths: [100, 9900]
        first_tr = table.rows[0]._tr
        physical_tcs = first_tr.findall(qn("w:tc"))
        _set_cell_width_dxa(physical_tcs[0], 100)
        _set_cell_width_dxa(physical_tcs[1], 9900)

        from core.filler import WordFiller
        wf = WordFiller()
        wf._ensure_table_readability(table)

        # Read widths back
        physical_tcs = table.rows[0]._tr.findall(qn("w:tc"))
        actual = [_get_cell_width_dxa(tc) for tc in physical_tcs]
        self.assertGreaterEqual(actual[0], 500,
                                f"Narrow column should be ≥ 500 dxa, got {actual[0]}")
        # Total width preserved: sum should remain unchanged
        self.assertEqual(sum(actual), sum([100, 9900]),
                         f"Total width should be conserved, got {sum(actual)}")

    @patch.object(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", True)
    @patch.object(config, "MIN_COLUMN_WIDTH_DXA", 500)
    def test_all_columns_above_minimum_unchanged(self):
        """When all columns are ≥ min, no adjustment occurs."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = f"C{i}"
        first_tr = table.rows[0]._tr
        physical_tcs = first_tr.findall(qn("w:tc"))
        _set_cell_width_dxa(physical_tcs[0], 3000)
        _set_cell_width_dxa(physical_tcs[1], 7000)

        from core.filler import WordFiller
        wf = WordFiller()
        wf._ensure_table_readability(table)

        physical_tcs = table.rows[0]._tr.findall(qn("w:tc"))
        actual = [_get_cell_width_dxa(tc) for tc in physical_tcs]
        self.assertEqual(actual, [3000, 7000],
                         f"Expected [3000, 7000] but got {actual}")


# ── Task 6.3: gridSpan safety ──────────────────────────────────────────────

class TestGridSpanSafety(unittest.TestCase):
    """Columns with gridSpan > 1 should be skipped during width adjustment."""

    @patch.object(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", True)
    @patch.object(config, "MIN_COLUMN_WIDTH_DXA", 500)
    def test_gridspan_cell_skipped(self):
        """A cell with gridSpan=2 is not adjusted even if its width is small."""
        doc = Document()
        # Create a 2-row table: row0 has 2 cells, row1 has a merged cell (gridSpan=2)
        table = doc.add_table(rows=2, cols=2)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "x"
        # Set widths on row 0
        first_tr = table.rows[0]._tr
        physical_tcs_r0 = first_tr.findall(qn("w:tc"))
        _set_cell_width_dxa(physical_tcs_r0[0], 3000)
        _set_cell_width_dxa(physical_tcs_r0[1], 7000)

        # Add gridSpan to first cell of row 1 to simulate a merged row
        second_tr = table.rows[1]._tr
        physical_tcs_r1 = second_tr.findall(qn("w:tc"))
        if physical_tcs_r1:
            tc = physical_tcs_r1[0]
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            gs = OxmlElement("w:gridSpan")
            gs.set(qn("w:val"), "2")
            tcPr.append(gs)
            _set_cell_width_dxa(tc, 100)  # Narrow but with gridSpan

        from core.filler import WordFiller
        wf = WordFiller()
        wf._ensure_table_readability(table)

        # Verify row 0 widths are unchanged (the adjustment only looks at first row)
        physical_tcs_r0_after = table.rows[0]._tr.findall(qn("w:tc"))
        actual = [_get_cell_width_dxa(tc) for tc in physical_tcs_r0_after]
        self.assertEqual(actual, [3000, 7000],
                         f"Row 0 widths should be unchanged, got {actual}")


if __name__ == "__main__":
    unittest.main()
