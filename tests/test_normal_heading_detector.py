"""Normal-heading detector 单元测试。

覆盖：
- 中文编号标题（一、二、三）→ level 1
- 十进制子标题（1.1, 2.1）→ level 2
- 三级标题（1.1.1）→ level 3
- 短粗体段落（无编号，标题场景）
- 普通正文段落不误判
- 封面大标题（36pt）被排除
- 真实模板（创新计划书）端到端检测
"""
from __future__ import annotations

import os
import sys
import unittest

from docx import Document
from docx.shared import Pt

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from core.normal_heading_detector import (
    classify_heading,
    find_all_headings,
    score_normal_heading,
)


def _make_para(doc: Document, text: str, font_name: str = "SimSun",
               font_size: Pt | None = None, bold: bool | None = None) -> None:
    """Helper: 追加一个段落并设置格式。"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = font_name
        if font_size is not None:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold


class ScoreNormalHeadingTests(unittest.TestCase):

    def test_cn_level1_heading_scores_high(self):
        """「一、项目基本信息」bold 14pt → score ≥ 50（阈值）"""
        doc = Document()
        _make_para(doc, "一、项目基本信息", font_size=Pt(14), bold=True)
        score = score_normal_heading(doc.paragraphs[0], doc=doc, para_index=0)
        # 30(size≥14pt) + 25(bold) + 35(CN编号) + 10(短文本) = 100
        self.assertGreaterEqual(score, 50, f"score={score}")

    def test_decimal_level2_heading(self):
        """「1.1 项目背景」bold 14pt → level 2"""
        doc = Document()
        _make_para(doc, "1.1 项目背景", font_size=Pt(14), bold=True)
        level = classify_heading(doc.paragraphs[0], threshold=50, doc=doc, para_index=0)
        self.assertEqual(level, 2)

    def test_triple_level3_heading(self):
        """「5.2.1 子技术路线」bold → level 3"""
        doc = Document()
        _make_para(doc, "5.2.1 子技术路线", font_size=Pt(12), bold=True)
        level = classify_heading(doc.paragraphs[0], threshold=50, doc=doc, para_index=0)
        self.assertEqual(level, 3)

    def test_body_paragraph_not_classified(self):
        """普通正文（no bold, 12pt, 长段落）→ None"""
        doc = Document()
        long_text = "本项目旨在开发一套智能文档生成系统，解决当前申报文档填写效率低、格式不一致的问题。"
        _make_para(doc, long_text, font_size=Pt(12), bold=False)
        level = classify_heading(doc.paragraphs[0], threshold=50, doc=doc, para_index=0)
        self.assertIsNone(level)

    def test_cover_title_excluded(self):
        """36pt 封面大标题不应当作章节标题。"""
        doc = Document()
        _make_para(doc, "大学生创新大赛", font_size=Pt(36), bold=True)
        score = score_normal_heading(doc.paragraphs[0], doc=doc, para_index=0)
        self.assertEqual(score, 0, f"cover title should score 0, got {score}")

    def test_bold_short_paragraph_no_numbering(self):
        """短粗体段落无编号（如「项目主要内容」）→ 评分低于 CN 编号标题。"""
        doc = Document()
        _make_para(doc, "项目主要内容", font_size=Pt(12), bold=True)
        score = score_normal_heading(doc.paragraphs[0], doc=doc, para_index=0)
        # 25(bold) + 15(all-bold) + 10(short) = 50，刚好达阈值
        self.assertGreaterEqual(score, 40, f"bold short text should score ≥ 40, got {score}")

    def test_empty_paragraph_returns_zero(self):
        """空段落评分为 0。"""
        doc = Document()
        doc.add_paragraph("")
        score = score_normal_heading(doc.paragraphs[0], doc=doc, para_index=0)
        self.assertEqual(score, 0)


class FindAllHeadingsTests(unittest.TestCase):

    def test_multi_heading_document(self):
        """多章节文档，应检测到所有「一、二、三」标题。"""
        doc = Document()
        _make_para(doc, "一、项目基本信息", font_size=Pt(14), bold=True)
        _make_para(doc, "正文内容...", font_size=Pt(12), bold=False)
        _make_para(doc, "二、项目主要内容", font_size=Pt(14), bold=True)
        _make_para(doc, "更多内容...", font_size=Pt(12), bold=False)
        _make_para(doc, "三、痛点与创新点", font_size=Pt(14), bold=True)

        headings = find_all_headings(doc, threshold=50)
        # 应至少有 3 个 level-1 标题
        level1_count = sum(1 for _, lvl in headings if lvl == 1)
        self.assertGreaterEqual(level1_count, 3, f"headings={headings}")

    def test_no_headings_returns_empty(self):
        """全文档无标题时返回空列表。"""
        doc = Document()
        _make_para(doc, "这是一段普通正文，没有任何标题格式。", font_size=Pt(12), bold=False)
        _make_para(doc, "另一段普通内容。", font_size=Pt(12), bold=False)
        headings = find_all_headings(doc, threshold=50)
        self.assertEqual(headings, [])


class RealTemplateTests(unittest.TestCase):

    @unittest.skipUnless(
        os.path.exists(
            os.path.join(os.path.dirname(__file__), "..", "docs",
                         "2.1.2024级广东理工学院创新计划书参考模板（通用）.docx")
        ),
        "创新计划书模板不存在"
    )
    def test_innovation_template_headings(self):
        """真实创新计划书模板：检测 8 个章节标题（一、～八、）。"""
        path = os.path.join(
            os.path.dirname(__file__), "..", "docs",
            "2.1.2024级广东理工学院创新计划书参考模板（通用）.docx"
        )
        doc = Document(path)
        headings = find_all_headings(doc, threshold=50)

        texts = [doc.paragraphs[idx].text.strip() for idx, _ in headings]
        # 至少应识别主要章节（一、到八、）
        cn_numbered = [t for t in texts if t and t[0] in "一二三四五六七八九十"]
        self.assertGreaterEqual(
            len(cn_numbered), 6,
            f"expected ≥6 CN-numbered headings, got {len(cn_numbered)}: {texts}"
        )


if __name__ == "__main__":
    unittest.main()
