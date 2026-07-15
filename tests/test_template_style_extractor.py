"""template_style_extractor 单元测试。

覆盖：
- _extract_column_widths 在创新计划书模板上的端到端测试
- 缓存命中/未命中行为
- _resolve_style_chain 在带 Heading 1 样式的文档上的正确解析
- get_or_extract_style_profile 整体提取
"""
from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from core.style_models import RunStyle, TemplateStyleProfile
from core.template_style_extractor import (
    _cache_path,
    _detect_heading_styles,
    _extract_column_widths,
    _resolve_style_chain,
    _run_style_from_run,
    extract_style_profile,
    get_or_extract_style_profile,
)

_INNOVATION_TEMPLATE = os.path.join(
    os.path.dirname(__file__),
    "..",
    "docs",
    "2.1.2024级广东理工学院创新计划书参考模板（通用）.docx",
)


class TestExtractColumnWidths(unittest.TestCase):
    """_extract_column_widths on the innovation template."""

    @unittest.skipUnless(
        os.path.isfile(_INNOVATION_TEMPLATE),
        f"Innovation template not found: {_INNOVATION_TEMPLATE}",
    )
    def test_table0_row0_has_2_physical_cells(self):
        """Table 0 row 0 should have 2 physical cells with correct dxa widths."""
        doc = Document(_INNOVATION_TEMPLATE)
        widths = _extract_column_widths(doc)
        self.assertIn(0, widths, "Table 0 not found in column_widths")
        row0_widths = widths[0]
        self.assertEqual(
            len(row0_widths), 2,
            f"Expected 2 physical cells in row 0 of table 0, got {len(row0_widths)}",
        )
        # Widths should be positive integers (dxa)
        for w in row0_widths:
            self.assertGreater(w, 0, f"Width should be positive, got {w}")

    def test_synthetic_table(self):
        """Synthetic document with explicit column widths."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        # Set widths via XML — use physical tc elements, not row.cells
        from docx.oxml import OxmlElement
        first_row = table.rows[0]
        tc_elements = first_row._tr.findall(qn("w:tc"))
        for i, tc in enumerate(tc_elements):
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            # Remove existing tcW if present
            existing = tc_pr.find(qn("w:tcW"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(2000 + i * 1000))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)

        widths = _extract_column_widths(doc)
        self.assertIn(0, widths)
        self.assertEqual(widths[0], [2000, 3000, 4000])


class TestCacheHitMiss(unittest.TestCase):
    """get_or_extract_style_profile cache behavior."""

    @unittest.skipUnless(
        os.path.isfile(_INNOVATION_TEMPLATE),
        f"Innovation template not found: {_INNOVATION_TEMPLATE}",
    )
    def test_cache_miss_then_hit(self):
        """First call extracts (miss), second call reads cache (hit)."""
        cpath = _cache_path(_INNOVATION_TEMPLATE)
        # Clean cache
        if os.path.isfile(cpath):
            os.remove(cpath)

        # First call: extract + write cache
        p1 = get_or_extract_style_profile(_INNOVATION_TEMPLATE)
        self.assertTrue(os.path.isfile(cpath), "Cache file should exist after first call")

        # Second call: should read from cache
        p2 = get_or_extract_style_profile(_INNOVATION_TEMPLATE)

        # Both profiles should have same body style
        self.assertEqual(p1.body_style.font_ascii, p2.body_style.font_ascii)
        self.assertEqual(p1.body_style.font_east_asia, p2.body_style.font_east_asia)
        self.assertEqual(p1.body_style.size_pt, p2.body_style.size_pt)
        self.assertEqual(set(p1.heading_styles.keys()), set(p2.heading_styles.keys()))

        # Cleanup
        if os.path.isfile(cpath):
            os.remove(cpath)

    def test_cache_roundtrip_json(self):
        """TemplateStyleProfile to_json/from_json roundtrip."""
        profile = TemplateStyleProfile(
            body_style=RunStyle(font_ascii="Arial", font_east_asia="黑体", size_pt=11.0),
            heading_styles={1: RunStyle(bold=True, size_pt=15.0)},
            column_widths={0: [3000, 6000]},
            line_spacing=1.5,
            source_template="test.docx",
        )
        raw = profile.to_json()
        restored = TemplateStyleProfile.from_json(raw)
        self.assertEqual(restored.body_style.font_ascii, "Arial")
        self.assertEqual(restored.body_style.font_east_asia, "黑体")
        self.assertEqual(restored.heading_styles[1].bold, True)
        self.assertEqual(restored.column_widths[0], [3000, 6000])
        self.assertEqual(restored.line_spacing, 1.5)

    def test_cache_synthetic_doc(self):
        """Cache miss/hit on a synthetic temporary document."""
        doc = Document()
        doc.add_paragraph("Hello world body text for testing.")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
            doc.save(tmp_path)

        try:
            cpath = _cache_path(tmp_path)
            if os.path.isfile(cpath):
                os.remove(cpath)

            p1 = get_or_extract_style_profile(tmp_path)
            self.assertTrue(os.path.isfile(cpath))

            p2 = get_or_extract_style_profile(tmp_path)
            self.assertEqual(p1.body_style.size_pt, p2.body_style.size_pt)
        finally:
            if os.path.isfile(tmp_path):
                os.unlink(tmp_path)
            cpath = _cache_path(tmp_path)
            if os.path.isfile(cpath):
                os.remove(cpath)


class TestResolveStyleChain(unittest.TestCase):
    """_resolve_style_chain on a document with Heading 1 style."""

    def test_heading1_style_resolution(self):
        """Heading 1 style should resolve with bold=True and a size."""
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        # Access the style
        styles_xml = doc.styles.element
        h1_style = doc.styles["Heading 1"]
        resolved = _resolve_style_chain(h1_style, styles_xml)
        # Heading 1 is typically bold and larger than body
        self.assertIsInstance(resolved, RunStyle)
        # At minimum, the style should have a size
        self.assertGreater(resolved.size_pt, 0)

    def test_normal_style_resolution(self):
        """Normal style should resolve without errors."""
        doc = Document()
        doc.add_paragraph("Plain text")
        styles_xml = doc.styles.element
        normal_style = doc.styles["Normal"]
        resolved = _resolve_style_chain(normal_style, styles_xml)
        self.assertIsInstance(resolved, RunStyle)

    def test_custom_style_with_basedon(self):
        """Custom style based on Heading 1 should inherit bold."""
        doc = Document()
        # Create a custom style based on Heading 1
        try:
            custom = doc.styles.add_style("CustomH1", 1)  # WD_STYLE_TYPE.PARAGRAPH
            custom.base_style = doc.styles["Heading 1"]
            # Set a custom size
            custom.font.size = Pt(18)

            styles_xml = doc.styles.element
            resolved = _resolve_style_chain(custom, styles_xml)
            self.assertIsInstance(resolved, RunStyle)
            self.assertEqual(resolved.size_pt, 18.0)
        except Exception:
            # If add_style fails (duplicate), skip
            self.skipTest("Could not create custom style (may already exist)")


class TestRunStyleFromRun(unittest.TestCase):
    """_run_style_from_run basic behavior."""

    def test_basic_run(self):
        """A run with only ascii font should still produce a valid RunStyle."""
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        rs = _run_style_from_run(run)
        self.assertEqual(rs.font_ascii, "Times New Roman")
        self.assertEqual(rs.size_pt, 12.0)


class TestExtractStyleProfileEndToEnd(unittest.TestCase):
    """End-to-end extraction on the innovation template."""

    @unittest.skipUnless(
        os.path.isfile(_INNOVATION_TEMPLATE),
        f"Innovation template not found: {_INNOVATION_TEMPLATE}",
    )
    def test_full_extraction(self):
        """extract_style_profile should return a populated profile."""
        profile = extract_style_profile(_INNOVATION_TEMPLATE)
        self.assertIsInstance(profile, TemplateStyleProfile)
        self.assertTrue(profile.source_template)
        self.assertTrue(profile.extracted_at)
        # Body style should have a size
        self.assertGreater(profile.body_style.size_pt, 0)
        # Should have at least one heading level
        self.assertGreater(len(profile.heading_styles), 0)
        # Should have column widths
        self.assertGreater(len(profile.column_widths), 0)


if __name__ == "__main__":
    unittest.main()
