"""document_type_detector + chapter_path_builder 单元测试。"""
from __future__ import annotations

import os
import sys
import unittest

from docx import Document
from docx.shared import Pt

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from core.document_type_detector import format_document_type_block, infer_document_type
from core.chapter_path_builder import build_chapter_path, find_target_heading_index


class InferDocumentTypeTests(unittest.TestCase):

    def test_explicit_type_takes_precedence(self):
        """explicit_type 非空时直接返回，不做推断。"""
        result = infer_document_type(
            "任何文件名.docx", explicit_type="国家自然科学基金申请书"
        )
        self.assertEqual(result, "国家自然科学基金申请书")

    def test_innovation_plan_from_filename(self):
        """创新计划书 → '创新创业计划书'"""
        result = infer_document_type("2.1.创新计划书参考模板.docx")
        self.assertEqual(result, "创新创业计划书")

    def test_innovation_entrepreneurship(self):
        """创新创业大赛模板 → '大学生创新创业项目计划书'"""
        result = infer_document_type("大学生创新创业大赛模板.docx")
        self.assertEqual(result, "大学生创新创业项目计划书")

    def test_course_report(self):
        """结课报告 → '课程实验/结课报告'"""
        result = infer_document_type("结课报告模板.docx")
        self.assertEqual(result, "课程实验/结课报告")

    def test_agent_practice(self):
        """智能体应用开发实践模板"""
        result = infer_document_type("智能体应用开发实践.docx")
        self.assertEqual(result, "智能体应用开发实践报告")

    def test_project_declaration(self):
        """项目申报书模板"""
        result = infer_document_type("项目申报书模板.docx")
        self.assertEqual(result, "项目申报书")

    def test_unknown_filename_fallback(self):
        """未匹配任何模式时返回兜底文案。"""
        result = infer_document_type("random_template_123.docx")
        self.assertEqual(result, "项目申报文档（通用）")

    def test_format_block(self):
        """format_document_type_block 返回带 emoji 的注入字符串。"""
        block = format_document_type_block("创新创业计划书")
        self.assertEqual(block, "⚠️ 文档类型：创新创业计划书")

    def test_format_block_empty(self):
        """空字符串返回空结果。"""
        self.assertEqual(format_document_type_block(""), "")


class FindTargetHeadingIndexTests(unittest.TestCase):

    def _make_doc_with_headings(self) -> Document:
        doc = Document()
        p1 = doc.add_paragraph("一、项目基本信息")
        for run in p1.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        doc.add_paragraph("一些正文...")
        p2 = doc.add_paragraph("二、项目主要内容")
        for run in p2.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        doc.add_paragraph("更多内容...")
        p3 = doc.add_paragraph("三、痛点与创新点")
        for run in p3.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        return doc

    def test_find_existing_chapter(self):
        """能找到存在的章节标题。"""
        doc = self._make_doc_with_headings()
        from core.normal_heading_detector import find_all_headings
        headings = find_all_headings(doc, threshold=50)
        idx = find_target_heading_index("二、项目主要内容", headings, doc)
        self.assertIsNotNone(idx)
        self.assertEqual(idx, 1)  # 第二个标题

    def test_find_with_normalization(self):
        """模糊匹配：target 是子串。"""
        doc = self._make_doc_with_headings()
        from core.normal_heading_detector import find_all_headings
        headings = find_all_headings(doc, threshold=50)
        idx = find_target_heading_index("项目基本信息", headings, doc)
        self.assertIsNotNone(idx)

    def test_not_found_returns_none(self):
        """target 不在 headings 中返回 None。"""
        doc = self._make_doc_with_headings()
        from core.normal_heading_detector import find_all_headings
        headings = find_all_headings(doc, threshold=50)
        idx = find_target_heading_index("七、不存在的章节", headings, doc)
        self.assertIsNone(idx)


class BuildChapterPathTests(unittest.TestCase):

    def test_top_level_chapter(self):
        """顶级章节返回原样的标题。"""
        doc = Document()
        p = doc.add_paragraph("一、项目基本信息")
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        doc.add_paragraph("一些正文...")

        from core.normal_heading_detector import find_all_headings
        headings = find_all_headings(doc, threshold=50)
        path = build_chapter_path("一、项目基本信息", doc, headings)
        # 顶级章节应返回标题本身（或含其文本）
        self.assertIn("项目基本信息", path)


if __name__ == "__main__":
    unittest.main()
