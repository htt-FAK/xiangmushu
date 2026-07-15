"""Normal-Style 章节标题检测器。

解决 80%+ 真实申报模板（创新计划书等）的段落全部为 Normal style 导致
`heading_level_from_style()` 返回 None、章节边界无法识别的问题。

核心方法：
    score_normal_heading(para, doc=None, para_index=None) -> int
        7 个加权信号融合评分

    classify_heading(para, threshold, doc=None, para_index=None) -> Optional[int]
        评分 >= threshold 返回 heading level (1-3)，否则 None

    find_all_headings(doc, threshold) -> list[tuple[int, int]]
        全文档扫描，返回 [(para_index, level), ...]

信号设计（加权评分，满分 155）：
    字号 ≥ 14pt           +30
    加粗                   +25
    中文编号（一、二、…）   +35  (level 1)
    十进制编号（1.1, 2.1） +30  (level 2)
    三级编号（1.1.1）      +25  (level 3)
    位于表格前             +10
    全段粗体               +15
    style 含"标题/Heading" +40  (兜底)
    短文本 ≤ 30 字         +10

level 推断逻辑：
    中文编号 "一、二、…、十"  → level 1
    小数编号 "1.1, 2.3"      → level 2
    三级编号 "1.1.1"         → level 3
    仅 bold + short + before_table → level 1（子标题场景）
"""
from __future__ import annotations

import logging
import re
from typing import Optional, TYPE_CHECKING

from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx import Document

_LOG = logging.getLogger(__name__)

# ── 信号权重 ──────────────────────────────────────────────────────────────────
_WEIGHT_SIZE_GE_14PT = 30
_WEIGHT_BOLD = 25
_WEIGHT_CN_NUMBERING_L1 = 35  # "一、二、三..."
_WEIGHT_DECIMAL_NUMBERING_L2 = 30  # "1.1, 2.3"
_WEIGHT_TRIPLE_NUMBERING_L3 = 25  # "1.1.1"
_WEIGHT_BEFORE_TABLE = 10
_WEIGHT_ALL_BOLD = 15
_WEIGHT_STYLE_HEADING = 40
_WEIGHT_SHORT_TEXT = 10

_SHORT_TEXT_MAX_CHARS = 30

# ── 编号模式 ──────────────────────────────────────────────────────────────────
_CN_L1_RE = re.compile(
    r"^[一二三四五六七八九十]+[、.．:：\s]"
)
_DECIMAL_L2_RE = re.compile(r"^\d+\.\d+(?:\s|$|[^.\d])")
_TRIPLE_L3_RE = re.compile(r"^\d+\.\d+\.\d+(?:\s|$)")

# ── 正文排除（字号过大的封面标题）─────────────────────────────────────────────
# 36pt = 457200 EMU，封面大标题不应被识别为章节标题
_COVER_TITLE_SIZE_THRESHOLD_EMU = 400000  # ~31.5pt


def _paragraph_element_index(para: Paragraph) -> int:
    """返回 para 在 body 中的位置（用于判断 next-sibling 是否是表格）。"""
    try:
        parent = para._p.getparent()
        if parent is None:
            return -1
        return list(parent).index(para._p)
    except (ValueError, AttributeError):
        return -1


def _next_sibling_is_table(doc: "Document | None", para: Paragraph) -> bool:
    """判断段落的下一个兄弟元素是否是 w:tbl。"""
    if doc is None:
        return False
    try:
        from docx.oxml.ns import qn
        idx = _paragraph_element_index(para)
        if idx < 0:
            return False
        body = doc.element.body
        children = list(body)
        if idx + 1 < len(children):
            return children[idx + 1].tag == qn("w:tbl")
    except Exception:
        pass
    return False


def _text_is_cover_title(size_emu: int | None, text: str) -> bool:
    """检测超大字号封面标题（不应当作章节标题）。"""
    if size_emu is None:
        return False
    return size_emu >= _COVER_TITLE_SIZE_THRESHOLD_EMU


def _infer_heading_level(text: str) -> int:
    """根据编号模式推断 heading level：1-3，无编号返回 1。"""
    t = text.strip()
    if _TRIPLE_L3_RE.match(t):
        return 3
    if _DECIMAL_L2_RE.match(t):
        return 2
    return 1  # CN numbering or plain bold sub-heading


def score_normal_heading(
    para: Paragraph,
    doc: "Document | None" = None,
    para_index: int | None = None,
) -> int:
    """7 个加权信号融合评分，返回 0–155 的整数。"""
    score = 0
    text = (para.text or "").strip()
    if not text:
        return 0

    # 排除超大封面标题
    first_size = para.runs[0].font.size if para.runs else None
    first_size_emu = int(first_size) if first_size is not None else None
    if _text_is_cover_title(first_size_emu, text):
        return 0

    # Signal 1: 字号 ≥ 14pt（14pt = 177800 EMU）
    if first_size is not None and int(first_size) >= 177800:
        score += _WEIGHT_SIZE_GE_14PT

    # Signal 2: 加粗
    has_bold = any(r.font.bold for r in para.runs if r.font.bold is not None)
    if has_bold:
        score += _WEIGHT_BOLD

    # Signal 3a: 中文编号（一、二、三...）→ level 1
    if _CN_L1_RE.match(text):
        score += _WEIGHT_CN_NUMBERING_L1

    # Signal 3b: 十进制编号（1.1, 2.3）→ level 2
    elif _DECIMAL_L2_RE.match(text):
        score += _WEIGHT_DECIMAL_NUMBERING_L2

    # Signal 3c: 三级编号（1.1.1）→ level 3
    elif _TRIPLE_L3_RE.match(text):
        score += _WEIGHT_TRIPLE_NUMBERING_L3

    # Signal 4: 位于表格前
    if _next_sibling_is_table(doc, para):
        score += _WEIGHT_BEFORE_TABLE

    # Signal 5: 全段粗体（非混排）
    if para.runs and all(
        r.font.bold is True for r in para.runs if (r.text or "").strip()
    ):
        score += _WEIGHT_ALL_BOLD

    # Signal 6: style name 含 "标题" 或 "Heading"
    style_name = (para.style.name or "") if para.style else ""
    if re.search(r"(标题|heading|title)", style_name, re.IGNORECASE):
        score += _WEIGHT_STYLE_HEADING

    # Signal 7: 短文本 ≤ 30 字
    if len(text) <= _SHORT_TEXT_MAX_CHARS:
        score += _WEIGHT_SHORT_TEXT

    return score


def classify_heading(
    para: Paragraph,
    threshold: int = 50,
    doc: "Document | None" = None,
    para_index: int | None = None,
) -> Optional[int]:
    """评分 >= threshold 返回 heading level (1-3)，否则 None。"""
    score = score_normal_heading(para, doc=doc, para_index=para_index)
    if score < threshold:
        return None
    text = (para.text or "").strip()
    level = _infer_heading_level(text)
    _LOG.debug(
        "normal_heading: score=%d level=%d text=%r",
        score, level, text[:40],
    )
    return level


def find_all_headings(
    doc: "Document",
    threshold: int = 50,
) -> list[tuple[int, int]]:
    """遍历文档所有段落，返回 [(para_index, level), ...]，按出现顺序。"""
    results: list[tuple[int, int]] = []
    for i, para in enumerate(doc.paragraphs):
        level = classify_heading(para, threshold, doc=doc, para_index=i)
        if level is not None:
            results.append((i, level))
    _LOG.info(
        "find_all_headings: found %d headings (threshold=%d)", len(results), threshold
    )
    return results


def find_heading_boundaries(
    doc: "Document",
    threshold: int = 50,
) -> dict[str, tuple[int, Optional[int]]]:
    """构建章节标题文本到段落索引区间的映射。

    返回 {chapter_text: (start_para_index, end_para_index_exclusive)}

    用于 `_collect_chapter_region()` 的 fallback。
    end_para_index_exclusive 为 None 表示延伸到文档结尾。
    """
    headings = find_all_headings(doc, threshold)
    if not headings:
        return {}

    n_paras = len(doc.paragraphs)
    boundaries: dict[str, tuple[int, Optional[int]]] = {}
    for idx, (para_idx, level) in enumerate(headings):
        chapter_text = (doc.paragraphs[para_idx].text or "").strip()
        end_idx: Optional[int] = None
        # 下一同级或更高级标题的位置
        for next_idx, (next_para_idx, next_level) in enumerate(headings[idx + 1:], idx + 1):
            if next_level <= level:
                end_idx = next_para_idx
                break
        boundaries[chapter_text] = (para_idx, end_idx if end_idx is not None else n_paras)

    return boundaries
