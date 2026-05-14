"""表格行批量生成：将同行的多个 table_cell 任务合并为一次 LLM 调用，输出 JSON。

降级策略：解析失败或 cell 数量超过上限时，自动回退到逐个生成。
"""
from __future__ import annotations

import json
import logging
import os
import re
from typing import Any, Dict, List, Optional

from openai import OpenAI

import config
from core.dashscope_chat import chat_completions_create
from core.evidence_planner import Evidence, compress_evidence
from core.fill_task import FillTask
from core.generator import SYSTEM_PROMPT
from core.table_context import get_column_header_text_from_table

_LOG = logging.getLogger(__name__)

BATCH_MAX_CELLS = int(config.__dict__.get("BATCH_MAX_CELLS", 8))
# 列数过多时一行 JSON 易混淆，降级逐格生成
BATCH_TABLE_MAX_COLS = int(os.getenv("BATCH_TABLE_MAX_COLS", "5"))

_BATCH_SYSTEM = SYSTEM_PROMPT + "\n输出纯 JSON 对象，键为单元格序号（从 0 开始），值为简短填写内容，不输出 JSON 以外内容。"


def _strip_json_fence(raw: str) -> str:
    s = (raw or "").strip()
    if s.startswith("```"):
        parts = s.split("```")
        if len(parts) >= 2:
            s = parts[1]
            if s.lstrip().startswith("json"):
                s = s.lstrip()[4:].lstrip()
    return s.strip()


def batch_generate_table_row(
    client: OpenAI,
    tasks: List[FillTask],
    evidence: Evidence,
    table_context: Optional[str] = None,
    enable_web: bool = False,
    template_path: Optional[str] = None,
) -> Optional[Dict[int, str]]:
    """
    批量生成同一表格行的所有单元格。

    返回 {cell_index: content} 字典；
    若无法批量（超过上限、证据不足等）返回 None（触发降级）。
    """
    if not tasks:
        return {}
    if len(tasks) > BATCH_MAX_CELLS:
        _LOG.info("batch_table_row: too many cells (%d), skip batch", len(tasks))
        return None
    if evidence.weak_kb and not enable_web:
        _LOG.info("batch_table_row: weak_kb without web, skip batch")
        return None

    doc = None
    ncols = 0
    tbl_idx = 0
    if template_path:
        try:
            from docx import Document as DocCls

            doc = DocCls(template_path)
            loc0 = tasks[0].location_hint or {}
            tbl_idx = int(loc0.get("table_index", 0))
            if 0 <= tbl_idx < len(doc.tables):
                t0 = doc.tables[tbl_idx]
                if t0.rows and t0.rows[0].cells:
                    ncols = len(t0.rows[0].cells)
        except Exception:
            doc = None
    if ncols >= BATCH_TABLE_MAX_COLS:
        _LOG.info("batch_table_row: wide table ncols=%d >= %d, skip", ncols, BATCH_TABLE_MAX_COLS)
        return None

    ref_texts = compress_evidence(
        evidence,
        tasks[0],
        max_chars=min(1200, config.RAG_SNIPPET_MAX_CHARS),
        all_tasks=tasks,
    )

    cell_descs: list[str] = []
    for idx, task in enumerate(tasks):
        loc = task.location_hint or {}
        col_i = int(loc.get("col", 0))
        hdr = ""
        if doc is not None and 0 <= tbl_idx < len(doc.tables):
            hdr = get_column_header_text_from_table(doc.tables[tbl_idx], col_i)
        cell_descs.append(
            f"  {idx}: 列表头={hdr!r} 章节={task.target_chapter!r} "
            f"要求={task.description!r} 字数上限={min(task.word_limit or 60, 120)}"
        )

    table_ctx_block = (
        "\n【表格上下文】\n" + table_context.strip()
        if table_context and table_context.strip()
        else ""
    )

    _tail = (
        "\n\n请输出 JSON 对象，键为上方序号（数字字符串），值为对应格的**单行**简短填写；"
        "直接依据参考资料，无依据写「资料未载明」。"
        "\n禁止：把左侧「问题」全文抄入答案列；输出「资料N：____」类模板骨架；"
        "一格内写多列内容或长段方案叙述。"
    )
    user_msg = (
        f"【批量表格填写】以下 {len(tasks)} 个单元格属于同一表格行，请一次性输出 JSON。\n"
        + "\n".join(cell_descs)
        + "\n\n【参考资料】\n"
        + ref_texts
        + table_ctx_block
        + _tail
    )

    extra_body: Dict[str, Any] = {}
    if enable_web and evidence.weak_kb:
        extra_body["enable_search"] = True
        model = config.VISION_WEB_MODEL
    else:
        model = config.SMALL_LLM_MODEL if (
            evidence.best_similarity is not None
            and evidence.best_similarity >= config.STRONG_RAG_SIMILARITY_FLOOR
        ) else config.LARGE_LLM_MODEL

    try:
        resp = chat_completions_create(
            client,
            model=model,
            messages=[
                {"role": "system", "content": _BATCH_SYSTEM},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.1,
            stream=False,
            extra_body=extra_body or None,
            max_tokens=min(800, len(tasks) * 150),
        )
        raw = (resp.choices[0].message.content if resp.choices else "") or ""
    except Exception as e:
        _LOG.warning("batch_table_row api_error: %s", e)
        return None

    try:
        data = json.loads(_strip_json_fence(raw))
    except json.JSONDecodeError:
        _LOG.warning("batch_table_row json_parse_fail raw_prefix=%r", raw[:200])
        return None

    result: Dict[int, str] = {}
    for k, v in data.items():
        try:
            idx = int(k)
            result[idx] = str(v).strip()
        except (ValueError, TypeError):
            pass

    if len(result) != len(tasks):
        _LOG.warning(
            "batch_table_row cell count mismatch: expected %d got %d",
            len(tasks),
            len(result),
        )
        return None

    _LOG.info(
        "batch_table_row done: model=%s cells=%d sim=%.2f",
        model,
        len(tasks),
        evidence.best_similarity or 0,
    )
    return result
