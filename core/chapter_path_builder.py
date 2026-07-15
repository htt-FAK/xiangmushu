"""章节层级路径构建器。

从文档 headings 列表和目标章节文本构建层级路径字符串，注入到 LLM Prompt
解决 "LLM 不知道当前章节在文档中的上下文位置" 的问题。

核心方法：
    build_chapter_path(target_chapter, doc, all_headings=None) -> str
        返回 "五、项目实施方案 > 项目实施方案（含时间安排）" 格式的层级路径

    find_target_heading_index(target_chapter, headings) -> Optional[int]
        在 headings 列表中找到最匹配 target_chapter 的索引
"""
from __future__ import annotations

import logging
import re
from typing import Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

_LOG = logging.getLogger(__name__)

# ── 标题归一化（用于模糊匹配）─────────────────────────────────────────────────
def _normalize_heading(text: str) -> str:
    """去掉空白和编号前缀，保留核心词以便比对。"""
    t = text.strip()
    # 移除中文编号 "一、二、..."  十进制编号  章节前缀
    t = re.sub(r"^[一二三四五六七八九十\d]+[、.．\s]*", "", t)
    t = re.sub(r"^第[一二三四五六七八九十\d]+[章节部]\s*", "", t)
    # 移除多余空白
    t = re.sub(r"\s+", "", t)
    return t


def _chapter_texts_match(chapter_a: str, chapter_b: str) -> bool:
    """两个章节标题是否指代同一章节（归一化后双向包含匹配）。"""
    na = _normalize_heading(chapter_a)
    nb = _normalize_heading(chapter_b)
    if not na or not nb:
        return False
    if na == nb:
        return True
    # 双向包含：任一包含另一个
    return na in nb or nb in na


def find_target_heading_index(
    target_chapter: str,
    headings: list[tuple[int, int]],
    doc: "Document | None" = None,
) -> Optional[int]:
    """在 headings 列表中找到最匹配 target_chapter 的索引。

    Args:
        target_chapter: 要匹配的章节标题文本（来自 FillTask.target_chapter）
        headings: find_all_headings 返回的 [(para_index, level), ...]
        doc: Document 对象，用于读取段落文本

    Returns:
        headings 列表中的索引 (0-based)，或 None 若未找到
    """
    if doc is None or not target_chapter:
        return None

    paras = doc.paragraphs
    for i, (para_idx, _level) in enumerate(headings):
        if para_idx >= len(paras):
            continue
        heading_text = (paras[para_idx].text or "").strip()
        if _chapter_texts_match(target_chapter, heading_text):
            return i
    return None


def build_chapter_path(
    target_chapter: str,
    doc: "Document | None" = None,
    all_headings: list[tuple[int, int]] | None = None,
) -> str:
    """构建章节层级路径字符串。

    Args:
        target_chapter: 当前 FillTask 的章节标题
        doc: Document 对象（用于获取 headings 和段落文本）
        all_headings: 预先计算的 headings 列表 (可省，否则内部调用 find_all_headings)

    Returns:
        层级路径，如：
        - 顶级章节："一、项目基本信息"
        - 嵌套章节："五、项目实施方案、技术路线及可行性分析 > 项目实施方案（含时间安排）"
        - 未知章节：target_chapter 原样返回

    示例：
        >>> build_chapter_path("项目实施方案", doc, headings)
        "五、项目实施方案、技术路线及可行性分析 > 项目实施方案（含时间安排）"
    """
    if not target_chapter or doc is None:
        return target_chapter or ""

    # 获取 headings 列表（若无传入则懒加载）
    if all_headings is None:
        try:
            from core.normal_heading_detector import find_all_headings
            import config
            threshold = int(getattr(config, "NORMAL_HEADING_THRESHOLD", 50))
            all_headings = find_all_headings(doc, threshold)
        except Exception as exc:
            _LOG.warning("build_chapter_path: heading detection failed: %s", exc)
            return target_chapter

    if not all_headings:
        return target_chapter

    # 找到当前章节在 headings 中的位置
    target_idx = find_target_heading_index(target_chapter, all_headings, doc)
    if target_idx is None:
        _LOG.debug("build_chapter_path: target %r not found in headings", target_chapter)
        return target_chapter

    target_para_idx, target_level = all_headings[target_idx]
    target_text = (doc.paragraphs[target_para_idx].text or "").strip()

    # 向上收集所有层级更高的祖先章节
    ancestors: list[str] = []
    for i in range(target_idx - 1, -1, -1):
        anc_para_idx, anc_level = all_headings[i]
        if anc_level < target_level:
            # 找到上级章节，加入祖先链（从近到远）
            anc_text = (doc.paragraphs[anc_para_idx].text or "").strip()
            if anc_text and not any(_chapter_texts_match(anc_text, a) for a in ancestors):
                ancestors.insert(0, anc_text)
                target_level = anc_level  # 继续向上找更高层

    # 组装路径
    if not ancestors:
        return target_text

    path_parts = ancestors + [target_text]
    return " > ".join(path_parts)


def format_chapter_path_block(chapter_path: str, target_chapter: str) -> str:
    """将章节路径格式化为注入 Prompt 的文本块。

    顶级章节（路径与章节相同）用简短格式，嵌套章节用完整路径。
    """
    if not chapter_path:
        return ""
    # 顶级章节（无路径层级）
    if chapter_path == (target_chapter or ""):
        return f"📍 当前章节：{chapter_path}"
    # 嵌套章节（有层级）
    return f"📍 当前章节路径：{chapter_path}"
