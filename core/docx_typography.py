"""Word 回填统一字体：宋体 + 正文小四、标题一二三级加粗规格。"""
from __future__ import annotations

import re
from copy import deepcopy
from typing import Optional

import config
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# 半磅单位 w:sz（pt * 2）
SZ_BODY = 24  # 小四 12pt
SZ_H1 = 30  # 小三 15pt
SZ_H2 = 28  # 四号 14pt
SZ_H3 = 24  # 小四 12pt

_FONT_ASCII = "SimSun"
_FONT_EAST_ASIA = "宋体"

_KEYWORD_PREFIXES = ("关键词", "Key words", "Keywords", "关键字")


def body_first_line_indent_pt() -> float:
    """正文首行缩进：小四宋体下 24pt ≈ 两个汉字宽。"""
    return float(getattr(config, "BODY_FIRST_LINE_INDENT_PT", 24) or 24)


def is_keyword_line(text: str) -> bool:
    t = (text or "").strip()
    return bool(t) and any(t.startswith(p) for p in _KEYWORD_PREFIXES)


def split_body_content_blocks(text: str) -> list[str]:
    """按换行拆成段落块，合并连续空行，段间不留空。"""
    if not text:
        return [""]
    s = re.sub(r"\r\n?", "\n", text.strip())
    s = re.sub(r"\n{2,}", "\n", s)
    parts = [p.strip() for p in s.split("\n") if p.strip()]
    return parts if parts else [""]


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """在段落后插入新段落，保留分节符等结构元素。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)

    # 检查原段落是否包含分节符属性，如果有则复制到新段落
    # 这样分节符不会被意外删除
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        sectPr = pPr.find(qn("w:sectPr"))
        if sectPr is not None:
            # 分节符应该在段落的 pPr 中，保留它
            new_pPr = OxmlElement("w:pPr")
            new_p.append(new_pPr)

    return Paragraph(new_p, paragraph._parent)


def apply_compact_paragraph_spacing(para: Paragraph) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def heading_level_from_style(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    s = style_name.strip()
    m = re.match(r"(?i)heading\s*(\d+)", s)
    if m:
        return int(m.group(1))
    m = re.match(r"标题\s*(\d+)", s)
    if m:
        return int(m.group(1))
    if s in ("Heading 1", "标题 1", "Heading1"):
        return 1
    if s in ("Heading 2", "标题 2", "Heading2"):
        return 2
    if s in ("Heading 3", "标题 3", "Heading3"):
        return 3
    return None


def _make_rPr(sz_half: int, *, bold: bool = False) -> OxmlElement:
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), _FONT_ASCII)
    rfonts.set(qn("w:hAnsi"), _FONT_ASCII)
    rfonts.set(qn("w:eastAsia"), _FONT_EAST_ASIA)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(sz_half))
    rpr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(sz_half))
    rpr.append(sz_cs)
    if bold:
        b = OxmlElement("w:b")
        rpr.append(b)
        b_cs = OxmlElement("w:bCs")
        rpr.append(b_cs)
    return rpr


def build_rPr_for_paragraph(para: Paragraph) -> OxmlElement:
    """按段落样式返回统一 rPr（正文/表格默认小四宋体）。"""
    style_name = para.style.name if para.style else ""
    lvl = heading_level_from_style(style_name or "")
    if lvl is None:
        t = (para.text or "").strip()
        if re.match(r"^摘\s*要\s*$", t):
            lvl = 1
    if lvl == 1:
        return _make_rPr(SZ_H1, bold=True)
    if lvl == 2:
        return _make_rPr(SZ_H2, bold=True)
    if lvl == 3:
        return _make_rPr(SZ_H3, bold=True)
    return _make_rPr(SZ_BODY, bold=False)


def build_body_rPr() -> OxmlElement:
    return deepcopy(_make_rPr(SZ_BODY, bold=False))


def apply_rPr_to_run(run, rpr: OxmlElement) -> None:
    clone = deepcopy(rpr)
    el = run._r
    if el.rPr is not None:
        el.remove(el.rPr)
    el.insert(0, clone)


def apply_typography_to_paragraph(para: Paragraph) -> None:
    rpr = build_rPr_for_paragraph(para)
    for run in para.runs:
        apply_rPr_to_run(run, rpr)


def apply_document_typography(doc: Document) -> None:
    """全文档段落与表格单元格 run 统一宋体字号。"""
    for para in doc.paragraphs:
        if not para.runs and not (para.text or "").strip():
            continue
        apply_typography_to_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.runs and not (para.text or "").strip():
                        continue
                    apply_typography_to_paragraph(para)
    apply_body_first_line_indent(doc)


def apply_body_first_line_indent(doc: Document) -> None:
    """正文段落首行缩进两字符（不处理标题、摘要标题行、关键词行）。"""
    pt_val = body_first_line_indent_pt()
    if pt_val <= 0:
        return
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if heading_level_from_style(style_name or "") is not None:
            continue
        t = (para.text or "").strip()
        if re.match(r"^摘\s*要\s*$", t):
            continue
        if is_keyword_line(t):
            continue
        if not t:
            continue
        para.paragraph_format.first_line_indent = Pt(pt_val)
        apply_compact_paragraph_spacing(para)


def apply_long_form_body_paragraph_format(para: Paragraph) -> None:
    """长文正文段：宋体小四 + 1.5 倍行距 + 首行缩进两字符 + 段前后无空距。"""
    apply_typography_to_paragraph(para)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pt_val = body_first_line_indent_pt()
    if pt_val > 0:
        pf.first_line_indent = Pt(pt_val)
    apply_compact_paragraph_spacing(para)


def _paragraph_has_break_element(para: Paragraph) -> bool:
    """检查段落是否包含分页符、分节符、分页断点、图片/水印等结构元素。"""
    p_el = para._element

    # w:br/w:lastRenderedPageBreak 通常嵌在 w:r 内，必须递归检查。
    for br in p_el.xpath(".//w:br"):
        br_type = br.get(qn("w:type"))
        if br_type and br_type.lower() in ("page", "column", "textwrapping"):
            return True
    if p_el.xpath(".//w:lastRenderedPageBreak"):
        return True
    if p_el.xpath(".//w:sectPr"):
        return True
    if p_el.xpath(".//w:drawing | .//w:pict"):
        return True
    return False


def remove_empty_body_paragraphs(doc: Document) -> None:
    """删除正文中无文字的空白段落（段间不留空行），但保护含分页符/分节符的段落。"""
    for para in reversed(doc.paragraphs):
        if (para.text or "").strip():
            continue
        # 保护：含分页符/分节符的段落不能删，否则排版结构崩溃
        if _paragraph_has_break_element(para):
            continue
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)


def apply_abstract_body_formats_in_document(doc: Document) -> None:
    """摘要章内各正文段：1.5 倍行距、首行缩进两字符、段间无空行。"""
    from core.filler import WordFiller

    wf = WordFiller()
    seen: set[int] = set()
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        if not (
            re.match(r"^摘\s*要\s*$", t)
            or WordFiller._heading_matches_chapter("摘要", t)
        ):
            continue
        _start, scope, _ = wf._collect_chapter_region(doc, t)
        if _start < 0:
            continue
        if any(i in seen for i in scope):
            continue
        seen.update(scope)
        for idx in scope:
            p = doc.paragraphs[idx]
            tx = (p.text or "").strip()
            if len(tx) < 40 or is_keyword_line(tx):
                continue
            if WordFiller._looks_like_writing_rubric(tx):
                continue
            if wf._classify_scope_paragraph(tx) in ("hint", "empty"):
                continue
            apply_long_form_body_paragraph_format(p)
