from __future__ import annotations

import re
from typing import Any, Dict, List

from docx import Document

from core.fill_task import FillTask
from core.parser import DocumentParser
from core.visual_auditor import DocxStructureExtractor

_PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{[^{}]+\}\}"),
    re.compile(r"【请在此填写[^】]*】"),
    re.compile(r"[_＿]{3,}"),
    re.compile(r"（\s*）|\(\s*\)"),
)


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _collect_leftover_placeholders(docx_path: str) -> List[str]:
    doc = Document(docx_path)
    leftovers: List[str] = []

    def _scan(text: str) -> None:
        raw = (text or "").strip()
        if not raw:
            return
        for pat in _PLACEHOLDER_PATTERNS:
            if pat.search(raw):
                leftovers.append(raw[:160])
                return

    for para in doc.paragraphs:
        _scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scan(cell.text)
    return leftovers


def _collect_heading_titles(docx_path: str) -> List[str]:
    parsed = DocumentParser().parse(docx_path)
    return [sec.title.strip() for sec in parsed.sections if sec.title.strip()]


def verify_filled_document(
    template_path: str,
    output_path: str,
    tasks: List[FillTask],
) -> Dict[str, Any]:
    template_struct = DocxStructureExtractor.extract(template_path)
    output_struct = DocxStructureExtractor.extract(output_path)

    leftovers = _collect_leftover_placeholders(output_path)
    heading_titles = _collect_heading_titles(output_path)
    heading_text = "\n".join(heading_titles)

    missing_chapters: List[str] = []
    seen_missing: set[str] = set()
    for task in tasks:
        chapter = (task.target_chapter or "").strip()
        if not chapter or chapter in seen_missing:
            continue
        if chapter not in heading_text:
            seen_missing.add(chapter)
            missing_chapters.append(chapter)

    template_cover = [_normalize_text(x) for x in template_struct.cover_elements if x.strip()]
    output_cover = "\n".join(_normalize_text(x) for x in output_struct.cover_elements if x.strip())
    cover_modified = any(item and item not in output_cover for item in template_cover[:3])

    template_rating = [_normalize_text(x) for x in template_struct.rating_tables if x.strip()]
    output_rating = "\n".join(_normalize_text(x) for x in output_struct.rating_tables if x.strip())
    rating_table_modified = any(item and item not in output_rating for item in template_rating[:3])

    protected_issues: List[str] = []
    if cover_modified:
        protected_issues.append("封面元素与模板不一致")
    if rating_table_modified:
        protected_issues.append("评分表或评价表与模板不一致")
    if template_struct.has_watermark and not output_struct.has_watermark:
        protected_issues.append("模板存在页眉/页脚内容，输出文档未检测到对应内容")

    ok = not leftovers and not missing_chapters and not protected_issues

    return {
        "ok": ok,
        "leftover_placeholders": leftovers,
        "missing_chapters": missing_chapters,
        "protected_issues": protected_issues,
        "cover_modified": cover_modified,
        "rating_table_modified": rating_table_modified,
        "template_words": template_struct.total_words,
        "output_words": output_struct.total_words,
        "template_tables": template_struct.total_tables,
        "output_tables": output_struct.total_tables,
    }
