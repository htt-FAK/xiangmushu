"""
模板锚点扫描与「装饰性空位」检测。
锚点格式：{{WORD_WORD}}（双花括号 + 字母数字下划线）
"""
from __future__ import annotations

import re
import uuid
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.fill_task import FillTask
from core.table_slot_expand import scan_table_fill_tasks
from core.template_slots import (
    cell_needs_fill,
    default_word_limit_for_paragraph,
    is_bracket_fill_slot,
    is_pure_hint_line,
    looks_like_fill_instruction_line,
    text_has_placeholder,
)

ANCHOR_PATTERN = re.compile(r"\{\{([A-Za-z0-9_]+)\}\}")

# 视为「无可见语义」的 Unicode 空白与零宽字符
_WS_CHARS = frozenset(
    " \t\n\r\u00a0\u1680\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007"
    "\u2008\u2009\u200a\u200b\u200c\u200d\u202f\u205f\u3000\ufeff"
)


def normalize_visible_text(s: str) -> str:
    if not s:
        return ""
    buf = []
    for ch in s:
        if ch in _WS_CHARS:
            continue
        buf.append(ch)
    return "".join(buf)


def is_semantic_empty_text(raw: str) -> bool:
    """无可见字符，或仅有下划线/横线类装饰。"""
    vis = normalize_visible_text(raw)
    if not vis:
        return True
    if set(vis) <= {"_", "＿", "-", "－", "—", "﹍", "＿"}:
        return True
    return False


def _heading_like(text: str, style_name: str) -> bool:
    from core.parser import DocumentParser  # 延迟避免环

    if style_name.startswith("Heading"):
        return True
    if DocumentParser._TITLE_PATTERN.match(text) and len(text) < 80:
        return True
    return False


def scan_anchor_tasks(template_path: str) -> List[FillTask]:
    """扫描段落与表格单元格中的 {{ANCHOR}}，生成稳定 FillTask。"""
    doc = Document(template_path)
    tasks: List[FillTask] = []
    current_chapter = "文档开头"

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc._body)
            text_full = para.text
            tstrip = text_full.strip()
            if tstrip:
                style_name = para.style.name if para.style else ""
                if _heading_like(tstrip, style_name):
                    current_chapter = tstrip
            for m in ANCHOR_PATTERN.finditer(text_full):
                anchor = m.group(0)
                tasks.append(
                    FillTask(
                        task_id=str(uuid.uuid4()),
                        target_chapter=current_chapter,
                        task_type="paragraph",
                        description=f"填写模板锚点 {anchor} 对应的申报内容（结合章节「{current_chapter}」上下文）。",
                        location_hint={"anchor": anchor},
                        word_limit=300,
                    )
                )

        elif child.tag == qn("w:tbl"):
            tidx = _table_index_for_element(doc, child)
            if tidx is None:
                continue
            table = doc.tables[tidx]
            seen_cell_ids: set[int] = set()
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    cid = id(cell._tc)
                    if cid in seen_cell_ids:
                        continue
                    seen_cell_ids.add(cid)
                    raw = cell.text or ""
                    if not ANCHOR_PATTERN.search(raw):
                        continue
                    for m in ANCHOR_PATTERN.finditer(raw):
                        anchor = m.group(0)
                        tasks.append(
                            FillTask(
                                task_id=str(uuid.uuid4()),
                                target_chapter=current_chapter,
                                task_type="table_cell",
                                description=f"填写表格中锚点 {anchor}（章节「{current_chapter}」）。",
                                location_hint={
                                    "anchor": anchor,
                                    "table_index": tidx,
                                    "row": r,
                                    "col": c,
                                },
                                word_limit=120,
                            )
                        )

    return tasks


def _table_index_for_element(doc: Document, tbl_el) -> Optional[int]:
    for i, t in enumerate(doc.tables):
        if t._element is tbl_el:
            return i
    return None


def _header_cells(table) -> List[str]:
    if not table.rows:
        return []
    return [(c.text or "").strip() for c in table.rows[0].cells]


def scan_placeholder_slots(template_path: str) -> List[FillTask]:
    """按 body 顺序扫描段落占位与表格待填格（不含 {{锚点}}）。"""
    doc = Document(template_path)
    tasks: List[FillTask] = []
    current_chapter = "文档开头"
    para_idx = 0

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            if para_idx >= len(doc.paragraphs):
                break
            para = doc.paragraphs[para_idx]
            text_full = (para.text or "").strip()
            style_name = para.style.name if para.style else ""
            if text_full and _heading_like(text_full, style_name):
                current_chapter = text_full

            if text_full and not ANCHOR_PATTERN.search(text_full):
                need = (
                    is_bracket_fill_slot(text_full)
                    or is_pure_hint_line(text_full)
                    or looks_like_fill_instruction_line(text_full)
                    or (
                        text_has_placeholder(text_full)
                        and len(text_full) <= 120
                    )
                )
                if need:
                    lh: dict = {"paragraph_text": text_full}
                    if is_bracket_fill_slot(text_full) or is_pure_hint_line(
                        text_full
                    ):
                        lh["replace_mode"] = "full"
                    wl = default_word_limit_for_paragraph(
                        current_chapter, text_full
                    )
                    desc = text_full[:80]
                    if is_bracket_fill_slot(text_full):
                        inner = text_full.strip("【】")
                        desc = f"填写：{inner}"
                    tasks.append(
                        FillTask(
                            task_id=str(uuid.uuid4()),
                            target_chapter=current_chapter,
                            task_type="paragraph",
                            description=desc,
                            location_hint=lh,
                            word_limit=wl,
                        )
                    )
            para_idx += 1

        elif child.tag == qn("w:tbl"):
            tidx = _table_index_for_element(doc, child)
            if tidx is None:
                continue
            table = doc.tables[tidx]
            tasks.extend(
                scan_table_fill_tasks(table, tidx, current_chapter)
            )

    return tasks


def scan_deterministic_fill_tasks(template_path: str) -> List[FillTask]:
    """锚点 + 占位槽，供 reconcile 与结构分析结果合并。"""
    seen: set[str] = set()
    out: List[FillTask] = []
    for t in scan_anchor_tasks(template_path) + scan_placeholder_slots(
        template_path
    ):
        key = (
            f"{t.task_type}|{t.target_chapter}|"
            f"{t.location_hint}|{t.description[:40]}"
        )
        if key in seen:
            continue
        seen.add(key)
        out.append(t)
    return out


def build_decorative_hints_for_llm(template_path: str) -> str:
    """供模板分析 LLM 参考：装饰性空单元格/段落（非锚点）。"""
    doc = Document(template_path)
    lines: List[str] = []
    for ti, table in enumerate(doc.tables):
        seen_cell_ids: set[int] = set()
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                cid = id(cell._tc)
                if cid in seen_cell_ids:
                    continue
                seen_cell_ids.add(cid)
                raw = cell.text or ""
                if ANCHOR_PATTERN.search(raw):
                    continue
                if is_semantic_empty_text(raw):
                    lines.append(
                        f"表格{ti} 第{r}行第{c}列: 程序判定为装饰性空白（仅空格/下划线等），应作为待填写。"
                    )
    return "\n".join(lines)
