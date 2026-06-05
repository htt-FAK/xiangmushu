from dataclasses import dataclass, field
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re

from core.document_models import DocumentBlock


@dataclass
class Section:
    level: int
    title: str
    content: str = ""
    tables: List[List[List[str]]] = field(default_factory=list)
    # 与 tables[i] 一一对应：在整篇 Document.tables 中的全局下标（供回填定位）
    table_doc_indices: List[int] = field(default_factory=list)


@dataclass
class ParsedDocument:
    filename: str
    sections: List[Section] = field(default_factory=list)
    raw_tables: List[List[List[str]]] = field(default_factory=list)
    kb_source_type: str = ""
    blocks: List[DocumentBlock] = field(default_factory=list)


class DocumentParser:
    """解析 Word：按 body 顺序遍历段落与表格，章节与表格归属一致。"""

    _TITLE_PATTERN = re.compile(r"^(第?[一二三四五六七八九十\d]+[、.．]\s*|(\d+\.)+\s*)")

    def parse(self, file_path: str) -> ParsedDocument:
        doc = Document(file_path)
        filename = file_path.rsplit("\\", 1)[-1].rsplit("/", 1)[-1]
        blocks: List[DocumentBlock] = []

        all_tables_raw: List[List[List[str]]] = []
        for tbl in doc.tables:
            table_data = []
            for row in tbl.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            all_tables_raw.append(table_data)

        sections: List[Section] = [Section(level=0, title="文档开头")]
        current_section = sections[0]

        def _heading_level(style_name: str) -> int:
            if not style_name.startswith("Heading"):
                return 0
            try:
                return int(style_name.replace("Heading", "").strip())
            except ValueError:
                return 1

        for child in doc.element.body:
            if child.tag == qn("w:p"):
                para = Paragraph(child, doc._body)
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                is_heading = False
                heading_level = 0

                if style_name.startswith("Heading"):
                    is_heading = True
                    heading_level = _heading_level(style_name)
                elif self._TITLE_PATTERN.match(text) and len(text) < 80:
                    is_heading = True
                    heading_level = 1

                if is_heading:
                    current_section = Section(level=heading_level, title=text)
                    sections.append(current_section)
                    blocks.append(
                        DocumentBlock(
                            text=text,
                            page=1,
                            block_type="heading",
                            source_type="docx",
                            chapter=text,
                            metadata={"style_name": style_name or ""},
                        )
                    )
                else:
                    if current_section.content:
                        current_section.content += "\n"
                    current_section.content += text
                    blocks.append(
                        DocumentBlock(
                            text=text,
                            page=1,
                            block_type="text",
                            source_type="docx",
                            chapter=current_section.title,
                            metadata={"style_name": style_name or ""},
                        )
                    )

            elif child.tag == qn("w:tbl"):
                table_idx = self._table_index_for_element(doc, child)
                if table_idx is not None and table_idx < len(all_tables_raw):
                    current_section.tables.append(all_tables_raw[table_idx])
                    current_section.table_doc_indices.append(table_idx)
                    header = " | ".join(all_tables_raw[table_idx][0]) if all_tables_raw[table_idx] else ""
                    blocks.append(
                        DocumentBlock(
                            text=self._table_to_text(all_tables_raw[table_idx]),
                            page=1,
                            block_type="table",
                            source_type="docx",
                            chapter=current_section.title,
                            table_index=table_idx,
                            content_format="markdown",
                            table_header=header,
                            metadata={"doc_table_index": table_idx},
                        )
                    )

        if sections[0].level == 0 and not sections[0].content and not sections[0].tables:
            sections = sections[1:]

        return ParsedDocument(
            filename=filename,
            sections=sections,
            raw_tables=all_tables_raw,
            kb_source_type="docx",
            blocks=blocks,
        )

    @staticmethod
    def _table_index_for_element(doc: Document, tbl_el) -> Optional[int]:
        for i, t in enumerate(doc.tables):
            if t._element is tbl_el:
                return i
        return None

    def _map_table_positions(self, doc: Document):
        """兼容旧调用：记录每个表格前的段落序号。"""
        positions = []
        table_idx = 0
        para_idx = 0
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "tbl":
                positions.append(("table", table_idx, para_idx))
                table_idx += 1
            elif tag == "p":
                para_idx += 1
        return positions

    @staticmethod
    def _table_to_text(table: List[List[str]]) -> str:
        lines = [" | ".join(row) for row in table if row]
        return "\n".join(lines)
