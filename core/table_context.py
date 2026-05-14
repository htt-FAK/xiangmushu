"""从 Word 模板抽取表格单元格周边的表头、左侧与行列文本，供生成与审核注入。"""
from __future__ import annotations

import re
from typing import Any, Optional

from docx import Document


def _cell_plain(cell) -> str:
    return (cell.text or "").strip().replace("\n", " ")


def get_column_header_text_from_table(table: Any, col: int, max_header_rows: int = 2) -> str:
    """取表第 0～(max_header_rows-1) 行、第 col 列合并文本（去重、截断），供批量生成与上下文共用。"""
    if not table.rows or col < 0:
        return ""
    bits: list[str] = []
    seen: set[str] = set()
    for hr in range(min(max_header_rows, len(table.rows))):
        cells = table.rows[hr].cells
        if col >= len(cells):
            continue
        t = _cell_plain(cells[col])
        if t and t not in seen:
            seen.add(t)
            bits.append(t)
    return (" / ".join(bits))[:220]


def get_column_header_from_path(template_path: str, table_index: int, col: int) -> str:
    """只读路径，解析表后返回第 col 列表头文本。"""
    try:
        doc = Document(template_path)
    except Exception:
        return ""
    if table_index < 0 or table_index >= len(doc.tables):
        return ""
    return get_column_header_text_from_table(doc.tables[table_index], col)


def _short_answer_hint_from_header(header: str) -> str:
    """表头关键词驱动的短答提示（无 LLM）。"""
    h = header or ""
    if not h:
        return ""
    if any(
        k in h
        for k in (
            "是/否",
            "是否",
            "页码",
            "片段",
            "标题",
            "准确",
            "资料出处",
            "出处",
        )
    ):
        return "（本列偏判定/定位：只输出极短短语或「是/否/部分准确/资料未载明」等，勿写长段。）"
    if "资料" in h and ("编号" in h or "名称" in h or "来源" in h):
        return "（本列填资料标识或名称：单行短答。）"
    return ""


def _placeholder_cell_hint(cell_preview: str) -> str:
    """当前格为下划线/资料占位时的提示。"""
    t = (cell_preview or "").strip()
    if not t:
        return ""
    if re.search(r"资料\s*\d+\s*[：:]\s*_+", t) or re.fullmatch(r"_+", t.replace(" ", "")):
        return "（本格原为占位符：请输出可替换进格的实质短答，勿保留下划线或「资料N：____」骨架。）"
    return ""


def build_table_cell_context(
    template_path: str,
    table_index: int,
    row: int,
    col: int,
    max_chars: int = 800,
) -> str:
    """
    只读打开模板，抽取：
    - 表头：第 0～1 行各列文本；
    - 本列表头（第 col 列 0～1 行）；
    - 本格左侧各列（不含本列）；
    - 本行第 0 列（常为行标题）；
    - 当前格原有占位文本预览；
    - 表头/占位驱动的短答提示。
    """
    try:
        doc = Document(template_path)
    except Exception:
        return ""
    if table_index < 0 or table_index >= len(doc.tables):
        return ""
    table = doc.tables[table_index]
    if not table.rows:
        return ""
    nrows = len(table.rows)
    ncols = len(table.rows[0].cells) if table.rows[0].cells else 0
    if row < 0 or row >= nrows or col < 0 or col >= ncols:
        return ""

    parts: list[str] = []
    seen: set[str] = set()

    def add_line(label: str, text: str) -> None:
        t = (text or "").strip()
        if not t or t in seen:
            return
        seen.add(t)
        parts.append(f"{label}{t}")

    for hr in range(min(2, nrows)):
        cells = table.rows[hr].cells
        row_bits = [_cell_plain(cells[c]) for c in range(len(cells)) if c < len(cells)]
        line = " | ".join(x for x in row_bits if x)
        if line:
            add_line(f"表头第{hr}行: ", line)

    col_hdr = get_column_header_text_from_table(table, col)
    if col_hdr:
        add_line("本列表头: ", col_hdr[:220])
    sh = _short_answer_hint_from_header(col_hdr)
    if sh:
        parts.append(sh)

    if col > 0:
        left_bits = []
        for c in range(0, col):
            if c < len(table.rows[row].cells):
                t = _cell_plain(table.rows[row].cells[c])
                if t:
                    left_bits.append(f"列{c}「{t}」")
        if left_bits:
            parts.append("本格左侧: " + "；".join(left_bits))

    if col != 0 and len(table.rows[row].cells) > 0:
        r0 = _cell_plain(table.rows[row].cells[0])
        if r0:
            add_line("本行行首列(常为行标题): ", r0[:240])

    cur = _cell_plain(table.rows[row].cells[col])
    if cur:
        add_line("当前格原有内容: ", cur[:400])
    ph = _placeholder_cell_hint(cur)
    if ph:
        parts.append(ph)

    text = "\n".join(parts)
    if len(text) > max_chars:
        return text[:max_chars] + "\n…(截断)"
    return text
