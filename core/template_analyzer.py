import logging
from typing import Any, Dict, List, Optional

from core.dashscope_chat import chat_completions_create
from core.parser import DocumentParser
from core.fill_task import FillTask
from core.slot_scanner import scan_anchor_tasks, build_decorative_hints_for_llm
from core.template_vision import apply_chapter_hints_to_tasks, compact_profile_for_analyzer
import config
import json
import re
import uuid

from docx import Document

from core.filler import WordFiller

_LOG = logging.getLogger(__name__)


class TemplateAnalyzer:
    """分析模板：优先锚点 {{NAME}} 扫描；否则 LLM + 装饰性空位提示。"""

    def __init__(self):
        self._client = config.openai_client_for_template_analyze()
        self._parser = DocumentParser()

    def analyze(
        self,
        template_path: str,
        vision_profile: Optional[Dict[str, Any]] = None,
    ) -> List[FillTask]:
        anchor_tasks = scan_anchor_tasks(template_path)
        if anchor_tasks:
            if vision_profile:
                apply_chapter_hints_to_tasks(anchor_tasks, vision_profile)
            return anchor_tasks

        doc = self._parser.parse(template_path)
        structure_text = self._build_structure_text(doc)
        deco = build_decorative_hints_for_llm(template_path)
        if deco.strip():
            structure_text += (
                "\n\n【程序化检测】以下位置为装饰性空白（仅空格/下划线等），"
                "必须输出为待填写，并给出正确的 table_index/row/col 或段落关键词：\n"
                + deco
            )

        vision_append = ""
        if vision_profile:
            compact = compact_profile_for_analyzer(vision_profile)
            if compact:
                vision_append = (
                    "\n\n【视觉版式摘要（由模板页渲染图分析，供区分说明区与待填区、表格语义）】\n"
                    + compact
                )
            fb = (vision_profile.get("ooxml_fallback") or "").strip()
            if fb and len(fb) > 80:
                vision_append += (
                    "\n\n【纯文本结构降级摘要（无 PDF/视觉时的 docx 文本抽取）】\n"
                    + fb[:6000]
                )

        max_prompt = int(getattr(config, "TEMPLATE_ANALYZE_MAX_PROMPT_CHARS", 12000))
        if len(structure_text) > max_prompt // 2:
            structure_text = (
                structure_text[: max_prompt // 2]
                + "\n…(上文结构已截断，请根据可见表格索引与标题继续识别填空位)…\n"
            )
        if len(vision_append) > max_prompt // 3:
            vision_append = vision_append[: max_prompt // 3] + "\n…(视觉摘要已截断)…\n"

        prompt = f"""你是一个文档分析助手。以下是项目计划书模板的结构：

{structure_text}
{vision_append}

请找出所有需要填写内容的空位（如空白段落、包含"请填写"/"（ ）"/"____"等占位符的段落或表格单元格）。
对每个空位输出 JSON 数组，每个元素包含：
- chapter: 所属章节标题
- type: "paragraph" 或 "table_cell"
- description: 应该填写什么内容（根据上下文推断）
- location_hint: 定位信息（段落用 {{"paragraph_text": "上下文关键词"}}，表格用 {{"table_index": 数字, "row": 行号, "col": 列号}}）
- word_limit: 建议字数
- replace_mode: （可选，**仅 type 为 paragraph 时有效**）"full" 或 "placeholder_only"。
  若同一段内左侧/前后为固定说明文字、仅「请填写」「请在此填写」「____」「（ ）」等为待填占位，必须填 **placeholder_only**；
  若整行仅为「（请在此填写…）」类短提示（独立一行），用 **full** 或省略（整行替换为正文）；
  若独立一行为「摘要：在以下填写…」「在以下填写」类章节填写指引（无长段成稿），必须用 **full**；
  若整段为「【请在此填写……】」类括号占位行，必须单独一条 paragraph、**full**、location_hint.paragraph_text 为整行原文；
  若整段几乎全为待生成正文，用 **full** 或省略。table_cell 省略本字段。
  摘要 chapter 的 word_limit 建议约 650 字；正文宜连贯分段，段首缩进两字符，段间勿留空行；须与程序扫槽一一对应，勿漏「【请在此填写】」行。

重要（表格）：
- location_hint.table_index **必须**等于上文中「doc.tables索引=」后面的整数，与整篇文档中表格的全局顺序一致，**绝不是**每个章节内从 0 重新编号。
- row、col 均为 **0 起算**。多列表格中，**左侧列多为标题/说明**，待填的简短答案、勾选说明通常在 **第 1 列或最右列**；不要把长段正文写进明显仅为标题的窄列。
- table_cell 类任务的 word_limit 建议不超过 120。
- **table_cell 的 description 必须是一句话**（≤80 字），只描述**本格**应填要点；须**显式带上该列表头中的关键词**（从表头第 0/1 行对应列抄录，勿整行合并左侧大段说明）。
- 同一行多列待填时，必须输出**多条** JSON，每条对应一个 (row,col)，**禁止**把整行说明塞进一个 table_cell 的 description。

只输出 JSON 数组，不要其他内容。"""

        analyze_model = config.TEMPLATE_ANALYZE_MODEL
        analyze_temp = config.TEMP_TEMPLATE_ANALYZE
        analyze_timeout = float(getattr(config, "TEMPLATE_ANALYZE_TIMEOUT", 90))
        prompt_chars = len(prompt)
        _LOG.info(
            "template_analyze request model=%s timeout=%.0fs prompt_chars=%d",
            analyze_model,
            analyze_timeout,
            prompt_chars,
        )

        def _call_analyze(model: str):
            return chat_completions_create(
                self._client,
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "你是文档分析助手，只输出 JSON，不要任何解释。",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=analyze_temp,
                max_tokens=8192,
                timeout=analyze_timeout,
            )

        response = _call_analyze(analyze_model)
        content = (response.choices[0].message.content or "").strip()
        if not content:
            fallback = getattr(config, "TEMPLATE_ANALYZE_FALLBACK_MODEL", "") or (
                config.FALLBACK_LLM_MODEL_1
            )
            if fallback and fallback != analyze_model:
                _LOG.warning(
                    "template_analyze 空回复，改用 %s 重试", fallback
                )
                response = _call_analyze(fallback)
                content = (response.choices[0].message.content or "").strip()
        _LOG.info(
            "template_analyze done model=%s content_len=%d",
            analyze_model,
            len(content),
        )
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]

        try:
            raw_tasks = json.loads(content)
        except json.JSONDecodeError:
            raw_tasks = []

        tasks = []
        for item in raw_tasks:
            ttype = item.get("type", "paragraph")
            try:
                wl = int(item.get("word_limit", 300) or 300)
            except (TypeError, ValueError):
                wl = 300
            if ttype == "table_cell":
                wl = min(max(wl, 1), 120)
            lh = dict(item.get("location_hint", {}) or {})
            rm = item.get("replace_mode") or lh.get("replace_mode")
            if isinstance(rm, str) and rm.strip().lower() in ("full", "placeholder_only"):
                rml = rm.strip().lower()
                if ttype == "paragraph":
                    lh["replace_mode"] = rml
            tasks.append(
                FillTask(
                    task_id=str(uuid.uuid4()),
                    target_chapter=item.get("chapter", ""),
                    task_type=ttype,
                    description=item.get("description", ""),
                    location_hint=lh,
                    word_limit=wl,
                )
            )
        self._supplement_section_paragraph_tasks(doc, tasks)
        if vision_profile:
            apply_chapter_hints_to_tasks(tasks, vision_profile)
        self._apply_replace_mode_heuristics(template_path, tasks)
        return tasks

    def _supplement_section_paragraph_tasks(self, doc, tasks: List[FillTask]) -> None:
        """补齐章节正文段落任务，避免只有表格任务导致小节正文为空。"""
        for sec in doc.sections:
            if self._skip_section_paragraph_task(sec):
                continue

            candidate = self._pick_section_body_paragraph(sec)
            if self._has_matching_paragraph_task(tasks, sec.title, candidate):
                continue

            location_hint: Dict[str, Any] = {"replace_mode": "full"}
            if candidate:
                location_hint["paragraph_text"] = candidate
                description = self._paragraph_description(sec.title, candidate)
            elif self._section_can_use_empty_paragraph(sec):
                location_hint["paragraph_text"] = ""
                description = self._paragraph_description(sec.title, "")
            else:
                continue

            tasks.append(
                FillTask(
                    task_id=str(uuid.uuid4()),
                    target_chapter=sec.title,
                    task_type="paragraph",
                    description=description,
                    location_hint=location_hint,
                    word_limit=self._word_limit_for_section(sec.title),
                )
            )

    @staticmethod
    def _compact_text(text: str) -> str:
        return re.sub(r"\s+", "", text or "")

    def _skip_section_paragraph_task(self, sec) -> bool:
        title = (sec.title or "").strip()
        compact = self._compact_text(title)
        if not title or title == "文档开头":
            return True
        if "目录" in compact:
            return True
        if compact.startswith("附录"):
            return True
        if "评分标准" in compact or "提交材料清单" in compact:
            return True
        if sec.level <= 0:
            return True
        if sec.level == 1 and not sec.content.strip():
            return True
        return False

    def _pick_section_body_paragraph(self, sec) -> Optional[str]:
        lines = [
            line.strip()
            for line in (sec.content or "").splitlines()
            if line.strip()
        ]
        if not lines:
            return None

        for line in lines:
            if self._is_body_fill_candidate(line, allow_placeholder=True):
                return line

        for line in lines:
            if self._is_body_fill_candidate(line, allow_placeholder=False):
                return line

        return None

    def _is_body_fill_candidate(self, text: str, *, allow_placeholder: bool) -> bool:
        t = (text or "").strip()
        if not t or self._is_non_body_template_line(t):
            return False
        if allow_placeholder and (
            WordFiller._is_pure_hint_line(t)
            or WordFiller._looks_like_fill_instruction_line(t)
            or WordFiller._text_has_placeholder(t)
        ):
            return True
        if allow_placeholder:
            return False
        if self._looks_like_section_body_guidance(t):
            return True
        return (
            WordFiller._looks_like_writing_rubric(t)
            or WordFiller._looks_like_template_guidance(t)
        )

    @staticmethod
    def _looks_like_section_body_guidance(text: str) -> bool:
        t = (text or "").strip()
        if len(t) < 12 or len(t) > 500:
            return False
        starters = (
            "说明",
            "描述",
            "填写",
            "本节应",
            "请",
            "用文字",
            "工作流应",
            "角色设定是",
            "知识库不是",
            "从",
            "客观分析",
            "围绕",
            "展望",
            "至少设计",
            "结果分析",
        )
        if not t.startswith(starters):
            return False
        guidance_markers = (
            "建议",
            "至少",
            "必须",
            "应",
            "不要",
            "不能",
            "避免",
            "如何",
            "哪些",
            "什么",
            "截图",
            "展示",
            "分析",
            "总结",
            "说明",
            "描述",
            "填写",
            "项目",
            "智能体",
            "工作流",
            "数据库",
            "知识库",
            "插件",
            "演示",
            "用户",
            "场景",
        )
        return any(marker in t for marker in guidance_markers)

    @staticmethod
    def _is_non_body_template_line(text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return True
        if re.match(r"^图\s*\d+(?:\.\d+)*\s*.+", t):
            return True
        if t.startswith(("关键词", "Key words", "Keywords", "关键字")):
            return True
        return False

    def _section_can_use_empty_paragraph(self, sec) -> bool:
        """无可见正文提示时，用小节内空段承载正文；主要覆盖表格型小节。"""
        if sec.level < 2:
            return False
        title = self._compact_text(sec.title)
        if not re.match(r"^\d+\.\d+", title):
            return False
        return bool(sec.tables or not (sec.content or "").strip())

    def _has_matching_paragraph_task(
        self,
        tasks: List[FillTask],
        chapter: str,
        paragraph_text: Optional[str],
    ) -> bool:
        chapter_key = self._compact_text(chapter)
        candidate_key = self._compact_text(paragraph_text or "")
        same_chapter_tasks = [
            t
            for t in tasks
            if t.task_type == "paragraph"
            and self._compact_text(t.target_chapter) == chapter_key
        ]
        if not same_chapter_tasks:
            return False
        if not candidate_key:
            return True

        for task in same_chapter_tasks:
            hint = task.location_hint or {}
            existing = str(hint.get("paragraph_text") or "")
            existing_key = self._compact_text(existing)
            if not existing_key:
                continue
            if (
                existing_key == candidate_key
                or existing_key in candidate_key
                or candidate_key in existing_key
            ):
                return True
        return False

    def _paragraph_description(self, chapter: str, candidate: str) -> str:
        if candidate:
            return f"填写「{chapter}」正文：{candidate[:100]}"
        return (
            f"补充「{chapter}」的小节正文，概括本节完成情况、关键证据、"
            "实现效果和可复核材料。"
        )

    def _word_limit_for_section(self, chapter: str) -> int:
        compact = self._compact_text(chapter)
        if "摘要" in compact:
            return int(getattr(config, "ABSTRACT_WORD_LIMIT", 650))
        if compact.startswith(("1.1", "1.2", "2.1", "2.2")):
            return 450
        if compact.startswith(("3.2", "3.3", "3.4", "3.5", "3.6")):
            return 520
        if compact.startswith("4.3"):
            return 500
        if compact.startswith(("5.2", "5.3", "6.1", "6.2", "6.3")):
            return 480
        if compact.startswith(("1.3", "2.3", "3.1", "4.2")):
            return 420
        if compact.startswith(("4.1", "5.1")):
            return 350
        return 400

    def _apply_replace_mode_heuristics(
        self, template_path: str, tasks: List[FillTask]
    ) -> None:
        """根据模板段落文本自动设置 placeholder_only（说明+占位混排）。"""
        try:
            doc = Document(template_path)
        except Exception:
            return
        paras = doc.paragraphs
        for task in tasks:
            if task.task_type != "paragraph":
                continue
            lh = task.location_hint or {}
            if lh.get("replace_mode"):
                continue
            hint = (lh.get("paragraph_text") or "").strip()
            para_text = ""
            for p in paras:
                t = p.text or ""
                if hint and hint in t:
                    para_text = t
                    break
                if task.target_chapter and WordFiller._heading_matches_chapter(
                    task.target_chapter, t
                ):
                    continue
            if not para_text and task.description:
                if "说明" in task.description and (
                    "请填写" in task.description or "占位" in task.description
                ):
                    lh["replace_mode"] = "placeholder_only"
                    task.location_hint = lh
                    continue
            if not para_text:
                continue
            if WordFiller._is_pure_hint_line(para_text):
                lh["replace_mode"] = "full"
                task.location_hint = lh
                continue
            if WordFiller._looks_like_fill_instruction_line(para_text):
                lh["replace_mode"] = "full"
                task.location_hint = lh
                continue
            if WordFiller._text_has_placeholder(para_text) and (
                "说明" in para_text or len(para_text) > 50
            ):
                lh["replace_mode"] = "placeholder_only"
                task.location_hint = lh

    def _build_structure_text(self, doc) -> str:
        lines = []
        for sec in doc.sections:
            prefix = "  " * max(sec.level - 1, 0)
            lines.append(f"{prefix}[Heading{sec.level}] {sec.title}")
            if sec.content:
                preview = sec.content[:200].replace("\n", " ")
                lines.append(f"{prefix}  内容预览: {preview}")
            for t_idx, table in enumerate(sec.tables):
                doc_ti = (
                    sec.table_doc_indices[t_idx]
                    if t_idx < len(sec.table_doc_indices)
                    else t_idx
                )
                lines.append(
                    f"{prefix}  [表格 doc.tables索引={doc_ti}] {len(table)}行x"
                    f"{len(table[0]) if table else 0}列"
                )
                if table:
                    for row in table[:5]:
                        lines.append(f"{prefix}    {' | '.join(row)}")
                    if len(table) > 5:
                        lines.append(f"{prefix}    ... (共{len(table)}行)")
        return "\n".join(lines)
