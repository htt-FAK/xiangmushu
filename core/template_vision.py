"""模板 docx → PDF → 页图 → 多模态：版式/填写策略/chapter_hints；失败时 OOXML 文本降级。"""
from __future__ import annotations

import base64
import hashlib
import json
import logging
import os
import re
import shutil
import sys
import subprocess
import tempfile
from copy import deepcopy
from typing import Any, Dict, List, Optional, Tuple, Union

import config
from core.dashscope_chat import chat_completions_create
from core.fill_task import FillTask

_LOG = logging.getLogger(__name__)

VISION_TEMPLATE_JSON_PROMPT = """你是 Word 排版与申报模板分析专家。下面将提供若干页由 Word 模板渲染成的页面截图（按页序，第 1 页对应下标 0）。
请根据**可见版式**（标题层级、表格、说明性文字 vs 空白/下划线待填区）输出**仅一个 JSON 对象**（不要 Markdown 围栏），键必须齐全：
- layout_notes: 字符串，概括整体版式、表格密度、明显「说明」与「待填」区域。
- fill_strategy: 字符串，建议哪些区域「整格/整段清空再写」、哪些「只替换占位符保留左侧说明」。
- style_observations: 字符串，定性描述正文/表格常见字号层级（不必精确 pt）。
- chapter_hints: 数组，每项为对象，含两个键：
  - chapter_anchor: 字符串，与文档中章节标题尽量一致的短语（便于程序匹配）。
  - hint: 字符串，该章节附近的**写作/填写要求**（摘录或概括截图中的说明文字）。
- table_page_hints: 数组，每项描述**一个表格**在截图中的位置（与 Word 中表格从上到下的顺序一致，第 1 个表为 0）：
  - table_ordinal: 整数，从 0 起，与文档第 1 个表、第 2 个表……对应。
  - likely_pages: 整数数组，该表**主要出现或跨页**的页下标（与截图顺序一致，从 0 起）；尽量 1～2 页，不要列出全部页。
  - note: 字符串，可选，如合并单元格、表头占多行等视觉提示。

若无某类信息，用空字符串或空数组。不要输出 JSON 以外的任何字符。"""


def _cache_dir() -> str:
    d = os.path.join(os.path.dirname(__file__), "..", "data", ".cache", "template_vision")
    abs_path = os.path.abspath(d)
    os.makedirs(abs_path, exist_ok=True)
    return abs_path


def _cache_key(docx_path: str) -> str:
    ap = os.path.abspath(docx_path)
    try:
        mt = int(os.path.getmtime(ap))
    except OSError:
        mt = 0
    raw = f"{ap}|{mt}".encode("utf-8", errors="replace")
    return hashlib.sha256(raw).hexdigest()[:40]


def _cache_path(docx_path: str) -> str:
    """旧版：单文件 profile.json（与 bundle 内 profile 二选一读取）。"""
    return os.path.join(_cache_dir(), f"{_cache_key(docx_path)}.json")


def cache_bundle_dir(docx_path: str) -> str:
    """模板视觉缓存目录：内含 profile.json 与 page_XXX.png。"""
    return os.path.join(_cache_dir(), _cache_key(docx_path))


def _profile_path_in_bundle(bundle_dir: str) -> str:
    return os.path.join(bundle_dir, "profile.json")


def _list_cached_page_png_paths(bundle_dir: str) -> List[str]:
    if not os.path.isdir(bundle_dir):
        return []
    names = [n for n in os.listdir(bundle_dir) if n.startswith("page_") and n.endswith(".png")]
    names.sort()
    return [os.path.join(bundle_dir, n) for n in names]


def _write_png_pages(bundle_dir: str, pages: List[bytes]) -> None:
    os.makedirs(bundle_dir, exist_ok=True)
    for i, blob in enumerate(pages):
        p = os.path.join(bundle_dir, f"page_{i:03d}.png")
        with open(p, "wb") as wf:
            wf.write(blob)


def _load_png_paths(paths: List[str]) -> List[bytes]:
    out: List[bytes] = []
    for p in paths:
        try:
            with open(p, "rb") as rf:
                out.append(rf.read())
        except OSError:
            continue
    return out


def ensure_template_page_pngs(docx_path: str, *, force_refresh: bool = False) -> int:
    """确保缓存目录下存在模板页 PNG（仅 PDF 栅格化，不调视觉 LLM）。返回页数。"""
    bundle = cache_bundle_dir(docx_path)
    existing = _list_cached_page_png_paths(bundle)
    if existing and not force_refresh:
        return len(existing)
    with tempfile.TemporaryDirectory(prefix="tplpng_") as td:
        pdf_path = os.path.join(td, "tpl.pdf")
        if not export_docx_to_pdf(docx_path, pdf_path):
            return 0
        pages = pdf_to_png_pages(
            pdf_path,
            int(config.TEMPLATE_VISION_MAX_PAGES),
            float(config.TEMPLATE_VISION_ZOOM),
        )
        if not pages:
            return 0
        os.makedirs(bundle, exist_ok=True)
        _write_png_pages(bundle, pages)
        return len(pages)


def get_cached_vision_profile(docx_path: str) -> Dict[str, Any]:
    """读取已缓存的 profile（bundle 内 profile.json 优先，其次旧版平铺 json）。"""
    bundle = cache_bundle_dir(docx_path)
    pp = _profile_path_in_bundle(bundle)
    if os.path.isfile(pp):
        try:
            with open(pp, "r", encoding="utf-8") as f:
                blob = json.load(f)
            prof = blob.get("profile") or blob
            return prof if isinstance(prof, dict) else {}
        except Exception as e:
            _LOG.debug("read bundle profile: %s", e)
    legacy = _cache_path(docx_path)
    if os.path.isfile(legacy):
        try:
            with open(legacy, "r", encoding="utf-8") as f:
                blob = json.load(f)
            prof = blob.get("profile") or blob
            return prof if isinstance(prof, dict) else {}
        except Exception as e:
            _LOG.debug("read legacy profile: %s", e)
    return {}


def page_indices_for_table(
    profile: Dict[str, Any],
    table_index: int,
    n_pages: int,
) -> List[int]:
    """根据 table_page_hints 选择页下标；失败则 0..min(n)-1 截断到 TABLE_VISION_MAX_PAGES。"""
    cap = max(1, int(getattr(config, "TABLE_VISION_MAX_PAGES", 6) or 6))
    if n_pages <= 0:
        return []
    hints = profile.get("table_page_hints") or []
    picked: List[int] = []
    if isinstance(hints, list):
        for h in hints:
            if not isinstance(h, dict):
                continue
            try:
                ordi = int(h.get("table_ordinal", h.get("table_index", -1)))
            except (TypeError, ValueError):
                continue
            if ordi != int(table_index):
                continue
            raw = h.get("likely_pages") or h.get("pages") or []
            if isinstance(raw, list):
                for x in raw:
                    try:
                        pi = int(x)
                    except (TypeError, ValueError):
                        continue
                    if 0 <= pi < n_pages and pi not in picked:
                        picked.append(pi)
            break
    if not picked:
        picked = list(range(min(n_pages, cap)))
    return picked[:cap]


def load_table_cell_vision_pngs(
    docx_path: str,
    table_index: int,
) -> List[bytes]:
    """供 table_cell 生成：按 profile 选取页 PNG；无缓存则先 ensure_template_page_pngs。"""
    if not getattr(config, "TABLE_CELL_VISION", True):
        return []
    bundle = cache_bundle_dir(docx_path)
    paths = _list_cached_page_png_paths(bundle)
    if not paths:
        ensure_template_page_pngs(docx_path)
        paths = _list_cached_page_png_paths(bundle)
    if not paths:
        return []
    prof = get_cached_vision_profile(docx_path)
    idxs = page_indices_for_table(prof, int(table_index), len(paths))
    chosen = [paths[i] for i in idxs if 0 <= i < len(paths)]
    if not chosen:
        chosen = paths[: max(1, int(getattr(config, "TABLE_VISION_MAX_PAGES", 6) or 6))]
    return _load_png_paths(chosen)


def build_table_cell_user_content(
    text_body: str,
    png_pages: Optional[List[bytes]],
) -> Union[str, List[Dict[str, Any]]]:
    """table_cell 用户侧：无图则纯文本，有图则 OpenAI 多模态 content 数组。"""
    if not png_pages:
        return text_body
    parts: List[Dict[str, Any]] = [
        {
            "type": "text",
            "text": (
                "【模板页面截图】以下为 Word 模板渲染后的页面图像，请结合图中的表格线、合并格与表头文字理解版式；"
                "并与下文「本格表格上下文」、参考资料一并判断，只输出本格应填的简短答案。\n\n"
                + text_body
            ),
        },
    ]
    for blob in png_pages:
        b64 = base64.standard_b64encode(blob).decode("ascii")
        parts.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"},
            }
        )
    return parts


def _export_pdf_docx2pdf(docx_path: str, pdf_path: str) -> bool:
    try:
        from docx2pdf import convert  # type: ignore
    except ImportError:
        return False
    try:
        convert(docx_path, pdf_path)
        return os.path.isfile(pdf_path) and os.path.getsize(pdf_path) > 0
    except Exception as e:
        _LOG.debug("docx2pdf failed: %s", e)
        return False


def _find_soffice() -> Optional[str]:
    exe = "soffice.exe" if os.name == "nt" else "soffice"
    w = shutil.which(exe) or shutil.which("soffice")
    if w:
        return w
    if os.name == "nt":
        for base in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if os.path.isfile(base):
                return base
    if sys.platform == "darwin":
        for base in (
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/Applications/LibreOffice-still.app/Contents/MacOS/soffice",
        ):
            if os.path.isfile(base):
                return base
    return None


def _export_pdf_soffice(docx_path: str, pdf_path: str) -> bool:
    soffice = _find_soffice()
    if not soffice:
        return False
    outdir = os.path.dirname(pdf_path) or "."
    os.makedirs(outdir, exist_ok=True)
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=120,
        )
        base = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        candidate = os.path.join(outdir, base)
        if os.path.isfile(candidate):
            if os.path.abspath(candidate) != os.path.abspath(pdf_path):
                shutil.move(candidate, pdf_path)
            return os.path.isfile(pdf_path) and os.path.getsize(pdf_path) > 0
    except Exception as e:
        _LOG.debug("soffice pdf export failed: %s", e)
    return False


def _export_pdf_pandoc(docx_path: str, pdf_path: str) -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    try:
        subprocess.run(
            [pandoc, docx_path, "-o", pdf_path],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=120,
        )
        return os.path.isfile(pdf_path) and os.path.getsize(pdf_path) > 0
    except Exception as e:
        _LOG.debug("pandoc pdf export failed: %s", e)
    return False


def export_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """依次尝试 docx2pdf（需本机 Word）、LibreOffice、pandoc。"""
    if _export_pdf_docx2pdf(docx_path, pdf_path):
        return True
    if _export_pdf_soffice(docx_path, pdf_path):
        return True
    if _export_pdf_pandoc(docx_path, pdf_path):
        return True
    return False


def pdf_to_png_pages(pdf_path: str, max_pages: int, zoom: float) -> List[bytes]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        _LOG.warning("PyMuPDF (fitz) 未安装，无法将 PDF 栅格化")
        return []
    out: List[bytes] = []
    try:
        doc = fitz.open(pdf_path)
        try:
            n = min(max_pages, len(doc))
            mat = fitz.Matrix(zoom, zoom)
            for i in range(n):
                page = doc[i]
                pix = page.get_pixmap(matrix=mat, alpha=False)
                out.append(pix.tobytes("png"))
        finally:
            doc.close()
    except Exception as e:
        _LOG.warning("pdf_to_png_pages: %s", e)
    return out


def extract_ooxml_text_fallback(docx_path: str, max_chars: int = 8000) -> str:
    """无 PDF 时从 docx 抽取纯文本作降级上下文。"""
    try:
        from docx import Document

        doc = Document(docx_path)
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_t = " | ".join(
                    (c.text or "").strip().replace("\n", " ") for c in row.cells
                )
                if row_t.strip():
                    parts.append(row_t)
        blob = "\n".join(parts)
        if len(blob) > max_chars:
            blob = blob[:max_chars] + "\n…(已截断)"
        return blob
    except Exception as e:
        _LOG.debug("ooxml fallback: %s", e)
        return ""


def _strip_json_fence(raw: str) -> str:
    s = (raw or "").strip()
    if s.startswith("```"):
        parts = s.split("```")
        if len(parts) >= 2:
            s = parts[1]
            if s.lstrip().startswith("json"):
                s = s.lstrip()[4:].lstrip()
    return s.strip()


def parse_vision_profile_json(raw: str) -> Dict[str, Any]:
    """解析视觉模型返回的 JSON；失败则返回带 error 的空壳。"""
    s = _strip_json_fence(raw)
    try:
        data = json.loads(s)
        if not isinstance(data, dict):
            return {"error": "not_object", "raw_prefix": s[:200]}
        for k in ("layout_notes", "fill_strategy", "style_observations"):
            if k not in data:
                data[k] = ""
        if "chapter_hints" not in data or not isinstance(data["chapter_hints"], list):
            data["chapter_hints"] = []
        if "table_page_hints" not in data or not isinstance(data["table_page_hints"], list):
            data["table_page_hints"] = []
        return data
    except json.JSONDecodeError:
        return {"error": "json_parse", "raw_prefix": s[:400]}


def describe_template_pages_with_vision(png_pages: List[bytes]) -> Dict[str, Any]:
    """多页 PNG → 多模态 → 结构化 profile。"""
    if not config.chat_llm_configured():
        return {"error": "no_api_key"}
    if not png_pages:
        return {"error": "no_images"}

    client = config.openai_client_for_chat()

    content: List[Dict[str, Any]] = [
        {"type": "text", "text": VISION_TEMPLATE_JSON_PROMPT},
    ]
    for blob in png_pages:
        b64 = base64.standard_b64encode(blob).decode("ascii")
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"},
            }
        )

    resp = chat_completions_create(
        client,
        model=getattr(config, "TEMPLATE_VISION_MODEL", None) or config.VISION_WEB_MODEL,
        messages=[{"role": "user", "content": content}],
        temperature=config.TEMP_VISION,
        max_tokens=4096,
    )
    ch0 = resp.choices[0] if resp.choices else None
    raw = (ch0.message.content if ch0 and ch0.message else "") or ""
    return parse_vision_profile_json(raw)


def compact_profile_for_analyzer(profile: Dict[str, Any], max_chars: int = 5000) -> str:
    """供模板分析 LLM 追加的压缩文本。"""
    if profile.get("error"):
        return ""
    slim = {
        "layout_notes": (profile.get("layout_notes") or "")[:2000],
        "fill_strategy": (profile.get("fill_strategy") or "")[:1500],
        "style_observations": (profile.get("style_observations") or "")[:800],
        "chapter_hints": profile.get("chapter_hints") or [],
        "table_page_hints": (profile.get("table_page_hints") or [])[:20],
    }
    s = json.dumps(slim, ensure_ascii=False)
    if len(s) > max_chars:
        s = s[:max_chars] + "…"
    return s


def compact_profile_for_generator(profile: Dict[str, Any], max_chars: int = 900) -> str:
    """写入每个任务的 location_hint，供正文/表格生成拼接。"""
    if profile.get("error"):
        return ""
    parts = [
        (profile.get("layout_notes") or "")[:400],
        (profile.get("fill_strategy") or "")[:350],
        (profile.get("style_observations") or "")[:200],
    ]
    s = " ".join(p for p in parts if p).strip()
    if len(s) > max_chars:
        s = s[:max_chars].rstrip() + "…"
    return s


def _hint_score(chapter: str, anchor: str) -> float:
    ca = (chapter or "").strip()
    an = (anchor or "").strip()
    if not ca or not an:
        return 0.0
    if an in ca or ca in an:
        return 1.0
    if an[: min(8, len(an))] in ca:
        return 0.85
    # 去编号后比较
    ca2 = re.sub(r"^第[一二三四五六七八九十百千零〇0-9]+章\s*", "", ca)
    an2 = re.sub(r"^第[一二三四五六七八九十百千零〇0-9]+章\s*", "", an)
    if an2 and (an2 in ca2 or ca2 in an2):
        return 0.9
    hits = 0
    step = max(1, len(an) // 12)
    for i in range(0, max(len(an) - 1, 0), step):
        if an[i : i + 2] in ca:
            hits += 1
    return min(0.35, hits / max(len(an) // max(step, 1), 1))


def pick_chapter_style_hint(chapter: str, profile: Dict[str, Any]) -> str:
    hints = profile.get("chapter_hints") or []
    if not isinstance(hints, list) or not chapter:
        return ""
    best_txt = ""
    best_sc = 0.0
    for h in hints:
        if not isinstance(h, dict):
            continue
        anchor = (h.get("chapter_anchor") or h.get("chapter") or "").strip()
        hint = (h.get("hint") or h.get("hint_text") or "").strip()
        if not hint:
            continue
        sc = _hint_score(chapter, anchor)
        if sc > best_sc:
            best_sc = sc
            best_txt = hint
    if best_sc >= 0.75 and best_txt:
        return best_txt[:600]
    return ""


def apply_chapter_hints_to_tasks(tasks: List[FillTask], profile: Dict[str, Any]) -> None:
    if profile.get("error"):
        return
    gen_compact = compact_profile_for_generator(profile)
    for t in tasks:
        h = pick_chapter_style_hint(t.target_chapter, profile)
        if h:
            t.location_hint["chapter_style_hint"] = h
        if gen_compact:
            t.location_hint["template_vision_compact"] = gen_compact


def get_or_build_template_vision_profile(
    docx_path: str,
    *,
    force_refresh: bool = False,
) -> Tuple[Dict[str, Any], str]:
    """
    返回 (profile_dict, 人类可读状态说明)。
    profile 含 layout_notes / fill_strategy / style_observations / chapter_hints / table_page_hints；
    失败时含 error 键，仍可能含 ooxml_fallback。
    成功时在 cache_bundle_dir 下写入 profile.json 与 page_XXX.png。
    """
    bundle = cache_bundle_dir(docx_path)
    profile_json = _profile_path_in_bundle(bundle)
    legacy_flat = _cache_path(docx_path)

    if not force_refresh and os.path.isfile(profile_json):
        try:
            with open(profile_json, "r", encoding="utf-8") as f:
                blob = json.load(f)
            prof = blob.get("profile") or {}
            if isinstance(prof, dict) and (
                prof.get("layout_notes")
                or prof.get("chapter_hints")
                or prof.get("table_page_hints")
            ):
                if not _list_cached_page_png_paths(bundle):
                    n = ensure_template_page_pngs(docx_path)
                    extra = f" 已补写 {n} 页 PNG。" if n else " 未写入页图（请检查 PDF 导出）。"
                else:
                    extra = ""
                return prof, "已使用本地缓存的模板视觉摘要（bundle）。" + extra
        except Exception as e:
            _LOG.debug("bundle cache read fail: %s", e)

    if not force_refresh and os.path.isfile(legacy_flat):
        try:
            with open(legacy_flat, "r", encoding="utf-8") as f:
                blob = json.load(f)
            prof = blob.get("profile") or {}
            if isinstance(prof, dict) and (
                prof.get("layout_notes") or prof.get("chapter_hints")
            ):
                os.makedirs(bundle, exist_ok=True)
                try:
                    shutil.copy2(legacy_flat, profile_json)
                except OSError:
                    pass
                ensure_template_page_pngs(docx_path)
                return prof, "已使用旧版 JSON 缓存并迁移到 bundle；已尝试补写页 PNG。"
        except Exception as e:
            _LOG.debug("legacy cache read fail: %s", e)

    profile: Dict[str, Any] = {
        "layout_notes": "",
        "fill_strategy": "",
        "style_observations": "",
        "chapter_hints": [],
        "table_page_hints": [],
        "ooxml_fallback": "",
    }
    status_parts: List[str] = []
    pages_saved: List[bytes] = []

    with tempfile.TemporaryDirectory(prefix="tplvis_") as td:
        pdf_path = os.path.join(td, "tpl.pdf")
        if export_docx_to_pdf(docx_path, pdf_path):
            status_parts.append("已导出 PDF。")
            pages = pdf_to_png_pages(
                pdf_path,
                int(config.TEMPLATE_VISION_MAX_PAGES),
                float(config.TEMPLATE_VISION_ZOOM),
            )
            if pages:
                pages_saved = pages
                status_parts.append(f"已栅格化 {len(pages)} 页，调用视觉模型…")
                vision_prof = describe_template_pages_with_vision(pages)
                if not vision_prof.get("error"):
                    profile.update(vision_prof)
                    status_parts.append("视觉分析完成。")
                else:
                    status_parts.append(
                        f"视觉分析失败（{vision_prof.get('error')}），已降级为纯文本摘要。"
                    )
            else:
                status_parts.append("PDF 栅格化失败（请安装 PyMuPDF：pip install pymupdf）。")
        else:
            status_parts.append(
                "未检测到可用的 docx→PDF 工具（可安装：本机 Word + pip install docx2pdf，"
                "或安装 LibreOffice 并确保 soffice 在 PATH）。"
            )

    if profile.get("error") or not (
        profile.get("layout_notes") or profile.get("chapter_hints")
    ):
        fb = extract_ooxml_text_fallback(docx_path)
        profile["ooxml_fallback"] = fb
        if fb:
            status_parts.append("已附加 OOXML 纯文本降级摘要供分析使用。")

    substantial = (
        bool((profile.get("layout_notes") or "").strip())
        or bool((profile.get("fill_strategy") or "").strip())
        or bool(profile.get("chapter_hints"))
        or bool(profile.get("table_page_hints"))
        or len((profile.get("ooxml_fallback") or "").strip()) > 80
    )
    if substantial:
        try:
            os.makedirs(bundle, exist_ok=True)
            if pages_saved:
                _write_png_pages(bundle, pages_saved)
            elif not _list_cached_page_png_paths(bundle):
                ensure_template_page_pngs(docx_path)
            with open(profile_json, "w", encoding="utf-8") as f:
                json.dump({"profile": profile}, f, ensure_ascii=False, indent=0)
            try:
                if os.path.isfile(legacy_flat):
                    os.remove(legacy_flat)
            except OSError:
                pass
        except Exception as e:
            _LOG.warning("vision bundle cache write failed: %s", e)
    else:
        status_parts.append(
            "未写入视觉缓存（摘要过少）。安装 Word+docx2pdf 或 LibreOffice 后重试分析；"
            "或删除 data/.cache/template_vision 下旧缓存。"
        )

    return profile, " ".join(status_parts) if status_parts else "模板视觉：无额外状态。"
