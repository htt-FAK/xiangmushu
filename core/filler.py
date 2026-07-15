from typing import Any, Dict, List, Optional, Tuple
import logging
import re
from copy import deepcopy

from core.fill_intent import FillIntent

_LOG = logging.getLogger(__name__)

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


class WatermarkPreserver:
    """水印保留器：备份并恢复整个页眉/页脚 XML + 正文段落中的图片/Drawing。"""

    # VML 命名空间常量（保留用于兼容检测）
    WATERMARK_TAG = qn("w:pict")
    VML_SHAPETYPE = "{urn:schemas-microsoft-com:vml}shapetype"
    VML_SHAPE = "{urn:schemas-microsoft-com:vml}shape"
    WATERMARK_ALTERNATIVE_TAGS = [VML_SHAPETYPE, VML_SHAPE]

    @staticmethod
    def extract_watermarks(doc: Document) -> Dict[str, List[Any]]:
        """备份页眉/页脚 + 正文段落的图片/Drawing 元素。

        Returns:
            dict: {
                "headers": [...],
                "footers": [...],
                "body_drawings": [(para_index, run_index, drawing_xml), ...]
            }
        """
        watermarks: Dict[str, List[Any]] = {
            "headers": [],
            "footers": [],
            "body_drawings": [],
        }
        try:
            # 备份页眉/页脚
            for section in doc.sections:
                header = section.header
                if header is not None:
                    watermarks["headers"].append(deepcopy(header._element))
                else:
                    watermarks["headers"].append(None)

                footer = section.footer
                if footer is not None:
                    watermarks["footers"].append(deepcopy(footer._element))
                else:
                    watermarks["footers"].append(None)

            # 备份正文段落中的图片/Drawing
            for pi, para in enumerate(doc.paragraphs):
                for ri, run in enumerate(para.runs):
                    # 查找 w:drawing 元素
                    drawings = run._element.xpath(".//w:drawing")
                    for di, d in enumerate(drawings):
                        watermarks["body_drawings"].append(
                            (pi, ri, di, deepcopy(d))
                        )

            _LOG.info(
                "水印/图片备份完成: sections=%s, body_drawings=%s",
                len(watermarks["headers"]),
                len(watermarks["body_drawings"]),
            )
        except Exception as e:
            _LOG.warning("水印/图片备份失败: %s", e)
        return watermarks

    @staticmethod
    def restore_watermarks(doc: Document, watermarks: Dict[str, List[Any]]) -> None:
        """将备份的页眉/页脚 + 正文图片还原到文档中。"""
        try:
            # 还原页眉/页脚
            for i, section in enumerate(doc.sections):
                if i < len(watermarks["headers"]) and watermarks["headers"][i] is not None:
                    header_el = section.header._element
                    backup_el = watermarks["headers"][i]
                    for child in list(header_el):
                        header_el.remove(child)
                    for child in backup_el:
                        header_el.append(deepcopy(child))

                if i < len(watermarks["footers"]) and watermarks["footers"][i] is not None:
                    footer_el = section.footer._element
                    backup_el = watermarks["footers"][i]
                    for child in list(footer_el):
                        footer_el.remove(child)
                    for child in backup_el:
                        footer_el.append(deepcopy(child))

            # 还原正文图片
            for pi, ri, di, drawing_xml in watermarks["body_drawings"]:
                if pi < len(doc.paragraphs):
                    para = doc.paragraphs[pi]
                    if ri < len(para.runs):
                        run = para.runs[ri]
                        # 找到对应位置的 drawing 并替换
                        existing = run._element.xpath(".//w:drawing")
                        if di < len(existing):
                            # 替换原有 drawing
                            existing[di].addprevious(deepcopy(drawing_xml))
                            existing[di].getparent().remove(existing[di])
                        else:
                            # 追加到 run 末尾
                            run._element.append(deepcopy(drawing_xml))

            _LOG.info("页眉/页脚/正文图片已完整还原")
        except Exception as e:
            _LOG.warning("水印/图片还原失败: %s", e)

    @staticmethod
    def _is_watermark_element(element) -> bool:
        """判断元素是否为水印元素（保留用于兼容）。"""
        tag = element.tag
        if tag == WatermarkPreserver.WATERMARK_TAG:
            return True
        if tag in WatermarkPreserver.WATERMARK_ALTERNATIVE_TAGS:
            return True
        if "watermark" in str(tag).lower():
            return True
        return False

import config
from core.docx_typography import (
    apply_abstract_body_formats_in_document,
    apply_document_typography,
    apply_rPr_to_run,
    build_body_rPr,
    build_rPr_for_paragraph,
    heading_level_from_style,
    insert_paragraph_after,
    remove_empty_body_paragraphs,
    split_body_content_blocks,
)
from core.fill_task import FillTask
from core.table_context import get_column_header_from_path
from core.table_slot_expand import max_chars_for_column_header
from core.template_slots import (
    is_bracket_fill_slot,
    is_pure_hint_line as slot_is_pure_hint_line,
    looks_like_fill_instruction_line as slot_fill_instruction,
    paragraph_slot_score,
)


class WordFiller:
    """将生成内容回填到 Word：支持 {{锚点}} 精确替换 + 传统占位符。"""

    _PLACEHOLDER_PATTERNS = [
        re.compile(r"请在此填写"),
        re.compile(r"在此填写"),
        re.compile(r"请在此"),
        re.compile(r"填写.{0,8}正文"),
        re.compile(r"（\s*请[^）]{0,16}填写[^）]*）"),
        re.compile(r"请填写"),
        re.compile(r"（\s*）"),
        re.compile(r"\(\s*\)"),
        re.compile(r"_{3,}"),
        re.compile(r"待填写"),
        re.compile(r"待补充"),
        re.compile(r"此处填写"),
        re.compile(r"在以下填写"),
        re.compile(r"以下填写"),
        re.compile(r"以下空白"),
        re.compile(r"[X×]{4,}"),
        re.compile(r"请.{0,12}填写"),
    ]

    _PURE_HINT_MAX_LEN = 40
    _FILL_INSTRUCTION_MAX_LEN = 120
    _KEYWORD_LINE_PREFIXES = ("关键词", "Key words", "Keywords", "关键字")

    @staticmethod
    def _chapter_text_compact(chapter: str) -> str:
        return re.sub(r"\s+", "", chapter or "")

    @classmethod
    def _heading_matches_chapter(cls, target_chapter: str, para_text: str) -> bool:
        """章节标题与任务名比对（去空白），解决「摘要」任务无法命中「摘  要」标题的问题。"""
        ta = cls._chapter_text_compact(target_chapter)
        tb = cls._chapter_text_compact(para_text)
        if not ta or not tb:
            return False
        if ta in tb:
            return True
        if len(tb) >= 2 and tb in ta:
            return True
        return False

    @classmethod
    def _is_abstract_chapter(cls, target_chapter: str) -> bool:
        c = cls._chapter_text_compact(target_chapter).lower()
        return "摘要" in c or c == "abstract"

    @staticmethod
    def _strip_leading_abstract_label(content: str) -> str:
        """去掉模型在正文里重复的「摘要」标题行（已有独立 Heading 时避免正文段再以宋体小四重复）。"""
        if not content:
            return content
        s = content.strip()
        # 勿匹配「摘要成稿…」等正文词：仅当「摘…要」后接冒号、换行或多空格时再剥
        s = re.sub(r"^摘\s*要(\s*[:：]\s*|\s*\n+|\s{2,})", "", s)
        s = re.sub(r"^【\s*摘\s*要\s*】\s*", "", s)
        return s.strip()

    @classmethod
    def _looks_like_writing_rubric(cls, text: str) -> bool:
        """「撰写要求」+ 分条 bullet 类写作说明（无占位符），与说明段同级用于候选打分。"""
        t = (text or "").strip()
        head = t[:200]
        if len(t) < 10:
            return False
        # 已成稿的长摘要：勿当 rubric
        if len(t) > 400:
            looks_done = (
                "本项目" in t
                or "本系统" in t
                or "系统旨在" in t
                or "系统实现了" in t
            )
            has_rubric_header = bool(
                re.search(r"(撰写要求|写作说明|摘要要求)", head)
                or re.search(
                    r"用\s*\d+\s*[—\-~～]\s*\d+\s*字", head
                )
            )
            if looks_done and not has_rubric_header:
                return False
        # 明确的撰写要求标题
        if re.match(r"^\s*(撰写要求|写作说明|摘要要求)", t):
            return True
        # 独立标题行可视为 rubric；正文里提到“写作说明”不应被误清理。
        if re.search(r"(?m)^\s*(撰写要求|写作说明|摘要要求)\s*[:：]?\s*$", t[:120]):
            return True
        # 检测 bullet 点 + 关键词
        bullet_marks = "•·●◦‣⁃"
        n_bullets = sum(t.count(c) for c in bullet_marks)
        n_bullets += len(re.findall(r"(?m)^\s*[-*+－—]\s+\S", t))
        if n_bullets >= 2 and len(t) < 1200:
            rubric_kw = (
                "300",
                "500",
                "概述",
                "真实场景",
                "目标用户",
                "创新点",
                "不足",
                "核心能力",
                "工作流",
                "知识库",
                "演示",
                "插件",
                "数据库",
                "角色设定",
                "避免只写",
            )
            if sum(1 for k in rubric_kw if k in t) >= 2:
                return True
        return False

    @classmethod
    def _looks_like_example_or_hint(cls, text: str) -> bool:
        """检测表格单元格中的「例如：...」「示例：...」「请填写...」等提示文字。"""
        t = (text or "").strip()
        if not t or len(t) > 400:
            return False
        # 明显的示例/提示前缀
        if re.match(r"^\s*(例如|示例|举例|如|请填写|请描述|请说明|填写|描述|说明)[：:]", t):
            return True
        # 包含"例如"或"示例"且长度较短（说明是提示不是正文）
        if ("例如" in t or "示例" in t) and len(t) < 100:
            return True
        # 下划线占位符（如"用户1：_______"）
        if re.search(r"[_]{3,}", t) and len(t) < 50:
            return True
        # 表格中的功能说明类提示（如"基于角色设定..."、"通过工作流..."）
        hint_patterns = [
            r"基于.*，回答",
            r"通过.*完成",
            r"把.*写入或读取",
            r"调用.*完成",
            r"组合.*实现",
            r"截图[:：]",
            r"验证.*质量",
        ]
        for pattern in hint_patterns:
            if re.search(pattern, t):
                return True
        # 表格中的简短功能描述（通常是提示）
        if len(t) < 80 and ("回答" in t or "完成" in t or "写入" in t or "读取" in t):
            return True
        return False

    @classmethod
    def _looks_like_template_guidance(cls, text: str) -> bool:
        """模板「说明……建议……」或「撰写要求」类填写指引（无占位符），与占位同级、优先于空段。"""
        t = (text or "").strip()
        if not t or len(t) > 1200:
            return False

        compact = re.sub(r"\s+", "", t)

        # 图片/截图占位说明通常很短，且会包含“请在此”等占位语，必须优先识别。
        if re.match(r"^请在此处粘贴.+", t):
            return True
        if re.match(r"^图\s*\d+(?:\.\d+)*\s*.+?(截图|图片|示意图|界面图|流程图)$", t):
            return True
        if re.match(r"^图\d+(?:\.\d+)*.+?(截图|图片|示意图|界面图|流程图)$", compact):
            return True
        if re.match(r"^(截图|图片|示意图|界面图|流程图)\s*[:：]", t):
            return True

        if cls._text_has_placeholder(t):
            return False
        if cls._is_pure_hint_line(t):
            return False
        if cls._looks_like_writing_rubric(t):
            return True

        if len(t) < 8:
            return False

        guidance_patterns = (
            r"^说明.{2,}(项目来源|项目背景|现实痛点|应用价值|调用了哪些|插件|外部工具|输入参数|输出结果|工作流)",
            r"^描述.{2,}(目标用户|使用场景|核心功能|业务流程|应用场景)",
            r"^展示.{2,}(智能体|角色设定|提示词|工作流|知识库|运行效果)",
            r"^填写.{2,}(应用|入口|平台|版本|账号|权限|演示设备|必要说明)",
            r"^本节应.{2,}",
            r"^学生提交时.{2,}",
            r"^结果分析[:：].{2,}",
            r"^客观分析.{2,}",
            r"^展望.{2,}",
            r"^总分\s*\d+\s*分.{2,}",
            r"建议从.+展开",
            r"建议至少.+",
            r"建议包含.+",
            r"建议写.+",
            r"至少(给出|设计|包含|上传|展示).+",
            r"(不要|不能|避免)只写.+",
            r"应(完整)?(体现|展示|包含|说明|把).+",
            r"请在此处粘贴.+",
        )
        if any(re.search(p, t) for p in guidance_patterns):
            return True

        if len(t) < 18:
            return False
        starters = (
            "说明",
            "描述",
            "列出",
            "填写",
            "简述",
            "阐述",
            "介绍",
            "展示",
            "分析",
            "展望",
            "建议",
            "应从",
            "应把",
            "本节应",
            "需说明",
            "请围绕",
            "围绕",
        )
        if not any(t.startswith(s) for s in starters):
            return False
        markers = (
            "建议",
            "避免",
            "至少",
            "不得",
            "模板",
            "示例",
            "字数",
            "几点",
            "层次",
            "展开",
            "格式",
            "截图",
            "图片",
            "提示词",
            "工作流",
            "知识库",
            "插件",
        )
        return any(m in t for m in markers)

    @staticmethod
    def strip_markdown_light(text: str) -> str:
        if not text:
            return ""
        s = text.replace("\r\n", "\n")
        s = re.sub(r"^#{1,6}\s*", "", s, flags=re.MULTILINE)
        s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
        s = re.sub(r"\*([^*]+)\*", r"\1", s)
        s = re.sub(r"^\s*[-*+]\s+", "", s, flags=re.MULTILINE)
        s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
        return s.strip()

    @classmethod
    def clean_table_answer(
        cls,
        text: str,
        word_limit: int = 120,
        max_chars: Optional[int] = None,
    ) -> str:
        """对表格单元格答案做深度清洗：去前缀、占位/下划线噪声、换行压平、超长截断。"""
        if not text:
            return ""
        s = text.strip()

        _TABLE_PREFIXES = [
            r"^(答案|答|填写内容|该单元格应填写|该格填写|应填写|内容)：?\s*",
            r"^(根据参考资料[，,]?|根据上述资料[，,]?|根据检索片段[，,]?)",
            r"^(综上所述[，,]?|综合来看[，,]?)",
            r"^(以下是|如下|答：)\s*",
        ]
        for pattern in _TABLE_PREFIXES:
            s = re.sub(pattern, "", s, flags=re.IGNORECASE).strip()

        s = re.sub(r"资料\s*\d+\s*[：:]\s*_+", "", s)
        s = re.sub(r"_{4,}", "", s)
        s = re.sub(r"[\r\n]+", " ", s)
        s = re.sub(r"\s{2,}", " ", s).strip()

        if cls._looks_like_template_guidance(s):
            return ""

        parts = re.split(r"(?<=[。！？!?；;])\s*", s)
        kept = [p.strip() for p in parts if p.strip() and not cls._looks_like_template_guidance(p)]
        if kept:
            s = "".join(kept).strip()
        elif parts:
            s = ""

        cap = max_chars if max_chars is not None else max(20, int(word_limit * 1.5))
        if (word_limit or 120) <= 45:
            cap = min(cap, 60)
        if len(s) > cap:
            s = s[:cap].rstrip("，。,. ")
        return s

    @classmethod
    def _text_has_placeholder(cls, text: str) -> bool:
        if not text:
            return False
        return any(p.search(text) for p in cls._PLACEHOLDER_PATTERNS)

    @classmethod
    def _is_pure_hint_line(cls, text: str) -> bool:
        """整段几乎全是填写指引（含【请在此填写…】）。"""
        if is_bracket_fill_slot(text):
            return True
        return slot_is_pure_hint_line(text)

    @classmethod
    def _looks_like_fill_instruction_line(cls, text: str) -> bool:
        """申报模板口语指引（如「摘要：在以下填写…」）。"""
        if is_bracket_fill_slot(text):
            return True
        if slot_is_pure_hint_line(text):
            return True
        return slot_fill_instruction(text)

    @staticmethod
    def _hint_wants_full_replace(hint: Dict[str, Any]) -> bool:
        rm = (hint.get("replace_mode") or "").strip().lower()
        fs = (hint.get("fill_strategy") or "").strip().lower()
        return rm == "full" or fs in ("full", "full_replace")

    @classmethod
    def _classify_scope_paragraph(cls, text: str) -> str:
        t = (text or "").strip()
        if not t:
            return "empty"
        if any(t.startswith(p) for p in cls._KEYWORD_LINE_PREFIXES):
            return "keyword"
        if cls._looks_like_writing_rubric(t) or cls._looks_like_template_guidance(t):
            return "rubric"
        if (
            cls._looks_like_fill_instruction_line(t)
            or cls._is_pure_hint_line(t)
            or cls._text_has_placeholder(t)
        ):
            return "hint"
        return "other"

    def _clear_non_body_scope_paragraphs(
        self,
        doc: Document,
        scope: List[int],
        body_idx: int,
    ) -> None:
        paras = doc.paragraphs
        for idx in scope:
            if idx == body_idx:
                continue
            kind = self._classify_scope_paragraph(paras[idx].text or "")
            if kind in ("hint", "rubric", "empty"):
                self._set_paragraph_text_keep_style(paras[idx], "")

    @staticmethod
    def _table_index_for_element(doc: Document, tbl_el) -> Optional[int]:
        for i, t in enumerate(doc.tables):
            if t._element is tbl_el:
                return i
        return None

    def _collect_chapter_region(
        self, doc: Document, target_chapter: str
    ) -> Tuple[int, List[int], List[int]]:
        """按 body 顺序返回 (标题段落下标, 本章段落下标, 本章表格 doc.tables 下标)。

        Normal-style fallback: 当所有段落均无 Heading style 时，调用
        normal_heading_detector 识别加粗编号段落作为章节边界。
        """
        paras = doc.paragraphs
        start_idx = -1
        chapter_lvl: Optional[int] = 1
        para_scope: List[int] = []
        table_scope: List[int] = []
        in_chapter = False
        para_idx = 0

        # ── 优先扫描 Heading style ──────────────────────────────────────
        has_any_heading_style = any(
            self._para_heading_level(p) is not None for p in paras
        )

        # ── Heading style 全缺时启用 Normal-style heading 检测 ──────────
        import config as cfg
        normal_headings: Dict[int, int] = {}
        if not has_any_heading_style and getattr(cfg, "APPLY_TEMPLATE_STYLE", True):
            try:
                from core.normal_heading_detector import find_all_headings
                threshold = int(getattr(cfg, "NORMAL_HEADING_THRESHOLD", 50))
                normal_headings = {
                    idx: lvl for idx, lvl in find_all_headings(doc, threshold)
                }
                _LOG.info(
                    "_collect_chapter_region: Normal-heading fallback, found %d headings",
                    len(normal_headings),
                )
            except Exception as exc:
                _LOG.warning("normal_heading_detector failed: %s", exc)

        def _effective_heading_level(p: Paragraph) -> Optional[int]:
            lvl = self._para_heading_level(p)
            if lvl is not None:
                return lvl
            # Normal-heading fallback lookup
            try:
                idx = paras.index(p)
                return normal_headings.get(idx)
            except ValueError:
                return None

        for child in doc.element.body:
            if child.tag == qn("w:p"):
                if para_idx >= len(paras):
                    break
                para = paras[para_idx]
                t = (para.text or "").strip()
                lvl = _effective_heading_level(para)
                if not in_chapter:
                    if target_chapter and self._heading_matches_chapter(
                        target_chapter, t
                    ):
                        in_chapter = True
                        start_idx = para_idx
                        chapter_lvl = lvl or 1
                else:
                    if (
                        lvl is not None
                        and lvl <= (chapter_lvl or 1)
                        and t
                    ):
                        return start_idx, para_scope, table_scope
                    para_scope.append(para_idx)
                para_idx += 1
            elif child.tag == qn("w:tbl") and in_chapter:
                ti = self._table_index_for_element(doc, child)
                if ti is not None:
                    table_scope.append(ti)

        if not in_chapter:
            return -1, [], []
        return start_idx, para_scope, table_scope

    @classmethod
    def _classify_scope_cell(cls, text: str) -> str:
        t = (text or "").strip()
        if not t:
            return "empty"
        if any(t.startswith(p) for p in cls._KEYWORD_LINE_PREFIXES):
            return "keyword"
        if cls._looks_like_writing_rubric(t) or cls._looks_like_template_guidance(
            t
        ):
            return "rubric"
        if (
            cls._looks_like_fill_instruction_line(t)
            or cls._is_pure_hint_line(t)
            or cls._text_has_placeholder(t)
        ):
            return "hint"
        return "other"

    def _clear_abstract_non_body(
        self,
        doc: Document,
        para_scope: List[int],
        table_scope: List[int],
        body_para_idx: int,
        body_table: Optional[Tuple[int, int, int]],
    ) -> None:
        """摘要章：清空非正文槽的段落与表内 rubric/提示单元格。"""
        paras = doc.paragraphs
        for idx in para_scope:
            if idx == body_para_idx:
                continue
            kind = self._classify_scope_paragraph(paras[idx].text or "")
            if kind in ("hint", "rubric", "empty"):
                self._set_paragraph_text_keep_style(paras[idx], "")

        body_ti, body_r, body_c = body_table if body_table else (-1, -1, -1)
        for ti in table_scope:
            if ti >= len(doc.tables):
                continue
            table = doc.tables[ti]
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if ti == body_ti and ri == body_r and ci == body_c:
                        continue
                    kind = self._classify_scope_cell(cell.text or "")
                    if kind in ("hint", "rubric", "empty"):
                        self._set_cell_text_keep_style(cell, "")

    def _sweep_abstract_chapter_table_rubrics(self, doc: Document) -> None:
        """回填结束后：清除摘要章范围内表内残留的撰写要求/提示。"""
        for i, para in enumerate(doc.paragraphs):
            t = (para.text or "").strip()
            if not self._heading_matches_chapter("摘要", t):
                continue
            _start, _ps, table_scope = self._collect_chapter_region(doc, t)
            if _start < 0:
                continue
            for ti in table_scope:
                if ti >= len(doc.tables):
                    continue
                for row in doc.tables[ti].rows:
                    for cell in row.cells:
                        txt = cell.text or ""
                        if (
                            self._looks_like_writing_rubric(txt)
                            or self._looks_like_template_guidance(txt)
                            or self._looks_like_fill_instruction_line(txt)
                            or self._is_pure_hint_line(txt)
                        ):
                            self._set_cell_text_keep_style(cell, "")
            break

    def _remove_chapter_rubric_tables(
        self,
        doc: Document,
        table_scope: List[int],
        body_table: Optional[Tuple[int, int, int]] = None,
    ) -> None:
        """删除章内单行单列且为 rubric/提示/已空的表格（去掉撰写要求框）。"""
        body_ti, body_r, body_c = body_table if body_table else (-1, -1, -1)
        for ti in sorted(table_scope, reverse=True):
            if ti >= len(doc.tables):
                continue
            if ti == body_ti:
                continue
            table = doc.tables[ti]
            if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
                continue
            cell = table.rows[0].cells[0]
            txt = cell.text or ""
            kind = self._classify_scope_cell(txt)
            if kind not in ("rubric", "hint", "empty") and txt.strip():
                continue
            tbl_el = table._tbl
            parent = tbl_el.getparent()
            if parent is not None:
                parent.remove(tbl_el)

    def _remove_rubric_tables_in_abstract_chapters(self, doc: Document) -> None:
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if not (
                re.match(r"^摘\s*要\s*$", t)
                or self._heading_matches_chapter("摘要", t)
            ):
                continue
            _start, _ps, table_scope = self._collect_chapter_region(doc, t)
            if _start < 0 or not table_scope:
                continue
            self._remove_chapter_rubric_tables(doc, table_scope, None)

    @staticmethod
    def _para_heading_level(para: Paragraph) -> Optional[int]:
        if not para.style:
            return None
        return heading_level_from_style(para.style.name or "")

    def __init__(self) -> None:
        # 三级合并后的样式档案，由 fill_template() 在运行前注入
        # APPLY_TEMPLATE_STYLE=True 时生效，False 时保留旧逻辑
        self._merged_profile: Optional[Any] = None

    def fill_template(
        self,
        template_path: str,
        tasks: List[FillTask],
        contents: List[str],
        output_path: str,
        merged_profile: Optional[Any] = None,
    ):
        self._merged_profile = merged_profile
        doc = Document(template_path)

        # 提取水印（如果启用）
        watermarks = None
        if getattr(config, "PRESERVE_WATERMARK", True):
            try:
                watermarks = WatermarkPreserver.extract_watermarks(doc)
                _LOG.info("已提取水印，准备保留")
            except Exception as e:
                _LOG.warning("水印提取失败，继续处理: %s", e)

        for task, raw in zip(tasks, contents):
            content = self.strip_markdown_light(raw)
            if task.task_type == "paragraph" and self._is_abstract_chapter(
                task.target_chapter or ""
            ):
                content = self._strip_leading_abstract_label(content)
            anchor = task.location_hint.get("anchor")
            if anchor:
                self._replace_anchor_everywhere(doc, anchor, content)
                continue
            if task.task_type == "table_cell":
                wl = int(task.word_limit or 80)
                loc = task.location_hint or {}
                hdr = get_column_header_from_path(
                    template_path,
                    int(loc.get("table_index", 0)),
                    int(loc.get("col", 0)),
                )
                cap = max_chars_for_column_header(hdr)
                content = self.clean_table_answer(
                    content, wl, max_chars=cap
                )
                self._fill_table_cell(doc, task, content)
            else:
                self._fill_paragraph(doc, task, content)

        self._sweep_residual_hint_paragraphs(doc)
        self._sweep_residual_hint_table_cells(doc)
        self._sweep_abstract_chapter_table_rubrics(doc)
        self._remove_rubric_tables_in_abstract_chapters(doc)
        remove_empty_body_paragraphs(doc)

        # 恢复水印（在排版调整之前，确保页眉/页脚 XML 完整还原）
        if watermarks is not None and getattr(config, "PRESERVE_WATERMARK", True):
            try:
                WatermarkPreserver.restore_watermarks(doc, watermarks)
                _LOG.info("已恢复水印到输出文档")
            except Exception as e:
                _LOG.warning("水印恢复失败: %s", e)

        if getattr(config, "ADJUST_TABLE_READABILITY", True):
            for table in doc.tables:
                self._ensure_table_readability(table)

        # ── 排版：智能样式（doc-gen-revamp）vs 旧硬编码 ─────────────────────
        from core.smart_style import apply_smart_style_to_document, should_use_smart_style
        if should_use_smart_style() and self._merged_profile is not None:
            apply_smart_style_to_document(doc, self._merged_profile)
        elif getattr(config, "APPLY_UNIFIED_TYPOGRAPHY", True):
            apply_document_typography(doc)
            apply_abstract_body_formats_in_document(doc)

        self._merged_profile = None  # 清理实例属性，避免泄漏到下次调用
        doc.save(output_path)

    def _replace_anchor_everywhere(self, doc: Document, anchor: str, content: str):
        for para in doc.paragraphs:
            self._replace_once_in_paragraph(para, anchor, content)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_once_in_paragraph(para, anchor, content)

    @staticmethod
    def _sample_rPr_from_paragraph(para: Paragraph):
        for r in para.runs:
            if (r.text or "").strip() and r._r.rPr is not None:
                return deepcopy(r._r.rPr)
        for r in para.runs:
            if r._r.rPr is not None:
                return deepcopy(r._r.rPr)
        return None

    @staticmethod
    def _clear_cell_body_keep_tcPr(cell) -> None:
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                continue
            tc.remove(child)

    def _set_paragraph_text_keep_style(self, para: Paragraph, text: str) -> None:
        """写入整段文本并应用统一规格（按段落样式分档），保留段落中的图片/Drawing。

        smart_style 模式：从 self._merged_profile 中选取 RunStyle 构造 rPr（模板提取字体）。
        回退模式：使用 build_rPr_for_paragraph(para)（硬编码宋体）。
        """
        if getattr(self, "_merged_profile", None) is not None:
            from core.smart_style import _build_rPr_from_runstyle

            lvl = self._para_heading_level(para)
            if lvl is not None:
                rs = self._merged_profile.heading_style_for_level(lvl)
            else:
                rs = self._merged_profile.body_style
            rpr = _build_rPr_from_runstyle(rs)
        else:
            rpr = build_rPr_for_paragraph(para)
        
        # 备份段落中的 drawing 元素（图片/WordArt）
        drawings_backup = []
        for r in para.runs:
            drawings = r._element.xpath(".//w:drawing")
            for d in drawings:
                drawings_backup.append((r, deepcopy(d)))
        
        # 清空文本但保留 drawing：只删除 w:t 节点
        for r in para.runs:
            for child in list(r._r):
                if child.tag == qn("w:t"):
                    r._r.remove(child)
        
        if para.runs:
            run = para.runs[0]
            # 添加新文本节点（插入到 drawing 之前）
            t_elem = run._r.makeelement(qn("w:t"))
            t_elem.text = text
            drawings = run._r.xpath(".//w:drawing")
            if drawings:
                drawings[0].addprevious(t_elem)
            else:
                run._r.insert(0, t_elem)
            apply_rPr_to_run(run, rpr)
        else:
            nr = para.add_run(text)
            apply_rPr_to_run(nr, rpr)
        
        # 如果原段落有 drawing 但当前 run 没有，追加回去
        for orig_run, d_xml in drawings_backup:
            if orig_run in para.runs:
                existing = orig_run._element.xpath(".//w:drawing")
                if not existing:
                    orig_run._element.append(d_xml)

    def _set_paragraph_blocks_keep_style(
        self, anchor_para: Paragraph, content: str
    ) -> Paragraph:
        """多段正文写入：首段替换锚点段，后续段紧接插入，不留空行。"""
        blocks = split_body_content_blocks(content)
        self._set_paragraph_text_keep_style(anchor_para, blocks[0])
        last = anchor_para
        for block in blocks[1:]:
            np = insert_paragraph_after(last)
            self._set_paragraph_text_keep_style(np, block)
            last = np
        return last

    def _replace_cell_text_preserve_format(self, cell, text: str) -> None:
        """替换单元格文本但保留原有格式（字体、颜色、大小等）。"""
        try:
            paragraphs = cell.paragraphs
            if not paragraphs:
                # 没有段落，回退到标准方法
                self._set_cell_text_keep_style(cell, text)
                return

            first_para = paragraphs[0]
            # 保留第一个 run 的格式
            preserved_rPr = None
            for run in first_para.runs:
                if run._r.rPr is not None:
                    preserved_rPr = deepcopy(run._r.rPr)
                    break

            # 清空所有 run 的文本
            for run in first_para.runs:
                run.text = ""

            # 在第一个 run 写入新文本
            if first_para.runs:
                first_run = first_para.runs[0]
                first_run.text = text or ""
                if preserved_rPr is not None:
                    first_run._r.insert(0, preserved_rPr)
            else:
                run = first_para.add_run(text or "")
                if preserved_rPr is not None:
                    run._r.insert(0, preserved_rPr)

            # 删除多余的段落（保留第一个）
            tc = cell._tc
            p_elements = [child for child in list(tc) if child.tag == qn("w:p")]
            for p_elem in p_elements[1:]:
                tc.remove(p_elem)

            _LOG.debug("表格单元格格式保留成功")
        except Exception as e:
            _LOG.warning("表格格式保留失败，回退到标准方法: %s", e)
            self._set_cell_text_keep_style(cell, text)

    @staticmethod
    def _sample_cell_rPr(cell):
        """采样单元格中第一个有内容的 run 的 rPr，用于保留原始字体/字号。
        如果 rPr 缺少 ascii/hAnsi 字体，从 run 的实际解析值或样式中补全。"""
        from docx.oxml.ns import qn

        def _enrich_rPr(rpr, run):
            """补全 rPr 中缺失的字体字段。"""
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.insert(0, rf)
            # 补全 ascii/hAnsi
            for attr in ["w:ascii", "w:hAnsi"]:
                if rf.get(qn(attr)) is None:
                    val = run.font.name
                    if val:
                        rf.set(qn(attr), val)
            # 补全 eastAsia
            if rf.get(qn("w:eastAsia")) is None:
                try:
                    style_el = getattr(run.style, "element", None) if run.style else None
                    if style_el is not None:
                        srpr = style_el.find(qn("w:rPr"))
                        if srpr is not None:
                            srf = srpr.find(qn("w:rFonts"))
                            if srf is not None:
                                ea = srf.get(qn("w:eastAsia"))
                                if ea:
                                    rf.set(qn("w:eastAsia"), ea)
                except Exception:
                    pass
            return rpr

        for para in cell.paragraphs:
            for run in para.runs:
                if (run.text or "").strip() and run._r.rPr is not None:
                    return _enrich_rPr(deepcopy(run._r.rPr), run)
        for para in cell.paragraphs:
            for run in para.runs:
                if run._r.rPr is not None:
                    return _enrich_rPr(deepcopy(run._r.rPr), run)
        # 无 rPr 时尝试从样式获取
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.name or run.font.size:
                    rpr = build_body_rPr()
                    return _enrich_rPr(rpr, run)
        return None

    def _set_cell_text_keep_style(self, cell, text: str) -> None:
        """清空单元格正文但保留 tcPr，写入单段并优先保留原始字体/字号。

        smart_style 模式：当单元格无原始 rPr 时，使用 self._merged_profile.table_cell_style 作为回退。
        回退模式：使用 build_body_rPr()（硬编码宋体）。
        """
        # 先采样原始格式
        preserved_rPr = self._sample_cell_rPr(cell)
        self._clear_cell_body_keep_tcPr(cell)
        p = cell.add_paragraph()
        run = p.add_run(text or "")
        if preserved_rPr is not None:
            apply_rPr_to_run(run, preserved_rPr)
        elif getattr(self, "_merged_profile", None) is not None:
            from core.smart_style import _build_rPr_from_runstyle

            rpr = _build_rPr_from_runstyle(self._merged_profile.table_cell_style)
            apply_rPr_to_run(run, rpr)
        else:
            apply_rPr_to_run(run, build_body_rPr())

    def _replace_once_in_paragraph(self, para: Paragraph, anchor: str, content: str):
        if anchor not in para.text:
            return
        merged = para.text.replace(anchor, content, 1)
        self._set_paragraph_text_keep_style(para, merged)

    def _set_paragraph_plain(self, para: Paragraph, text: str):
        self._set_paragraph_text_keep_style(para, text)

    @staticmethod
    def _first_placeholder_span(
        text: str, anchor: Optional[str] = None
    ) -> Optional[Tuple[int, int]]:
        for pat in WordFiller._PLACEHOLDER_PATTERNS:
            m = pat.search(text)
            if m:
                return (m.start(), m.end())
        if anchor:
            a = str(anchor).strip()
            if a and a in text:
                i = text.index(a)
                return (i, i + len(a))
        return None

    def _fill_paragraph_placeholder_only(
        self, para: Paragraph, content: str, hint: Dict[str, Any]
    ) -> bool:
        full = para.text or ""
        anchor = hint.get("anchor")
        anchor_s = str(anchor).strip() if anchor else None
        span = self._first_placeholder_span(full, anchor=anchor_s or None)
        if span is None:
            return False
        start, end = span
        new_text = full[:start] + (content or "") + full[end:]
        self._set_paragraph_text_keep_style(para, new_text)
        return True

    def _collect_chapter_scope(
        self, doc: Document, target_chapter: str
    ) -> Tuple[int, List[int]]:
        """返回 (章节标题段落下标, 本章内后续段落下标列表)。"""
        start_idx, para_scope, _ = self._collect_chapter_region(doc, target_chapter)
        return start_idx, para_scope

    def _score_paragraph_candidate(
        self, para: Paragraph, para_text_hint: str
    ) -> int:
        return paragraph_slot_score(
            para.text or "",
            para_text_hint,
            writing_rubric_fn=self._looks_like_writing_rubric,
            template_guidance_fn=self._looks_like_template_guidance,
        )

    def _score_table_cell_candidate(
        self, cell_text: str, para_text_hint: str
    ) -> int:
        return paragraph_slot_score(
            cell_text or "",
            para_text_hint,
            writing_rubric_fn=self._looks_like_writing_rubric,
            template_guidance_fn=self._looks_like_template_guidance,
        )

    def _find_best_table_body_slot(
        self, doc: Document, table_scope: List[int], para_text_hint: str
    ) -> Tuple[int, Tuple[int, int, int], int]:
        """返回 (score, (table_index, row, col), table_index)。"""
        best_score = 0
        best_slot: Tuple[int, int, int] = (0, 0, 0)
        best_ti = -1
        for ti in table_scope:
            if ti >= len(doc.tables):
                continue
            table = doc.tables[ti]
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    sc = self._score_table_cell_candidate(
                        cell.text or "", para_text_hint
                    )
                    if sc > best_score:
                        best_score = sc
                        best_slot = (ti, ri, ci)
                        best_ti = ti
        return best_score, best_slot, best_ti

    def _write_table_cell_content(
        self, cell, content: str, hint: Dict[str, Any]
    ) -> None:
        cell_text = cell.text or ""
        if (
            self._hint_wants_full_replace(hint)
            or self._looks_like_writing_rubric(cell_text)
            or self._looks_like_template_guidance(cell_text)
            or self._looks_like_fill_instruction_line(cell_text)
            or self._is_pure_hint_line(cell_text)
        ):
            self._set_cell_text_keep_style(cell, content)
            return
        self._set_cell_text_keep_style(cell, content)

    def _clear_residual_hint_paragraph(self, para: Paragraph) -> None:
        t = para.text or ""
        if (
            is_bracket_fill_slot(t)
            or self._is_pure_hint_line(t)
            or self._looks_like_fill_instruction_line(t)
        ):
            self._set_paragraph_text_keep_style(para, "")

    @staticmethod
    def _clear_paragraph_text_preserve_structure(para: Paragraph) -> None:
        """清除可见文本，保留段落、run、分页符、分节符、图片/水印等非文本结构。"""
        for node in para._element.xpath(".//w:t | .//w:instrText | .//w:delText"):
            node.text = ""

    def _clear_cell_text_preserve_structure(self, cell) -> None:
        """清除单元格可见文本，保留单元格属性、段落结构、图片和分页符。"""
        for para in cell.paragraphs:
            self._clear_paragraph_text_preserve_structure(para)

    def _clear_adjacent_pure_hints(self, doc: Document, filled_idx: int) -> None:
        paras = doc.paragraphs
        for j in (filled_idx - 1, filled_idx + 1):
            if 0 <= j < len(paras):
                self._clear_residual_hint_paragraph(paras[j])

    def _sweep_residual_hint_paragraphs(self, doc: Document) -> None:
        """邻接清理未覆盖的独立提示行，回填结束后再扫一遍正文段落。
        跳过封面段落和评分表段落。"""
        for para in doc.paragraphs:
            text = para.text or ""
            # 跳过封面段落
            if self._is_cover_paragraph(text):
                continue
            if self._looks_like_template_guidance(text):
                self._clear_paragraph_text_preserve_structure(para)
                continue
            # 跳过包含评分表关键词的段落
            if any(kw in text for kw in ["评分", "评价", "打分", "得分", "总分", "等级"]):
                continue
            self._clear_residual_hint_paragraph(para)

    def _sweep_residual_hint_table_cells(self, doc: Document) -> None:
        """清理未被任务覆盖的表格模板指引，保留表格结构和非文本对象。"""
        for table in doc.tables:
            if self._is_cover_table(table) or self._is_rating_table(table):
                continue
            for row in table.rows:
                for cell in row.cells:
                    text = (cell.text or "").strip()
                    if not text:
                        continue
                    if (
                        self._looks_like_template_guidance(text)
                        or self._looks_like_fill_instruction_line(text)
                        or self._is_pure_hint_line(text)
                        or self._looks_like_example_or_hint(text)
                    ):
                        self._clear_cell_text_preserve_structure(cell)

    def _write_paragraph_content(
        self, para: Paragraph, content: str, hint: Dict[str, Any]
    ) -> None:
        text = para.text or ""
        mode = (hint.get("replace_mode") or "").strip().lower()

        if is_bracket_fill_slot(text):
            blocks = split_body_content_blocks(content)
            if len(blocks) <= 1:
                self._set_paragraph_text_keep_style(para, blocks[0])
            else:
                self._set_paragraph_blocks_keep_style(para, content)
            return

        if self._hint_wants_full_replace(hint) and (
            self._looks_like_fill_instruction_line(text)
            or self._is_pure_hint_line(text)
            or self._looks_like_writing_rubric(text)
            or self._looks_like_template_guidance(text)
        ):
            self._write_body_content_to_paragraph(para, content)
            return

        # 检查是否是"撰写要求"类模板说明文字（需要完全替换）
        if self._looks_like_writing_rubric(text):
            self._write_body_content_to_paragraph(para, content)
            return

        # 检查是否是"说明...建议..."类模板指引文字
        if self._looks_like_template_guidance(text):
            self._write_body_content_to_paragraph(para, content)
            return

        if self._is_pure_hint_line(text) or self._looks_like_fill_instruction_line(
            text
        ):
            self._write_body_content_to_paragraph(para, content)
            return

        if mode == "placeholder_only":
            if not self._fill_paragraph_placeholder_only(para, content, hint):
                self._set_paragraph_text_keep_style(para, content)
            return

        # 混排说明+占位：无显式 full 时先尝试只换占位
        if self._text_has_placeholder(text) and len(text) > self._PURE_HINT_MAX_LEN:
            if "说明" in text or len(text) > 60:
                if self._fill_paragraph_placeholder_only(para, content, hint):
                    return

        self._write_body_content_to_paragraph(para, content)

    def _write_body_content_to_paragraph(
        self, para: Paragraph, content: str
    ) -> None:
        """写入正文：多段时用连续段落、无空行。"""
        blocks = split_body_content_blocks(content)
        if len(blocks) <= 1:
            self._set_paragraph_text_keep_style(para, blocks[0])
        else:
            self._set_paragraph_blocks_keep_style(para, content)

    @staticmethod
    def _is_cover_paragraph(text: str) -> bool:
        """检测段落是否为封面内容（不应被填充）。"""
        if not text:
            return False
        cover_keywords = [
            "广东理工学院", "结课报告", "课程报告", "实验报告",
            "作品名称", "应用平台", "信息技术学院",
            "人工智能与大模型技术应用微专业", "微专业 01 班",
            "学号", "姓名", "总分", "任课教师",
            "2026", "年", "月", "日",
        ]
        matched = sum(1 for kw in cover_keywords if kw in text)
        return matched >= 2

    def _fill_paragraph(self, doc: Document, task: FillTask, content: str):
        hint = task.location_hint or {}
        para_text_hint = (hint.get("paragraph_text") or "").strip()
        paras = doc.paragraphs

        if not task.target_chapter:
            # 无章节：按原逻辑找第一个占位/空段
            for para in paras:
                text = para.text.strip()
                # 跳过封面段落
                if self._is_cover_paragraph(text):
                    _LOG.debug("跳过封面段落: %s", text[:50])
                    continue
                if (
                    self._text_has_placeholder(text)
                    or self._looks_like_fill_instruction_line(text)
                    or (para_text_hint and para_text_hint in (para.text or ""))
                ):
                    self._write_paragraph_content(para, content, hint)
                    return
                if not text:
                    self._write_paragraph_content(para, content, hint)
                    return
            return

        is_abstract = self._is_abstract_chapter(task.target_chapter or "")
        if is_abstract:
            _start, para_scope, table_scope = self._collect_chapter_region(
                doc, task.target_chapter
            )
        else:
            _start, para_scope = self._collect_chapter_scope(
                doc, task.target_chapter
            )
            table_scope = []

        if not para_scope and not (is_abstract and table_scope):
            return

        best_idx = -1
        best_score = 0
        for idx in para_scope:
            sc = self._score_paragraph_candidate(paras[idx], para_text_hint)
            if is_abstract and self._classify_scope_paragraph(paras[idx].text or "") == "empty":
                sc = max(sc, 34)
            if sc > best_score:
                best_score = sc
                best_idx = idx

        table_score = 0
        table_slot: Tuple[int, int, int] = (-1, -1, -1)
        if is_abstract and table_scope:
            table_score, table_slot, _ = self._find_best_table_body_slot(
                doc, table_scope, para_text_hint
            )

        body_table: Optional[Tuple[int, int, int]] = None
        use_table = False
        if is_abstract:
            # 摘要：只要有可写段落（含空段），正文写在段落；表内撰写要求只清空
            if best_score >= 1:
                use_table = False
            elif table_score > 0:
                use_table = True
            else:
                return
        elif best_score < 10:
            return

        if use_table and table_slot[0] >= 0:
            ti, ri, ci = table_slot
            if ti < len(doc.tables):
                table = doc.tables[ti]
                if ri < len(table.rows) and ci < len(table.rows[ri].cells):
                    cell = table.rows[ri].cells[ci]
                    self._write_table_cell_content(cell, content, hint)
                    body_table = table_slot
                    try:
                        table.autofit = False
                    except Exception:
                        pass
        elif best_idx >= 0:
            target_para = paras[best_idx]
            self._write_paragraph_content(target_para, content, hint)
        else:
            return

        if is_abstract:
            body_para = best_idx if not use_table else -1
            if body_para >= 0 or body_table:
                self._clear_abstract_non_body(
                    doc,
                    para_scope,
                    table_scope,
                    body_para,
                    body_table,
                )
        elif best_idx >= 0 and best_score < 40:
            self._clear_adjacent_pure_hints(doc, best_idx)

    @staticmethod
    def _is_cover_table(table) -> bool:
        """检测表格是否为封面表格（包含作品名称、学号、姓名等封面信息）。"""
        try:
            # 提取表格所有文本
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text or ""

            # 封面表格关键词
            cover_keywords = [
                "作品名称", "应用平台", "学院", "专业", "班级",
                "学号", "姓名", "总分", "任课教师", "日期",
                "结课报告", "课程报告", "实验报告",
            ]

            # 如果包含多个封面关键词，则认为是封面表格
            matched = sum(1 for kw in cover_keywords if kw in table_text)
            return matched >= 3
        except Exception:
            return False

    @staticmethod
    def _is_rating_table(table) -> bool:
        """检测表格是否为评分表/评价表。"""
        try:
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text or ""

            rating_keywords = [
                "评分", "评价", "打分", "得分", "总分", "等级",
                "优秀", "良好", "中等", "及格", "不及格",
            ]

            matched = sum(1 for kw in rating_keywords if kw in table_text)
            return matched >= 2
        except Exception:
            return False

    def _fill_table_cell(self, doc: Document, task: FillTask, content: str):
        hint = task.location_hint or {}

        # --- fill_intent early skip (task 4.4) ---
        fill_intent_val = hint.get("fill_intent")
        if fill_intent_val is not None:
            # Normalize: accept FillIntent enum or raw str value
            intent_str = (
                fill_intent_val.value
                if isinstance(fill_intent_val, FillIntent)
                else str(fill_intent_val)
            )
            if intent_str != FillIntent.FILL.value:
                _LOG.debug(
                    "_fill_table_cell: skip cell, fill_intent=%r (not FILL)",
                    intent_str,
                )
                return

        table_idx = hint.get("table_index", 0)
        row_idx = hint.get("row", 0)
        col_idx = hint.get("col", 0)

        if table_idx >= len(doc.tables):
            _LOG.error(
                "fill_table_cell table_index=%s out of range n_tables=%s chapter=%r",
                table_idx,
                len(doc.tables),
                task.target_chapter,
            )
            return
        table = doc.tables[table_idx]

        # 保护封面表格和评分表
        if self._is_cover_table(table):
            _LOG.info("检测到封面表格，跳过填充 (table_idx=%s)", table_idx)
            return
        if self._is_rating_table(table):
            _LOG.info("检测到评分表，跳过填充 (table_idx=%s)", table_idx)
            return

        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return
        cell = row.cells[col_idx]

        cell_text = cell.text or ""

        # 检查是否是"例如：..."等示例/提示文字，如果是则直接替换
        if self._looks_like_example_or_hint(cell_text):
            if getattr(config, "PRESERVE_TABLE_FORMAT", True):
                self._replace_cell_text_preserve_format(cell, content)
            else:
                self._set_cell_text_keep_style(cell, content)
            try:
                table.autofit = False
            except Exception:
                pass
            return

        mode = (hint.get("replace_mode") or "").strip().lower()
        if mode == "placeholder_only" and self._text_has_placeholder(cell_text):
            span = self._first_placeholder_span(cell_text)
            if span:
                start, end = span
                new_text = cell_text[:start] + content + cell_text[end:]
                if getattr(config, "PRESERVE_TABLE_FORMAT", True):
                    self._replace_cell_text_preserve_format(cell, new_text)
                else:
                    self._set_cell_text_keep_style(cell, new_text)
                try:
                    table.autofit = False
                except Exception:
                    pass
                return

        if getattr(config, "PRESERVE_TABLE_FORMAT", True):
            self._replace_cell_text_preserve_format(cell, content)
        else:
            self._set_cell_text_keep_style(cell, content)
        try:
            table.autofit = False
        except Exception:
            pass

    def _ensure_table_readability(self, table) -> None:
        try:
            rows = table.rows
            if not rows:
                return
            ncols = len(rows[0].cells)
            if ncols <= 0:
                return
        except Exception:
            return

        try:
            table.autofit = False
        except Exception:
            pass

        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tl = tblPr.find(qn("w:tblLayout"))
        if tl is None:
            tl = OxmlElement("w:tblLayout")
            tl.set(qn("w:type"), "fixed")
            tblPr.append(tl)
        else:
            tl.set(qn("w:type"), "fixed")

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # ---- Column width handling ----
        preserve = getattr(config, "PRESERVE_ORIGINAL_COLUMN_WIDTHS", True)
        if not preserve:
            # OLD BEHAVIOR: equalize column widths
            try:
                usable = 6.35
                col_w = Inches(usable / ncols)
                for row in rows:
                    for cell in row.cells:
                        try:
                            cell.width = col_w
                        except Exception:
                            pass
            except Exception:
                pass

        # ---- Minimum column width guard (Task 6.2 + 6.3) ----
        try:
            min_width = int(getattr(config, "MIN_COLUMN_WIDTH_DXA", 500))
            first_row_tr = rows[0]._tr
            physical_tcs = first_row_tr.findall(qn("w:tc"))
            # Collect (tc, width_dxa, gridSpan) for each physical cell
            widths: list[tuple] = []
            for tc in physical_tcs:
                tcPr_el = tc.find(qn("w:tcPr"))
                gs = tcPr_el.find(qn("w:gridSpan")) if tcPr_el is not None else None
                span = int(gs.get(qn("w:val"))) if gs is not None else 1
                if span > 1:
                    # gridSpan safety: skip merged cells in adjustment
                    widths.append((tc, -1, span))
                    continue
                tw = tcPr_el.find(qn("w:tcW")) if tcPr_el is not None else None
                w_val = tw.get(qn("w:w")) if tw is not None else None
                if w_val is not None and w_val != "auto":
                    try:
                        w = int(w_val)
                    except (ValueError, TypeError):
                        w = 0
                else:
                    w = 0
                widths.append((tc, w, span))

            # Expand narrow columns to min_width, compensate from widest
            narrow_indices = [
                i for i, (tc, w, span) in enumerate(widths)
                if w >= 0 and w < min_width
            ]
            if narrow_indices:
                adjustable = [
                    i for i, (tc, w, span) in enumerate(widths)
                    if w >= 0 and w >= min_width
                ]
                for ni in narrow_indices:
                    tc_el, old_w, sp = widths[ni]
                    deficit = min_width - old_w
                    # Find the widest adjustable cell to compensate
                    if adjustable:
                        widest_i = max(adjustable, key=lambda i: widths[i][1])
                        tc_w, w_w, sp_w = widths[widest_i]
                        if w_w - deficit >= min_width:
                            # Expand narrow cell
                            tcPr_el = tc_el.find(qn("w:tcPr"))
                            if tcPr_el is None:
                                tcPr_el = OxmlElement("w:tcPr")
                                tc_el.insert(0, tcPr_el)
                            tw_el = tcPr_el.find(qn("w:tcW"))
                            if tw_el is None:
                                tw_el = OxmlElement("w:tcW")
                                tcPr_el.append(tw_el)
                            tw_el.set(qn("w:w"), str(min_width))
                            tw_el.set(qn("w:type"), "dxa")
                            # Shrink widest cell
                            tc_w2, w_w2, sp_w2 = widths[widest_i]
                            new_widest = w_w2 - deficit
                            tcPr_w = tc_w2.find(qn("w:tcPr"))
                            if tcPr_w is not None:
                                tw_w = tcPr_w.find(qn("w:tcW"))
                                if tw_w is not None:
                                    tw_w.set(qn("w:w"), str(new_widest))
                            # Update tracked widths
                            widths[ni] = (tc_el, min_width, sp)
                            widths[widest_i] = (tc_w2, new_widest, sp_w2)
        except Exception:
            pass

        # ---- noWrap removal + vAlign=top (always applied) ----
        for row in rows:
            seen_tc_ids: set = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)
                try:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                except Exception:
                    pass
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                for nw in tcPr.findall(qn("w:noWrap")):
                    tcPr.remove(nw)
                for child in list(tcPr):
                    if child.tag == qn("w:vAlign"):
                        tcPr.remove(child)
                valign = OxmlElement("w:vAlign")
                valign.set(qn("w:val"), "top")
                tcPr.append(valign)
