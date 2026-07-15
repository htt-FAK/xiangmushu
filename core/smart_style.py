"""智能样式回填模块（三级优先级合并 + 段落样式应用）。

核心功能（doc-gen-revamp Group 5）：
    - merge_style(profile, overrides) → 三级合并后的 TemplateStyleProfile
    - apply_smart_style(para, merged_profile, heading_detector) → 给 run.rPr 注入样式
    - apply_smart_style_to_document(doc, merged_profile) → 整文档逐段应用

三级优先级：user_overrides > TemplateStyleProfile > SystemDefaults

与现有 `docx_typography.py` 的关系：
    - `APPLY_TEMPLATE_STYLE=True`（默认）时使用本模块
    - `APPLY_TEMPLATE_STYLE=False` 时 fallback 到 `apply_document_typography()`（旧逻辑）

依赖：
    - core.style_models (RunStyle, TemplateStyleProfile, system_default_profile)
    - core.normal_heading_detector (classify_heading)
    - core.docx_typography (heading_level_from_style, is_cover_paragraph)
"""
from __future__ import annotations

import logging
from copy import deepcopy
from typing import Any, Optional, TYPE_CHECKING

import config
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.docx_typography import (
    heading_level_from_style,
    is_cover_paragraph,
    is_keyword_line,
)
from core.style_models import (
    RunStyle,
    TemplateStyleProfile,
    system_default_profile,
)

if TYPE_CHECKING:
    from docx import Document

_LOG = logging.getLogger(__name__)


# ── 三级合并 ──────────────────────────────────────────────────────────────
def merge_style(
    template_profile: Optional[TemplateStyleProfile],
    user_overrides: Optional[dict[str, Any]] = None,
) -> TemplateStyleProfile:
    """三级样式优先级合并，返回新 TemplateStyleProfile。

    优先级：user_overrides > template_profile > system_default_profile()

    Args:
        template_profile: 从模板提取的样式档案；None 时使用 system defaults
        user_overrides: 用户 API 传入的覆盖项（dict 或 FormatOverrides.to_merge_dict()）

    Returns:
        新 TemplateStyleProfile（不修改原始输入）
    """
    base = template_profile if template_profile is not None else system_default_profile()
    merged = base.merge_user_overrides(user_overrides)
    _LOG.debug(
        "merge_style: body=%s/%s/%.1fpt, %d heading levels, overrides=%s",
        merged.body_style.font_ascii,
        merged.body_style.font_east_asia,
        merged.body_style.size_pt,
        len(merged.heading_styles),
        list((user_overrides or {}).keys()),
    )
    return merged


# ── 样式 XML 构造 ──────────────────────────────────────────────────────
def _build_rPr_from_runstyle(rs: RunStyle) -> OxmlElement:
    """根据 RunStyle 构造 w:rPr XML 元素（与原 _make_rPr 等价但支持更多字段）。"""
    rpr = OxmlElement("w:rPr")

    # 字体：rFonts
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), rs.font_ascii or "SimSun")
    rf.set(qn("w:hAnsi"), rs.font_ascii or "SimSun")
    rf.set(qn("w:eastAsia"), rs.font_east_asia or "宋体")
    rpr.append(rf)

    # 字号（半磅值）
    sz_half = rs.half_point_size()
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(sz_half))
    rpr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(sz_half))
    rpr.append(sz_cs)

    # 粗体
    if rs.bold:
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))

    # 斜体
    if rs.italic:
        rpr.append(OxmlElement("w:i"))
        rpr.append(OxmlElement("w:iCs"))

    # 颜色
    if rs.color_rgb:
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), rs.color_rgb.lstrip("#"))
        rpr.append(color_el)

    return rpr


def _apply_rPr_to_run(run, rpr: OxmlElement) -> None:
    """将 rPr 应用到 run，替换现有 rPr。"""
    el = run._r
    if el.rPr is not None:
        el.remove(el.rPr)
    el.insert(0, deepcopy(rpr))


# ── 段落级样式应用 ──────────────────────────────────────────────────────
def apply_smart_style(
    para: Paragraph,
    merged_profile: TemplateStyleProfile,
    heading_detector: Optional[Any] = None,
    is_table_para: bool = False,
) -> None:
    """应用样式到段落的所有 run。

    决策优先级：
        1. para.style.name → heading_level_from_style（Heading style）
        2. 封面/关键词段落 → 保留原格式
        3. normal_heading_detector.classify_heading（Normal-style 标题检测）
        4. 默认使用 body_style

    Args:
        para: 目标段落
        merged_profile: 三级合并后的 TemplateStyleProfile
        heading_detector: 可选，用于 Normal-heading 检测
        is_table_para: 是否表格内段落（影响默认样式选择）
    """
    # 封面段落保护
    if is_cover_paragraph(para):
        return

    # 关键词行保护（不应用正文缩进/行距）
    text = (para.text or "").strip()
    if is_keyword_line(text):
        return

    # 空段落跳过
    if not text and not para.runs:
        return

    # 选择 RunStyle
    style_name = para.style.name if para.style else ""

    # 1. Heading style 优先
    lvl = heading_level_from_style(style_name)
    if lvl is not None:
        rs = merged_profile.heading_style_for_level(lvl)
    else:
        # 2. Normal-heading 检测
        normal_level: Optional[int] = None
        if heading_detector is not None:
            try:
                normal_level = heading_detector(para)
            except Exception:
                pass
        if normal_level is not None:
            rs = merged_profile.heading_style_for_level(normal_level)
        elif is_table_para:
            rs = merged_profile.table_cell_style
        else:
            rs = merged_profile.body_style

    # 构造 rPr 并应用到每个 run
    rpr = _build_rPr_from_runstyle(rs)
    for run in para.runs:
        _apply_rPr_to_run(run, rpr)


# ── 整文档样式应用 ──────────────────────────────────────────────────────
def apply_smart_style_to_document(
    doc: "Document",
    merged_profile: TemplateStyleProfile,
) -> None:
    """对所有段落逐段应用 smart_style（替代 apply_document_typography 的新入口）。

    保留：
    - 封面段落保护
    - 关键词行保护
    - 评分表/封面表格跳过
    """
    from core.docx_typography import apply_body_first_line_indent

    # 准备普通段落用的 heading_detector 包装（避免重复 import / 重复扫描）
    threshold = int(getattr(config, "NORMAL_HEADING_THRESHOLD", 50))

    def _detect_heading(para: Paragraph) -> Optional[int]:
        try:
            from core.normal_heading_detector import classify_heading
            return classify_heading(para, threshold=threshold)
        except Exception:
            return None

    # WordFiller 的封面表/评分表判断（懒加载避免循环导入）
    def _is_protected_table(table) -> bool:
        try:
            from core.filler import WordFiller
            return WordFiller._is_cover_table(table) or WordFiller._is_rating_table(table)
        except Exception:
            return False

    # 正文段落
    for para in doc.paragraphs:
        apply_smart_style(para, merged_profile, heading_detector=_detect_heading)

    # 表格单元格（保留已显式设置的 run rPr）
    for table in doc.tables:
        if _is_protected_table(table):
            continue
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.runs and not (para.text or "").strip():
                        continue
                    apply_smart_style(
                        para,
                        merged_profile,
                        heading_detector=_detect_heading,
                        is_table_para=True,
                    )

    # 首行缩进（仍用现有函数，但用 merged_profile.first_line_indent_pt）
    # 临时覆盖 config 值再调用，避免改 docx_typography 接口
    original_indent = getattr(config, "BODY_FIRST_LINE_INDENT_PT", 24.0)
    try:
        config.BODY_FIRST_LINE_INDENT_PT = merged_profile.first_line_indent_pt
        apply_body_first_line_indent(doc)
    finally:
        config.BODY_FIRST_LINE_INDENT_PT = original_indent


def should_use_smart_style() -> bool:
    """是否启用智能样式模式（由 APPLY_TEMPLATE_STYLE 配置决定）。"""
    return bool(getattr(config, "APPLY_TEMPLATE_STYLE", True))
