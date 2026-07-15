"""从 .docx 模板提取 TemplateStyleProfile。

扫描正文段落、标题、表格，提取字体/字号/列宽等样式信息，
并支持 SHA-256 文件缓存避免重复提取。
"""
from __future__ import annotations

import hashlib
import logging
import os
from collections import Counter
from datetime import datetime, timezone
from typing import Optional, TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn

from core.style_models import RunStyle, TemplateStyleProfile
from core import normal_heading_detector

if TYPE_CHECKING:
    from docx.styles.style import BaseStyle
    from lxml.etree import _Element

_LOG = logging.getLogger(__name__)


# ── Cache helpers ────────────────────────────────────────────────────────────

def _cache_dir() -> str:
    d = os.path.join(os.path.dirname(__file__), "..", "data", ".cache", "template_styles")
    abs_path = os.path.abspath(d)
    os.makedirs(abs_path, exist_ok=True)
    return abs_path


def _cache_key(docx_path: str) -> str:
    ap = os.path.abspath(docx_path)
    try:
        mt = int(os.path.getmtime(ap))
    except OSError:
        mt = 0
    raw = f"{ap}|{mt}".encode("utf-8", errors="replace")
    return hashlib.sha256(raw).hexdigest()[:40]


def _cache_path(docx_path: str) -> str:
    return os.path.join(_cache_dir(), f"{_cache_key(docx_path)}.json")


# ── RunStyle extraction from a single run ────────────────────────────────────

def _run_style_from_run(run) -> RunStyle:
    """Extract RunStyle from a python-docx Run, with eastAsia fallback."""
    font_ascii: Optional[str] = run.font.name
    font_east_asia: Optional[str] = None
    size_pt: Optional[float] = None

    # eastAsia fallback via XML
    try:
        r_pr = run._r.find(qn("w:rPr"))
        if r_pr is not None:
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is not None:
                font_east_asia = r_fonts.get(qn("w:eastAsia"))
                # If ascii not set via python-docx, try XML
                if not font_ascii:
                    font_ascii = r_fonts.get(qn("w:ascii"))
    except Exception:
        _LOG.debug("Failed to read rPr XML for run", exc_info=True)

    # Size
    if run.font.size is not None:
        size_pt = run.font.size.pt
    else:
        # Try XML fallback for size
        try:
            r_pr = run._r.find(qn("w:rPr"))
            if r_pr is not None:
                sz = r_pr.find(qn("w:sz"))
                if sz is not None:
                    half_pt = sz.get(qn("w:val"))
                    if half_pt:
                        size_pt = int(half_pt) / 2.0
        except Exception:
            pass

    bold = run.font.bold
    italic = run.font.italic

    color_rgb: Optional[str] = None
    try:
        if run.font.color and run.font.color.rgb:
            color_rgb = str(run.font.color.rgb)
    except Exception:
        pass

    return RunStyle(
        font_ascii=font_ascii or "SimSun",
        font_east_asia=font_east_asia or "宋体",
        size_pt=size_pt if size_pt and size_pt > 0 else 12.0,
        bold=bold if bold is not None else False,
        italic=italic if italic is not None else False,
        color_rgb=color_rgb,
    )


# ── Style chain resolution ───────────────────────────────────────────────────

def _resolve_style_chain(style, styles_xml_element) -> RunStyle:
    """Follow basedOn chain to merge inherited attributes into a RunStyle.

    Args:
        style: A python-docx Style object.
        styles_xml_element: The document's styles XML element (unused in
            simplified implementation; kept for API compatibility).

    Returns:
        Fully-resolved RunStyle with explicit attributes overriding inherited ones.
    """
    # Collect the chain from leaf to root
    chain: list = []
    current = style
    seen_ids: set[str] = set()
    while current is not None:
        style_id = getattr(current, "style_id", None) or getattr(current, "name", "") or id(current)
        sid = str(style_id)
        if sid in seen_ids:
            break  # cycle guard
        seen_ids.add(sid)
        chain.append(current)
        try:
            current = current.base_style
        except Exception:
            current = None

    # Build from root (most ancestral) down to leaf
    chain.reverse()

    base = RunStyle()
    for s in chain:
        try:
            font = getattr(s, "font", None)
            if font is None:
                continue

            vals = {}
            if font.name:
                vals["font_ascii"] = font.name

            # eastAsia via XML
            try:
                elem = s.element
                if elem is not None:
                    r_pr = elem.find(qn("w:rPr"))
                    if r_pr is not None:
                        r_fonts = r_pr.find(qn("w:rFonts"))
                        if r_fonts is not None:
                            ea = r_fonts.get(qn("w:eastAsia"))
                            if ea:
                                vals["font_east_asia"] = ea

                    # Size via XML
                    if r_pr is not None:
                        sz = r_pr.find(qn("w:sz"))
                        if sz is not None:
                            half_pt = sz.get(qn("w:val"))
                            if half_pt:
                                vals["size_pt"] = int(half_pt) / 2.0
            except Exception:
                pass

            if font.size is not None:
                vals["size_pt"] = font.size.pt

            if font.bold is not None:
                vals["bold"] = font.bold
            if font.italic is not None:
                vals["italic"] = font.italic

            if vals:
                override = RunStyle(
                    font_ascii=vals.get("font_ascii", ""),
                    font_east_asia=vals.get("font_east_asia", ""),
                    size_pt=vals.get("size_pt", 0),
                    bold=vals.get("bold"),
                    italic=vals.get("italic"),
                )
                base = base.merge(override)
        except Exception:
            _LOG.debug("Error resolving style %r", getattr(s, "name", "?"), exc_info=True)
            continue

    return base


# ── Column widths extraction ─────────────────────────────────────────────────

def _extract_column_widths(doc: Document) -> dict[int, list[int]]:
    """Extract physical cell widths from each table's first row.

    Uses XML-level access to avoid merged-cell deduplication by row.cells.

    Returns:
        {table_index: [width_dxa, ...]}
    """
    result: dict[int, list[int]] = {}

    for tbl_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        first_row = table.rows[0]
        # Physical cells via XML (not row.cells which deduplicates merged)
        try:
            tc_elements = first_row._tr.findall(qn("w:tc"))
        except Exception:
            _LOG.debug("Failed to read tc elements for table %d", tbl_idx, exc_info=True)
            continue

        widths: list[int] = []
        for tc in tc_elements:
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None:
                    w_val = tc_w.get(qn("w:w"))
                    if w_val:
                        try:
                            widths.append(int(w_val))
                        except ValueError:
                            widths.append(0)
                        continue
            widths.append(0)

        if widths:
            result[tbl_idx] = widths
            _LOG.debug("Table %d column widths: %s", tbl_idx, widths)

    return result


# ── Heading style detection ──────────────────────────────────────────────────

def _detect_heading_styles(
    doc: Document,
    headings_list: list[tuple[int, int]],
) -> dict[int, RunStyle]:
    """Sample RunStyle from each detected heading paragraph.

    Args:
        doc: The python-docx Document.
        headings_list: Output of normal_heading_detector.find_all_headings(doc).

    Returns:
        {level: RunStyle} via majority vote per level.
    """
    # Collect RunStyle samples per level
    level_samples: dict[int, list[RunStyle]] = {}

    for para_idx, level in headings_list:
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]
        # Find first non-empty run
        sample_run = None
        for run in para.runs:
            if (run.text or "").strip():
                sample_run = run
                break
        if sample_run is None and para.runs:
            sample_run = para.runs[0]
        if sample_run is None:
            continue

        rs = _run_style_from_run(sample_run)
        level_samples.setdefault(level, []).append(rs)

    # Majority vote per level
    result: dict[int, RunStyle] = {}
    for level, samples in level_samples.items():
        if not samples:
            continue
        # Vote on (font_ascii, font_east_asia, size_pt, bold)
        votes: Counter = Counter()
        for s in samples:
            key = (s.font_ascii, s.font_east_asia, s.size_pt, s.bold)
            votes[key] += 1
        winner = votes.most_common(1)[0][0]
        result[level] = RunStyle(
            font_ascii=winner[0],
            font_east_asia=winner[1],
            size_pt=winner[2],
            bold=winner[3],
        )

    return result


# ── Body style extraction ────────────────────────────────────────────────────

def _extract_body_style(doc: Document) -> RunStyle:
    """Scan body paragraphs for the most common RunStyle (majority vote).

    Excludes heading paragraphs and cover-style paragraphs (very large fonts).
    """
    # Detect heading para indices to exclude
    headings = normal_heading_detector.find_all_headings(doc)
    heading_indices = {idx for idx, _ in headings}

    samples: list[RunStyle] = []
    for i, para in enumerate(doc.paragraphs):
        if i in heading_indices:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        # Skip cover paragraphs (very large font)
        first_run = None
        for run in para.runs:
            if (run.text or "").strip():
                first_run = run
                break
        if first_run is None:
            continue

        # Skip cover-style (font > 28pt)
        if first_run.font.size is not None and first_run.font.size.pt > 28:
            continue

        rs = _run_style_from_run(first_run)
        samples.append(rs)

    if not samples:
        _LOG.warning("No body samples found; using default RunStyle")
        return RunStyle()

    # Majority vote on (font_ascii, font_east_asia, size_pt)
    votes: Counter = Counter()
    for s in samples:
        key = (s.font_ascii, s.font_east_asia, s.size_pt)
        votes[key] += 1
    winner = votes.most_common(1)[0][0]

    return RunStyle(
        font_ascii=winner[0],
        font_east_asia=winner[1],
        size_pt=winner[2],
    )


# ── Table cell styles ────────────────────────────────────────────────────────

def _extract_table_cell_style(doc: Document) -> RunStyle:
    """Extract typical table cell RunStyle from first table's first data row."""
    if not doc.tables:
        return RunStyle()

    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = (para.text or "").strip()
                if not text:
                    continue
                for run in para.runs:
                    if (run.text or "").strip():
                        return _run_style_from_run(run)
    return RunStyle()


def _extract_table_label_style(doc: Document) -> RunStyle:
    """Extract table label column style (col 0 of label-value tables)."""
    if not doc.tables:
        return RunStyle(bold=True)

    table = doc.tables[0]
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        first_cell = cells[0]
        for para in first_cell.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            for run in para.runs:
                if (run.text or "").strip():
                    rs = _run_style_from_run(run)
                    # Label is typically bold
                    return RunStyle(
                        font_ascii=rs.font_ascii,
                        font_east_asia=rs.font_east_asia,
                        size_pt=rs.size_pt,
                        bold=True,
                    )
    return RunStyle(bold=True)


# ── Line spacing extraction ──────────────────────────────────────────────────

def _extract_line_spacing(doc: Document) -> float:
    """Extract the most common line spacing from body paragraphs."""
    spacings: Counter = Counter()
    for para in doc.paragraphs:
        pf = para.paragraph_format
        if pf.line_spacing is not None:
            try:
                val = float(pf.line_spacing)
                if 0.5 <= val <= 5.0:
                    spacings[val] += 1
            except (ValueError, TypeError):
                # Exact value in EMU - convert to multiple
                pass
    if spacings:
        return spacings.most_common(1)[0][0]
    return 1.0


# ── Main extraction ──────────────────────────────────────────────────────────

def extract_style_profile(docx_path: str) -> TemplateStyleProfile:
    """Extract a complete TemplateStyleProfile from a .docx template.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Populated TemplateStyleProfile.
    """
    _LOG.info("Extracting style profile from %s", docx_path)
    doc = Document(docx_path)

    body_style = _extract_body_style(doc)
    headings = normal_heading_detector.find_all_headings(doc)
    heading_styles = _detect_heading_styles(doc, headings)
    table_cell_style = _extract_table_cell_style(doc)
    table_label_style = _extract_table_label_style(doc)
    column_widths = _extract_column_widths(doc)
    line_spacing = _extract_line_spacing(doc)

    profile = TemplateStyleProfile(
        body_style=body_style,
        heading_styles=heading_styles,
        table_cell_style=table_cell_style,
        table_label_style=table_label_style,
        column_widths=column_widths,
        line_spacing=line_spacing,
        source_template=os.path.basename(docx_path),
        extracted_at=datetime.now(timezone.utc).isoformat(),
    )
    _LOG.info(
        "Extracted profile: body=%s, headings=%s, tables=%d",
        body_style, list(heading_styles.keys()), len(column_widths),
    )
    return profile


# ── Cached entry point ───────────────────────────────────────────────────────

def get_or_extract_style_profile(docx_path: str) -> TemplateStyleProfile:
    """Check cache first; extract and cache on miss.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        TemplateStyleProfile (from cache or freshly extracted).
    """
    cpath = _cache_path(docx_path)
    if os.path.isfile(cpath):
        try:
            with open(cpath, "r", encoding="utf-8") as f:
                raw = f.read()
            profile = TemplateStyleProfile.from_json(raw)
            _LOG.info("Style profile cache hit: %s", cpath)
            return profile
        except Exception:
            _LOG.warning("Cache read failed, re-extracting", exc_info=True)

    profile = extract_style_profile(docx_path)

    try:
        with open(cpath, "w", encoding="utf-8") as f:
            f.write(profile.to_json())
        _LOG.info("Style profile cached to %s", cpath)
    except Exception:
        _LOG.warning("Failed to write cache", exc_info=True)

    return profile
