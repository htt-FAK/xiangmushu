"""doc-gen-revamp 端到端 showcase 验收脚本。

直接调用新管线核心模块（不走 server / 不依赖 ChromaDB / 不需要 LLM 调用），
模拟一次完整的模板分析和回填流程，验证 7 个核心改进点全部生效。

运行方式：
    python scripts/showcase_pipeline.py

期望：全部 7 个 CHECK 通过；输出 docx 文件供人工查看样式/列宽/字体。
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

# 让导入找到 core/
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "data" / "templates" / "智能体应用开发实践.docx"
ALT_TEMPLATE = (
    Path(__file__).resolve().parents[1] /
    "docs" / "2.1.2024级广东理工学院创新计划书参考模板（通用）.docx"
)


def _ok(label: str, detail: str = "") -> None:
    suffix = f" — {detail}" if detail else ""
    print(f"  [OK]    {label}{suffix}")


def _fail(label: str, detail: str = "") -> None:
    suffix = f" — {detail}" if detail else ""
    print(f"  [FAIL]  {label}{suffix}")


def run_check(label: str, fn) -> bool:
    try:
        ok = fn()
    except Exception as exc:
        _fail(label, f"exception: {exc}")
        import traceback
        traceback.print_exc()
        return False
    # 强制转 bool（防止 fn 返回 set/dict 等 truthy 非 bool）
    ok = bool(ok)
    if ok:
        _ok(label)
    else:
        _fail(label)
    return ok


def main() -> int:
    # 选模板
    template = TEMPLATE_PATH if TEMPLATE_PATH.exists() else ALT_TEMPLATE
    if not template.exists():
        print(f"[FATAL] 找不到模板: {template}")
        return 1
    print(f"=== doc-gen-revamp E2E showcase ===")
    print(f"    模板: {template.name}")
    print()

    out_dir = Path(__file__).resolve().parents[1] / "data" / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    passed = 0
    total = 0

    # ── CHECK 1: Normal-Heading Detection ──────────────────────────────────
    from core.normal_heading_detector import find_all_headings
    def _c1():
        doc = Document(str(template))
        headings = find_all_headings(doc, threshold=50)
        if not headings:
            return False
        # 至少 2 个标题且 level 合理
        levels = {lvl for _, lvl in headings}
        preview = [
            f"L{lvl}: {doc.paragraphs[idx].text.strip()[:25]}"
            for idx, lvl in headings[:5]
        ]
        print(f"          检测到 {len(headings)} 个标题: {preview}")
        return len(headings) >= 2 and levels
    passed += run_check("CHECK 1: Normal-heading detector 识别章节边界", _c1)
    total += 1

    # ── CHECK 2: Template Style Extraction ─────────────────────────────────
    from core.template_style_extractor import get_or_extract_style_profile
    style_profile = None
    def _c2():
        nonlocal style_profile
        style_profile = get_or_extract_style_profile(str(template))
        body = style_profile.body_style
        return (
            body.font_east_asia  # 有字体
            and body.size_pt > 0
            and style_profile.source_template
        )
    passed += run_check(
        "CHECK 2: Style extractor 提取模板真实样式",
        _c2,
    )
    total += 1
    if style_profile:
        b = style_profile.body_style
        print(f"          body_style: {b.font_east_asia or b.font_ascii} {b.size_pt}pt "
              f"bold={b.bold}")
        print(f"          heading levels: {sorted(style_profile.heading_styles.keys())}")
        print(f"          tables analyzed: {len(style_profile.column_widths)}")

    # ── CHECK 3: Table Semantic Classification ─────────────────────────────
    from core.table_semantic_analyzer import analyze_table
    from core.fill_intent import TableSemanticType
    analyses = []
    def _c3():
        doc = Document(str(template))
        for i, table in enumerate(doc.tables):
            analyses.append(analyze_table(table, i, "demo_chapter"))
        if not analyses:
            return False
        # 至少能识别出一种非 UNKNOWN 类型
        type_counts = {}
        for a in analyses:
            type_counts[a.table_type.value] = type_counts.get(a.table_type.value, 0) + 1
        print(f"          表格分类: {type_counts}")
        non_unknown = sum(v for k, v in type_counts.items() if k != "unknown")
        return non_unknown >= 1
    passed += run_check("CHECK 3: Table semantic analyzer 分类表格", _c3)
    total += 1

    # ── CHECK 4: Smart Style Merge (三级优先级) ─────────────────────────────
    from core.smart_style import merge_style
    merged = None
    def _c4():
        nonlocal merged
        from core.style_models import RunStyle
        # 用户 override 字体为黑体 14pt
        overrides = {"body_font_east_asia": "黑体", "body_size_pt": 14.0}
        merged = merge_style(style_profile, overrides)
        return (
            merged.body_style.font_east_asia == "黑体"
            and merged.body_style.size_pt == 14.0
        )
    passed += run_check("CHECK 4: Smart style 三级合并 (user > template > default)", _c4)
    total += 1
    if merged:
        print(f"          merged body: {merged.body_style.font_east_asia} "
              f"{merged.body_style.size_pt}pt")

    # ── CHECK 5: Fill Template with merged profile ──────────────────────────
    output_path = out_dir / f"showcase_{template.stem}_e2e.docx"
    def _c5():
        from core.filler import WordFiller
        from core.fill_task import FillTask
        # 模拟 2 个 FillTask（段落 + 表格），不依赖 LLM
        # 找一个非封面表做内容
        target_chapter = "demo_chapter"
        # 找一个 label-value 表的空单元格
        lv_task = None
        for a in analyses:
            if a.table_type == TableSemanticType.LABEL_VALUE_PAIR:
                for (r, c), intent in a.fill_intents.items():
                    if intent.value == "fill" and r > 0 and c == 1:
                        lv_task = FillTask(
                            task_id="demo_lv",
                            target_chapter=target_chapter,
                            task_type="table_cell",
                            description="demo fill",
                            location_hint={
                                "table_index": a.table_index,
                                "row": r,
                                "col": c,
                                "fill_intent": "fill",
                            },
                            word_limit=60,
                        )
                        break
                if lv_task:
                    break

        tasks = []
        contents = []
        if lv_task:
            tasks.append(lv_task)
            contents.append("这是一段端到端验收文本，验证智能样式管线正确性。")

        filler = WordFiller()
        filler.fill_template(
            str(template), tasks, contents, str(output_path),
            merged_profile=merged,
        )
        return output_path.exists() and output_path.stat().st_size > 1000
    passed += run_check("CHECK 5: Filler 完整回填流程 (smart style enabled)", _c5)
    total += 1
    if output_path.exists():
        print(f"          输出: {output_path.name} ({output_path.stat().st_size} bytes)")

    # ── CHECK 6: Column Width Preservation (正向验证) ─────────────────────
    import config as cfg
    MIN_WIDTH_GUARD = int(getattr(cfg, "MIN_COLUMN_WIDTH_DXA", 500))
    def _c6():
        # 正向验证: 检查输出表格列宽是否合理 (不被强制等宽破坏)
        # (rubric 删除会改变表格总数, 所以不能用 index 对比)
        out_doc = Document(str(output_path))
        orig_doc = Document(str(template))
        # 先验证: 表格数量差在合理范围 (rubric 删除通常 ≤ 3 个)
        delta = abs(len(orig_doc.tables) - len(out_doc.tables))
        if delta > 5:
            print(f"          tables delta 过大: orig={len(orig_doc.tables)} "
                  f"out={len(out_doc.tables)} delta={delta}")
            return False
        # 找输出中第一个 ≥ 2 列的 label_value 形表格, 验证列宽合理性
        n_checked = 0
        for table in out_doc.tables:
            if not table.rows:
                continue
            first_row_tcs = table.rows[0]._tr.findall(qn("w:tc"))
            n_cols = len(first_row_tcs)
            if n_cols < 2:
                continue
            widths = []
            for tc in first_row_tcs:
                tc_pr = tc.find(qn("w:tcPr"))
                if tc_pr is None:
                    continue
                tw = tc_pr.find(qn("w:tcW"))
                if tw is None:
                    continue
                raw = tw.get(qn("w:w"))
                if raw and raw != "auto":
                    try:
                        widths.append(int(raw))
                    except ValueError:
                        pass
            if len(widths) != n_cols:
                continue
            # 验证: 至少有 1 列 ≥ MIN_WIDTH_GUARD; 没有列被强制等分(即全等)
            if not any(w >= MIN_WIDTH_GUARD for w in widths):
                continue
            n_checked += 1
            # 列宽不应全部相同 (强制等分会让所有列相等)
            if len(set(widths)) == 1 and n_cols > 1:
                print(f"          列宽被强制等分: widths={widths}")
                return False
        print(f"          表格数量: orig={len(orig_doc.tables)} out={len(out_doc.tables)} "
              f"(delta={delta}); {n_checked} 个表格列宽合理")
        return True
    passed += run_check("CHECK 6: 回填后表格列宽未被强制等分", _c6)
    total += 1

    # ── CHECK 7: Fill task 写到正确的位置 (正向验证) ───────────────────────
    def _c7():
        # 正向验证: 我们写的 FILL 内容确实在输出文件中 (位置正确)
        # 以及我们没把内容写到 col 0 (典型的 LABEL 位置)
        out_doc = Document(str(output_path))
        FILL_TEXT = "这是一段端到端验收文本，验证智能样式管线正确性。"
        # 在输出文件中搜索我们写入的文本, 找到它在哪个位置
        fill_cells_found = []
        for ti, table in enumerate(out_doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if FILL_TEXT in (cell.text or ""):
                        fill_cells_found.append((ti, ri, ci))
        if not fill_cells_found:
            print("          FILL 内容未在输出中找到")
            return False
        # 验证: 我们写入的位置 col > 0 (LABEL_VALUE_PAIR 的 col 0 是 label)
        for ti, ri, ci in fill_cells_found:
            if ci == 0:
                # 检查这一行的 col 0 是不是 label
                for a in analyses:
                    if a.table_type != TableSemanticType.LABEL_VALUE_PAIR:
                        continue
                    if (ri, 0) in a.fill_intents and a.fill_intents[(ri, 0)].value == "label":
                        print(f"          FILL 内容写到了 LABEL 位置: table[{ti}][{ri},{ci}]")
                        return False
        print(f"          FILL 内容写入位置: {fill_cells_found} (均正确)")
        return True
    passed += run_check("CHECK 7: FILL 内容写到了正确位置", _c7)
    total += 1

    # ── 总结 ──────────────────────────────────────────────────────────────
    print()
    if passed == total:
        print("=== ALL " + str(total) + " CHECKS PASSED ===")
        print(f"    输出示例: {output_path}")
        print(f"    该 .docx 可用 Word 打开，查看样式 / 列宽 / 字体是否正确。")
        return 0
    print(f"=== FAILURES: {passed}/{total} checks passed ===")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
